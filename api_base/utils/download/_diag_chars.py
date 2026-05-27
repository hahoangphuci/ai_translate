# -*- coding: utf-8 -*-
import os
import re
import sys
import tempfile
from collections import Counter

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

import fitz
import docx
from pdf2docx import Converter

from app.services.document_v2.pdf_docx_pipeline.layout_recovery import sanitize_converted_docx


def doc_text(doc):
    parts = []
    for p in doc.paragraphs:
        parts.append("".join(r.text or "" for r in p.runs))
    return "\n".join(parts)


def count_chars(s):
    return {
        "len": len(s),
        "+": s.count("+"),
        "(": s.count("("),
        ")": s.count(")"),
        "en_dash": s.count("\u2013"),
        "pua_f02b": s.count("\uf02b"),
        "bullet": s.count("\u2022"),
    }


pdf_path = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "app", "uploads", "HoangPhuc222683.pdf")
)

# PDF text
doc = fitz.open(pdf_path)
pdf_text = ""
for page in doc:
    pdf_text += page.get_text()

with tempfile.TemporaryDirectory() as td:
    raw_docx = os.path.join(td, "raw.docx")
    san_docx = os.path.join(td, "san.docx")
    cv = Converter(pdf_path)
    cv.convert(raw_docx)
    cv.close()

    raw = docx.Document(raw_docx)
    raw_s = doc_text(raw)

    import shutil

    shutil.copy(raw_docx, san_docx)
    sanitize_converted_docx(san_docx, pdf_path=pdf_path, layout_mode="academic")
    san = docx.Document(san_docx)
    san_s = doc_text(san)

print("=== char counts ===")
for label, s in [("pdf", pdf_text), ("raw_docx", raw_s), ("sanitized", san_s)]:
    print(label, count_chars(s))

# Find lines in PDF missing from sanitized
pdf_lines = [ln.strip() for ln in pdf_text.splitlines() if ln.strip()]
san_lines = [ln.strip() for ln in san_s.splitlines() if ln.strip()]

print("\n=== PDF lines with PUA or method names ===")
for ln in pdf_lines:
    if "\uf02b" in ln or "()" in ln or "capNhat" in ln or "Sach" in ln:
        norm = ln.replace("\uf02b", "+")
        found = any(norm in sl or ln.replace("\uf02b", "+").strip() in sl for sl in san_lines)
        status = "OK" if found else "MISSING"
        print(f"[{status}] {repr(ln[:100])}")

print("\n=== glue samples (missing space) in sanitized ===")
for ln in san_lines:
    if re.search(r"[a-zà-ỹ][A-ZÀ-Ỹ]", ln) or re.search(r"thực thể[a-z]", ln):
        print(repr(ln[:120]))

print("\n=== standalone + lines in PDF vs sanitized ===")
for ln in pdf_lines:
    if ln.strip() in ("+", "\uf02b", "+\uf02b") or re.match(r"^[\uf02b+]\s*$", ln.strip()):
        print("PDF:", repr(ln))
for ln in san_lines:
    if re.match(r"^[+\u2022\u25aa]\s*$", ln.strip()) or ln.strip() == "+":
        print("SAN:", repr(ln))
