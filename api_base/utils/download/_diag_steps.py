# -*- coding: utf-8 -*-
import os
import re
import sys
import tempfile
import shutil

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

import docx
from pdf2docx import Converter

from app.services.document_v2.pdf_docx_pipeline import layout_recovery as lr


def doc_text(doc):
    return "\n".join("".join(r.text or "" for r in p.runs) for p in doc.paragraphs)


def counts(s):
    return len(s), s.count("+"), s.count("("), s.count("\uf02b")


pdf_path = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "app", "uploads", "HoangPhuc222683.pdf")
)

with tempfile.TemporaryDirectory() as td:
    raw_docx = os.path.join(td, "raw.docx")
    cv = Converter(pdf_path)
    cv.convert(raw_docx)
    cv.close()

    path = os.path.join(td, "step.docx")
    shutil.copy(raw_docx, path)
    doc = docx.Document(path)
    print("raw", counts(doc_text(doc)))

    lr.normalize_pdf_symbol_chars_in_doc(doc)
    print("after symbols", counts(doc_text(doc)))

    lr.split_merged_pdf_paragraphs(doc)
    print("after split", counts(doc_text(doc)))

    lr.unify_pdf_docx_fonts(doc)
    lr._normalize_body_run_sizes(doc)
    print("after fonts/sizes", counts(doc_text(doc)))

    n = lr.sync_bold_spans_from_pdf_to_doc(doc, pdf_path)
    print(f"after bold sync (changed={n})", counts(doc_text(doc)))

    noise = lr.strip_noise_paragraphs(doc)
    print(f"after noise {noise}", counts(doc_text(doc)))

    lr.normalize_converted_docx_layout_in_doc(doc, pdf_path=pdf_path, layout_mode="academic")
    print("after layout norm", counts(doc_text(doc)))

    # find paragraphs where bold sync might have changed text
    doc2 = docx.Document(raw_docx)
    before = ["".join(r.text or "" for r in p.runs) for p in doc2.paragraphs]
    lr.normalize_pdf_symbol_chars_in_doc(doc2)
    lr.split_merged_pdf_paragraphs(doc2)
    lr.unify_pdf_docx_fonts(doc2)
    lr._normalize_body_run_sizes(doc2)
    pre_bold = ["".join(r.text or "" for r in p.runs) for p in doc2.paragraphs]
    lr.sync_bold_spans_from_pdf_to_doc(doc2, pdf_path)
    post_bold = ["".join(r.text or "" for r in p.runs) for p in doc2.paragraphs]
    out = open(os.path.join(os.path.dirname(__file__), "_diag_step_out.txt"), "w", encoding="utf-8")
    for i, (a, b) in enumerate(zip(pre_bold, post_bold)):
        if a != b:
            out.write(f"--- para {i} ---\nBEFORE: {a!r}\nAFTER:  {b!r}\n\n")
    out.close()
    print("wrote text diffs from bold sync")
