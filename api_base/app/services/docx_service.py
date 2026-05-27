"""DOCX translation service — extracted from FileService._process_docx.

This module contains the full DOCX translation pipeline including:
- Bilingual modes (inline, newline)
- OCR image handling (overlay, text insertion, auto mode)
- Format-group translation preserving per-run styling
- Table, header/footer, TOC processing
- Post-processing fixes (hyperlinks, leaders, phrase fixes)
"""

import os
import re
import unicodedata
import uuid
import io
import zipfile
import docx
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, RGBColor, Pt
from .file_service import ProviderRateLimitError


def _docx_run_boundary_needs_space(prev: str, next_t: str) -> bool:
    """Heuristic: pdf2docx often emits word-sized <w:r> without spaces between them."""
    ps = (prev or "").rstrip()
    ns = (next_t or "").lstrip()
    if not ps or not ns:
        return False
    if ps.endswith((" ", "\n", "\t")) or ns.startswith((" ", "\n", "\t")):
        return False
    pc = ps[-1]
    nc = ns[0]
    if ps.endswith(("-", "'", "\u2019")):
        return False
    # Comma / clause punctuation → next token (2025,bao | giá.Kết split across runs)
    if pc in ",.;:!?":
        if nc.isalnum():
            if pc in ",." and len(ps) >= 2 and ps[-2].isdigit() and nc.isdigit():
                return False
            return True
        return False
    if pc in ")]}" and nc.isalnum():
        return True
    if pc.isdigit() and nc.isdigit():
        return False
    if pc.isalnum() and nc.isalnum():
        return True
    return False


def _intrarun_insert_missing_spaces(text: str) -> str:
    """Inside a single run: add space after punctuation if pdf2docx glued (e.g. 'giá.Kết', '2025,bao')."""
    if not text:
        return text
    out = []
    n = len(text)
    for i, c in enumerate(text):
        out.append(c)
        if i + 1 >= n:
            continue
        nxt = text[i + 1]
        if nxt in " \n\t\r":
            continue
        prev_ch = text[i - 1] if i > 0 else ""
        if c in ",;:":
            if c == "," and prev_ch.isdigit() and nxt.isdigit():
                continue
            if nxt.isalnum():
                out.append(" ")
        elif c == ".":
            if prev_ch.isdigit() and nxt.isdigit():
                continue
            if nxt.isalpha() and unicodedata.category(nxt) == "Lu":
                out.append(" ")
        elif c in "!?":
            if nxt.isalnum():
                out.append(" ")
    return "".join(out)


def _fix_vn_word_glue(text: str) -> str:
    """Insert space when pdf2docx glued a Vietnamese particle to the previous word."""
    if not text:
        return text
    return re.sub(
        r"([a-zà-ỹđ])(?=(?:và|của|cho|trong|theo|với|hoặc|hay|được|bởi)\b)",
        r"\1 ",
        text,
        flags=re.IGNORECASE,
    )


def _fix_section_number_spacing(text: str) -> str:
    if not text:
        return text
    return re.sub(r"(?<=\d\.)([A-Za-zÀ-ỹ])", r" \1", text)


def apply_docx_paragraph_spacing(paragraph) -> None:
    """Mutate runs only (preserve bold per run); fixes pdf2docx glue + punctuation."""
    try:
        runs = list(paragraph.runs)
    except Exception:
        return
    if not runs:
        return
    for i in range(len(runs) - 1):
        try:
            a = runs[i].text or ""
            b = runs[i + 1].text or ""
            if _docx_run_boundary_needs_space(a, b):
                runs[i].text = a + " "
        except Exception:
            continue
    for r in runs:
        try:
            t = r.text or ""
            nt = _fix_section_number_spacing(
                _fix_vn_word_glue(_intrarun_insert_missing_spaces(t))
            )
            if nt != t:
                r.text = nt
        except Exception:
            continue


def merged_paragraph_plain(paragraph) -> str:
    """Plain text as seen after spacing heuristics (for translation + cache keys)."""
    try:
        texts = [r.text or "" for r in paragraph.runs]
    except Exception:
        return ""
    return _intrarun_insert_missing_spaces(join_docx_run_texts(texts))


def join_docx_run_texts(run_texts):
    """Join DOCX run strings; infer spaces between runs when pdf2docx glued words."""
    seq = [t or "" for t in (run_texts or [])]
    if not seq:
        return ""
    parts = [seq[0]]
    for t in seq[1:]:
        prev = parts[-1]
        if _docx_run_boundary_needs_space(prev, t):
            parts.append(" ")
        parts.append(t)
    return "".join(parts)


def paragraph_plain_literal(paragraph) -> str:
    """Join run text exactly as stored — no inferred spaces or punctuation fixes."""
    try:
        return "".join(r.text or "" for r in paragraph.runs)
    except Exception:
        return ""


def collapse_pdf_prose_paragraph_runs(paragraph) -> bool:
    """Merge pdf2docx word-sized runs into one run (fixes justify spacing gaps)."""
    try:
        runs = [r for r in paragraph.runs if (r.text or "")]
    except Exception:
        return False
    if len(runs) <= 1:
        return False
    text = merged_paragraph_plain(paragraph)
    if not text.strip():
        return False
    template_rpr = None
    try:
        from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
            pick_translation_template_rpr,
        )
        template_rpr = pick_translation_template_rpr(paragraph)
    except Exception:
        template_rpr = None
    if template_rpr is None:
        from docx.oxml.ns import qn as _qn
        for run in runs:
            try:
                rpr = run._element.find(_qn("w:rPr"))
                if rpr is not None:
                    import copy as _copy
                    template_rpr = _copy.deepcopy(rpr)
                    break
            except Exception:
                continue
    for run in list(paragraph.runs):
        try:
            run.text = ""
        except Exception:
            pass
    primary = runs[0]
    primary.text = text
    if template_rpr is not None:
        try:
            import copy as _copy
            from docx.oxml.ns import qn as _qn
            old = primary._element.find(_qn("w:rPr"))
            if old is not None:
                primary._element.remove(old)
            primary._element.insert(0, _copy.deepcopy(template_rpr))
        except Exception:
            pass
    for run in list(paragraph.runs)[1:]:
        try:
            if not (run.text or "").strip():
                run._element.getparent().remove(run._element)
        except Exception:
            pass
    return True


def _apply_translation_with_linebreaks(paragraph, translated_text) -> None:
    """Write translated text back preserving \\n as Word line breaks (w:br)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import copy as _copy

    text = "" if translated_text is None else str(translated_text)
    runs = list(paragraph.runs)

    template_rpr = None
    try:
        from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
            pick_translation_template_rpr,
        )
        template_rpr = pick_translation_template_rpr(paragraph)
    except Exception:
        template_rpr = None

    if template_rpr is None:
        for r in runs:
            try:
                rpr = r._element.find(qn("w:rPr"))
                if rpr is not None:
                    template_rpr = _copy.deepcopy(rpr)
                    break
            except Exception:
                continue

    for r in runs:
        try:
            r.text = ""
        except Exception:
            pass

    def _apply_rpr(run) -> None:
        if template_rpr is None:
            return
        try:
            old = run._element.find(qn("w:rPr"))
            if old is not None:
                run._element.remove(old)
            run._element.insert(0, _copy.deepcopy(template_rpr))
        except Exception:
            pass

    chunks = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not chunks:
        return

    first_run = runs[0] if runs else paragraph.add_run("")
    first_run.text = chunks[0]
    _apply_rpr(first_run)

    for chunk in chunks[1:]:
        br_run = paragraph.add_run("")
        _apply_rpr(br_run)
        br_run.add_break()
        if chunk:
            content_run = paragraph.add_run(chunk)
            _apply_rpr(content_run)


def process_docx(service, file_path, target_lang, progress_callback=None, *, ocr_images=False, ocr_langs=None, ocr_mode=None, bilingual_mode=None, bilingual_delimiter=None, from_pdf=False, pdf_layout_mode=None):
    """Translate DOCX while preserving original formatting, layout, images.
    
    Bilingual modes:
      - none: normal translation (replace original with translation)
      - inline: song ngữ liền kề (Original | Translated in same paragraph)
      - newline: song ngữ xuống dòng (keep original, add translated paragraph below)
      - preserve_layout: alias for 'inline' mode (dịch song ngữ liền kề, giữ layout)
      - line_by_line: alias for 'newline' mode (dịch song ngữ xuống dòng)
    """
    # python-docx rejects macro-enabled documents (.docm content type).
    # Strip the macros by re-packaging as a standard .docx before opening.
    try:
        doc = docx.Document(file_path)
    except ValueError:
        import tempfile, zipfile as _zf
        _tmp = tempfile.mktemp(suffix='.docx')
        with _zf.ZipFile(file_path, 'r') as zin, _zf.ZipFile(_tmp, 'w') as zout:
            for item in zin.infolist():
                # Skip VBA / macro parts entirely
                if 'vbaProject' in item.filename or 'vbaData' in item.filename:
                    continue
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    data = data.replace(
                        b'application/vnd.ms-word.document.macroEnabled.main+xml',
                        b'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml',
                    )
                # Remove relationship entries that reference VBA parts
                if item.filename.endswith('.rels'):
                    data = re.sub(rb'<Relationship[^>]*Target="[^"]*vba[^"]*"[^/]*/>', b'', data, flags=re.IGNORECASE)
                zout.writestr(item, data)
        doc = docx.Document(_tmp)
        os.remove(_tmp)

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy
        pass
    except Exception:
        pass

    api_only = str(os.getenv('AI_DISABLE_FALLBACK', '0')).strip().lower() in ('1', 'true', 'yes', 'on')

    try:
        from app.services.document_v2.pdf_docx_pipeline.layout_recovery import normalize_bilingual_mode
    except Exception:
        def normalize_bilingual_mode(mode):
            bi = (str(mode or "").strip().lower() or "none")
            if bi in ("preserve_layout", "inline"):
                return "inline"
            if bi in ("line_by_line", "newline"):
                return "newline"
            return "none"

    bi_mode = normalize_bilingual_mode(bilingual_mode)

    try:
        from app.services.document_v2.pdf_docx_pipeline.layout_recovery import should_preserve_pdf_lines
    except Exception:
        def should_preserve_pdf_lines(_layout_mode=None):
            return str(os.getenv("PDF_DOCX_PRESERVE_LINES", "1")).strip().lower() not in ("0", "false", "no", "off")

    preserve_pdf_lines = bool(from_pdf) and should_preserve_pdf_lines(pdf_layout_mode)
    is_academic_pdf = bool(from_pdf) and str(pdf_layout_mode or "").strip().lower() == "academic"

    def _paragraph_source_text(paragraph) -> str:
        if bool(from_pdf) and (is_academic_pdf or not preserve_pdf_lines):
            return merged_paragraph_plain(paragraph)
        if preserve_pdf_lines:
            return paragraph_plain_literal(paragraph)
        return merged_paragraph_plain(paragraph)

    mode = (str(ocr_mode).strip().lower() if ocr_mode else 'image')
    if mode not in ('image', 'text', 'both', 'auto'):
        mode = 'auto'

    def _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode=None):
        try:
            raw = (ocr_text or '').strip()
            if not raw:
                return 'text'

            char_count = len(raw)
            words = re.findall(r'\w+', raw, flags=re.UNICODE)
            word_count = len(words)

            if char_count >= 120 or word_count >= 25:
                print(f"  [MODE] Prose detected (chars={char_count}, words={word_count}), AI={ai_recommended_mode} -> text")
                return 'text'

            lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
            low_raw = raw.lower()
            low_trans = (translated_text or '').lower()
            promo_keywords = (
                'sale', 'discount', 'offer', 'book now', 'vacation',
                'summer', 'up to', '% off', 'promo', 'hotline',
                'free', 'limited', 'special', 'deal', 'subscribe',
            )
            has_promo = any(k in low_raw or k in low_trans for k in promo_keywords)

            alpha_chars = [ch for ch in raw if ch.isalpha()]
            upper_ratio = (
                sum(1 for ch in alpha_chars if ch == ch.upper()) / max(1, len(alpha_chars))
                if alpha_chars else 0.0
            )
            line_word_counts = [len(re.findall(r'\w+', ln, flags=re.UNICODE)) for ln in lines] if lines else [0]
            avg_wpl = (sum(line_word_counts) / len(line_word_counts)) if line_word_counts else 0.0
            short_lines = sum(1 for c in line_word_counts if c <= 3)

            looks_banner = (
                has_promo or
                (upper_ratio >= 0.50 and avg_wpl <= 4) or
                (short_lines >= 3 and avg_wpl <= 3)
            )

            ai_mode = (ai_recommended_mode or '').lower()

            if looks_banner or ai_mode == 'image':
                final = 'image'
            else:
                final = 'text'

            print(
                f"  [MODE] AI={ai_mode}, banner={looks_banner}, "
                f"chars={char_count}, words={word_count}, upper={upper_ratio:.2f} -> {final}"
            )
            return final
        except Exception:
            return 'text'

    def iter_all_paragraphs(document):
        paras = []
        try:
            paras.extend(list(document.paragraphs))
        except Exception:
            pass
        try:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paras.extend(list(cell.paragraphs))
        except Exception:
            pass
        try:
            for section in document.sections:
                paras.extend(list(section.header.paragraphs))
                paras.extend(list(section.footer.paragraphs))
        except Exception:
            pass
        return paras

    def paragraph_image_rids(paragraph):
        rids = []
        try:
            runs = list(paragraph.runs)
        except Exception:
            runs = []
        if not runs:
            return rids

        rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        for run in runs:
            try:
                blips = run._element.xpath('.//*[local-name()="blip"]')
            except Exception:
                blips = []
            for blip in blips:
                try:
                    rid = blip.get(rel_attr)
                except Exception:
                    rid = None
                if rid:
                    rids.append(rid)
        seen = set()
        out = []
        for rid in rids:
            if rid in seen:
                continue
            seen.add(rid)
            out.append(rid)
        return out

    def rid_to_image_part(paragraph, rid):
        try:
            part = paragraph.part
            related = getattr(part, 'related_parts', None)
            if isinstance(related, dict) and rid in related:
                return related[rid]
        except Exception:
            pass
        try:
            rels = getattr(paragraph.part, 'rels', None)
            if rels and rid in rels:
                return rels[rid].target_part
        except Exception:
            pass
        return None

    def _collect_header_footer_image_partnames(document):
        protected = set()
        try:
            for section in document.sections:
                for hf in (section.header, section.footer):
                    part = getattr(hf, 'part', None)
                    related = getattr(part, 'related_parts', None)
                    if not isinstance(related, dict):
                        continue
                    for _rid, target in related.items():
                        try:
                            ct = str(getattr(target, 'content_type', '') or '').lower()
                            if not ct.startswith('image/'):
                                continue
                            pn = str(getattr(target, 'partname', '') or '').lstrip('/')
                            if pn:
                                protected.add(pn)
                        except Exception:
                            continue
        except Exception:
            pass
        return protected

    def replace_image_with_text(paragraph, rid, translated_text):
        txt = (translated_text or '').strip()
        if not txt:
            return False

        rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
        replaced = False
        try:
            runs = list(paragraph.runs)
        except Exception:
            runs = []

        for run in runs:
            try:
                blips = run._element.xpath('.//*[local-name()="blip"]')
            except Exception:
                blips = []
            if not blips:
                continue

            has_target = False
            for blip in blips:
                try:
                    if blip.get(rel_attr) == rid:
                        has_target = True
                        break
                except Exception:
                    continue
            if not has_target:
                continue

            try:
                drawings = run._element.xpath('./*[local-name()="drawing"]')
                for dr in drawings:
                    parent = dr.getparent()
                    if parent is not None:
                        parent.remove(dr)
            except Exception:
                pass
            run.text = ""
            replaced = True
            break

        try:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        except Exception:
            pass

        try:
            new_run = paragraph.add_run(txt)
            _ = new_run
            replaced = True
        except Exception:
            pass

        if not replaced:
            try:
                paragraph.add_run(txt)
                replaced = True
            except Exception:
                replaced = False
        return replaced

    def _normalize_ocr_text_for_docx(text):
        raw = (text or '').replace('\r\n', '\n').replace('\r', '\n')
        if not raw.strip():
            return ''

        cleaned_lines = []
        for ln in raw.split('\n'):
            ln2 = re.sub(r'\s+', ' ', (ln or '').strip())
            if not ln2:
                continue
            if len(ln2) <= 1 and not re.search(r'[0-9]', ln2):
                continue
            cleaned_lines.append(ln2)

        if not cleaned_lines:
            return ''

        out_parts = []
        cur = ''
        for ln in cleaned_lines:
            if not cur:
                cur = ln
                continue

            end_punct = cur.endswith(('.', '!', '?', ':', ';'))
            starts_bullet = bool(re.match(r'^(\-|\*|\d+[\.)])\s+', ln))
            if end_punct or starts_bullet:
                out_parts.append(cur)
                cur = ln
            else:
                cur = f"{cur} {ln}".strip()

        if cur:
            out_parts.append(cur)

        normalized = '\n'.join(out_parts)
        return normalized.strip()

    DB_IDENTIFIER_MAP = {
        "ma_khach_hang": "customer_id",
        "ngay_ban": "sale_date",
    }

    def _apply_db_identifier_map(text: str) -> str:
        out = text or ""
        for src, dst in DB_IDENTIFIER_MAP.items():
            out = re.sub(rf"\b{re.escape(src)}\b", dst, out, flags=re.IGNORECASE)
        return out

    def _cleanup_translated_text(text: str) -> str:
        out = "" if text is None else str(text)
        out = re.sub(r"\bNAMEBUILDING\b", "NAME BUILDING", out, flags=re.IGNORECASE)
        out = _apply_db_identifier_map(out)
        out = re.sub(r"\bNot\s+nul\b", "Not null", out, flags=re.IGNORECASE)
        out = re.sub(r"\bInfo\s+tin\s+basic\b", "Basic Information", out, flags=re.IGNORECASE)
        out = re.sub(r"\bWhere\s+the\s+topic\s+is\s+applied\b", "Application of the project", out, flags=re.IGNORECASE)
        out = re.sub(r"\bDevelopment\s+direction\s*:\s*Is\s+there\b", "Development direction: Yes", out, flags=re.IGNORECASE)
        out = re.sub(r"\bSTUDENT\s+ID\b", "Student ID", out, flags=re.IGNORECASE)
        out = re.sub(r"\bData\s+types\b", "Data Types", out, flags=re.IGNORECASE)
        out = re.sub(r"\bKHOA\s*CÔNG\s*NGHỆ\s*THÔNG\s*TIN\b", "FACULTY OF INFORMATION TECHNOLOGY", out, flags=re.IGNORECASE)
        out = re.sub(r"Gửi\s+lại\s+phiếu\s+đăng\s+ký\s+qua\s+Email\s*:", "Resubmit the registration form via Email:", out, flags=re.IGNORECASE)
        return out

    _guard_enabled = str(os.getenv("DOCX_TRANSLATION_GUARD", "1")).strip().lower() in ("1", "true", "yes", "on")
    _guard_url_re = re.compile(r"(https?://[^\s<>()]+|www\.[^\s<>()]+)", re.IGNORECASE)
    _guard_email_re = re.compile(r"\b[\w.+\-]+@[\w.\-]+\.[A-Za-z]{2,}\b")
    _guard_doi_re = re.compile(r"\b10\.\d{4,9}/[^\s\"'<>]+", re.IGNORECASE)
    _guard_doi_label_re = re.compile(r"\bdoi\s*:\s*10\.\d{4,9}/[^\s\"'<>]+", re.IGNORECASE)
    _guard_arxiv_re = re.compile(r"\barxiv\s*:\s*\S+", re.IGNORECASE)
    _guard_ref_line_re = re.compile(r"^\s*(\[\d+\]|\d{1,3}[.)])\s+\S+")
    _guard_code_block_re = re.compile(r"```.+?```", re.DOTALL)
    _guard_inline_code_re = re.compile(r"`[^`]+`")
    _guard_method_call_re = re.compile(r"\b[a-z][A-Za-z0-9_]*\s*\(\s*\)")
    _guard_type_attr_re = re.compile(
        r"\((?:PK|FK|string|int|date|boolean|float|double|varchar)[^)]{0,80}\)",
        re.IGNORECASE,
    )
    _guard_line_marker_re = re.compile(r"(?m)(^\s*(?:[+•➤▪]|o)\s*)")
    _guard_unicode_dash_re = re.compile(r"[–—−]")
    _guard_camel_ident_re = re.compile(r"\b[a-z]+[A-Z][A-Za-z0-9_]*\b")

    def _looks_like_formula(text: str) -> bool:
        s = (text or "").strip()
        if not s:
            return False
        math_syms = re.findall(r"[=+\-*/^_{}\\<>]", s)
        if len(math_syms) >= 2:
            letters = len(re.findall(r"[A-Za-zÀ-ỹ]", s))
            if letters <= 2:
                return True
            if len(math_syms) >= 4 and letters < len(math_syms):
                return True
        if re.search(r"\b\d+\s*[+\-*/]\s*\d+\b", s) and not re.search(r"[A-Za-zÀ-ỹ]", s):
            return True
        return False

    def _looks_like_code(text: str) -> bool:
        s = (text or "").strip()
        if not s:
            return False
        if _guard_code_block_re.search(s) or _guard_inline_code_re.search(s):
            return True
        if re.search(r"\b(def|class|function|var|let|const|public|private)\b", s):
            return True
        if re.search(r"[{}<>;]|->|=>|::", s) and len(s) <= 200:
            letters = len(re.findall(r"[A-Za-zÀ-ỹ]", s))
            if letters <= 4:
                return True
        return False

    _ref_heading_re = re.compile(
        r"^\s*(references|bibliography|works\s+cited|tài\s+liệu\s+tham\s+khảo|tai\s+lieu\s+tham\s+khao)\s*:?\s*$",
        re.IGNORECASE | re.UNICODE,
    )
    _ref_entry_line_re = re.compile(r"^\s*(?:\[\d+\]|\d{1,3}[.)])\s+\S")

    def _looks_like_english_reference_block(text: str) -> bool:
        """Whole paragraph looks like an English bibliography — pointless to translate to EN."""
        s = (text or "").strip()
        if len(s) < 80:
            return False
        try:
            if service._looks_vietnamese_like_text(s):
                return False
        except Exception:
            if re.search(r"[à-ỹđ]", s, flags=re.IGNORECASE):
                return False
        low = s.lower()
        if "doi:" in low or "doi.org/" in low or "arxiv:" in low:
            return True
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        if len(lines) < 2:
            return False
        numbered = sum(1 for ln in lines if _ref_entry_line_re.match(ln))
        if numbered >= 2:
            return True
        if len(re.findall(r"\(\s*\d{4}\s*[a-z]?\)", s)) >= 5:
            return True
        return False

    def _references_skip_translation_whole(text: str) -> bool:
        tl = str(target_lang).strip().lower()
        if not tl.startswith("en"):
            return False
        return _looks_like_english_reference_block(text)

    def _split_reference_entries(text: str):
        lines = (text or "").splitlines()
        chunks = []
        cur = []
        for ln in lines:
            if _ref_entry_line_re.match(ln):
                if cur:
                    chunks.append("\n".join(cur).strip())
                cur = [ln]
            else:
                cur.append(ln)
        if cur:
            chunks.append("\n".join(cur).strip())
        out = [c for c in chunks if c.strip()]
        return out if len(out) >= 2 else []

    def _paragraph_is_multi_reference_list(text: str) -> bool:
        return len(_split_reference_entries(text)) >= 2

    def _skip_reference_entry_translation(entry: str) -> bool:
        """Skip API translation for headings / citation lines already in English (or non‑VI when target is VI)."""
        e = (entry or "").strip()
        if not e:
            return True
        lines = e.splitlines()
        first = lines[0].strip() if lines else ""
        if _ref_heading_re.match(first) and not _ref_entry_line_re.search(e):
            return True
        citation_like = bool(
            _ref_entry_line_re.search(e)
            or "doi:" in e.lower()
            or "doi.org" in e.lower()
            or re.search(r"\b(?:19|20)\d{2}\b", e)
        )
        if not citation_like or len(e) < 22:
            return False
        try:
            looks_vi = service._looks_vietnamese_like_text(e)
        except Exception:
            looks_vi = bool(re.search(r"[à-ỹđ]", e, flags=re.IGNORECASE))
        return not looks_vi

    def _should_skip_translation(text: str) -> bool:
        s = (text or "").strip()
        if _references_skip_translation_whole(s):
            return True
        if not _guard_enabled:
            return False
        if not s:
            return False
        if len(s) <= 2 and s in {"+", "-", "*", "•", "▪", "➕", "➤", "◆", "o"}:
            return True
        if _guard_method_call_re.search(s) and s.count("(") >= 2:
            return True
        if _guard_url_re.fullmatch(s) or _guard_email_re.fullmatch(s):
            return True
        if _guard_doi_label_re.fullmatch(s) or _guard_doi_re.fullmatch(s):
            return True
        if _guard_arxiv_re.fullmatch(s):
            return True
        if _guard_ref_line_re.match(s):
            low = s.lower()
            if "doi" in low or "arxiv" in low or re.search(r"\b(19|20)\d{2}\b", s) or "," in s:
                return True
        if _looks_like_formula(s):
            return True
        if _looks_like_code(s):
            return True
        return False

    def _guard_mask_tokens(text: str) -> tuple[str, dict]:
        if not _guard_enabled:
            return text, {}
        placeholders = {}

        def _sub(pattern, src: str) -> str:
            def repl(match):
                key = f"__KEEP_{len(placeholders)}__"
                placeholders[key] = match.group(0)
                return key

            return pattern.sub(repl, src)

        out = text
        out = _sub(_guard_doi_label_re, out)
        out = _sub(_guard_doi_re, out)
        out = _sub(_guard_arxiv_re, out)
        out = _sub(_guard_url_re, out)
        out = _sub(_guard_email_re, out)
        out = _sub(_guard_method_call_re, out)
        out = _sub(_guard_type_attr_re, out)
        out = _sub(_guard_camel_ident_re, out)
        out = _sub(_guard_line_marker_re, out)
        out = _sub(_guard_unicode_dash_re, out)
        return out, placeholders

    def _guard_restore_tokens(text: str, placeholders: dict) -> str:
        out = text or ""
        for key, value in (placeholders or {}).items():
            out = out.replace(key, value)
        return out

    def _translate_preserve_exact_lines(text):
        raw = text or ""
        if not raw:
            return raw
        parts = re.split(r"(\r\n|\r|\n)", raw)
        out = []
        for part in parts:
            if part in ("\r\n", "\r", "\n"):
                out.append(part)
                continue
            if not part.strip():
                out.append(part)
                continue
            m = re.match(r"^(\s*)(.*?)(\s*)$", part, flags=re.DOTALL)
            if m:
                lead, core, tail = m.group(1), m.group(2), m.group(3)
            else:
                lead, core, tail = "", part, ""
            if not core.strip():
                out.append(part)
                continue
            translated_core = _translate_preserve_form_leaders(core)
            out.append(f"{lead}{translated_core}{tail}")
        return "".join(out)

    def _handle_multi_reference_bilingual(paragraph, paragraph_text: str) -> None:
        """Split numbered bibliography into entries; skip EN citations; newline → one § per entry below."""
        if bi_mode not in ("inline", "newline"):
            return
        entries = _split_reference_entries(paragraph_text)
        if len(entries) < 2:
            return
        if _references_skip_translation_whole(paragraph_text):
            return

        if bi_mode == "newline":
            anchor = paragraph
            for entry in entries:
                es = entry.strip()
                if _skip_reference_entry_translation(es):
                    continue
                try:
                    te = _translate_preserve_exact_lines(es).strip()
                except ProviderRateLimitError:
                    raise
                except Exception:
                    continue
                if not te or te == es:
                    continue
                np = _insert_paragraph_after(anchor, te, italic=False)
                if np is not None:
                    anchor = np
                    try:
                        _seen_para_elems.add(id(np))
                    except Exception:
                        pass
                else:
                    _append_translation_linebreak(anchor, te, italic=False)
            return

        # inline: one line per reference — "entry | trans"
        parts_out = []
        for entry in entries:
            es = entry.strip()
            if _skip_reference_entry_translation(es):
                parts_out.append(es)
                continue
            try:
                te = _translate_preserve_exact_lines(es).strip()
            except ProviderRateLimitError:
                raise
            except Exception:
                parts_out.append(es)
                continue
            if not te or te == es:
                parts_out.append(es)
            else:
                parts_out.append(service._join_inline_bilingual(es, te, bilingual_delimiter))
        combined = "\n".join(parts_out)
        _set_paragraph_text_preserve_runs(paragraph, combined)

    def image_part_ext(image_part):
        try:
            partname = str(getattr(image_part, 'partname', '') or '')
            base = os.path.basename(partname)
            ext = os.path.splitext(base)[1].lower()
            if ext:
                return ext
        except Exception:
            pass
        try:
            ct = str(getattr(image_part, 'content_type', '') or '').lower()
            mapping = {
                'image/png': '.png',
                'image/jpeg': '.jpg',
                'image/jpg': '.jpg',
                'image/gif': '.gif',
                'image/bmp': '.bmp',
                'image/tiff': '.tif',
                'image/webp': '.webp',
            }
            return mapping.get(ct, '.png')
        except Exception:
            return '.png'

    def _overlay_bytes_to_original_format(png_bytes: bytes, desired_ext: str) -> bytes:
        desired_ext = (desired_ext or '.png').lower()
        try:
            from PIL import Image
        except Exception:
            return png_bytes

        fmt_map = {
            '.png': 'PNG',
            '.jpg': 'JPEG',
            '.jpeg': 'JPEG',
            '.bmp': 'BMP',
            '.tif': 'TIFF',
            '.tiff': 'TIFF',
            '.webp': 'WEBP',
            '.gif': 'PNG',
        }
        out_fmt = fmt_map.get(desired_ext, 'PNG')
        try:
            img = Image.open(io.BytesIO(png_bytes))
            if out_fmt in ('JPEG', 'BMP', 'TIFF'):
                if img.mode not in ('RGB', 'L'):
                    img = img.convert('RGB')
            buf = io.BytesIO()
            img.save(buf, format=out_fmt)
            return buf.getvalue()
        except Exception:
            return png_bytes

    def _apply_translation_to_runs(paragraph, translated_text):
        runs = list(paragraph.runs)
        if not runs:
            paragraph.add_run(translated_text or "")
            return

        template_rpr = None
        apply_template_fn = None
        if bool(from_pdf):
            try:
                from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                    pick_translation_template_rpr,
                    _apply_template_rpr_to_run,
                )
                apply_template_fn = _apply_template_rpr_to_run
                template_rpr = pick_translation_template_rpr(paragraph)
            except Exception:
                template_rpr = None

        content_indices = []
        structural_indices = []
        for i, run in enumerate(runs):
            rt = run.text or ""
            if _is_structural_text(rt):
                structural_indices.append(i)
            else:
                content_indices.append(i)

        if not content_indices:
            return

        if structural_indices:
            primary = content_indices[0]
            cleaned = translated_text or ""
            if bool(from_pdf):
                try:
                    from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                        strip_leading_markers_from_text,
                    )
                    cleaned = strip_leading_markers_from_text(cleaned)
                except Exception:
                    pass
            runs[primary].text = cleaned
            for i in content_indices[1:]:
                runs[i].text = ""
        else:
            primary = content_indices[0]
            runs[primary].text = translated_text or ""
            for i, r in enumerate(runs):
                if i != primary:
                    r.text = ""

        if template_rpr is not None and apply_template_fn is not None:
            try:
                apply_template_fn(runs[primary], template_rpr)
            except Exception:
                pass

    def _apply_translated_text(paragraph, translated_text) -> None:
        if preserve_pdf_lines and ("\n" in (translated_text or "") or "\r" in (translated_text or "")):
            _apply_translation_with_linebreaks(paragraph, translated_text)
        else:
            _apply_translation_to_runs(paragraph, translated_text)
        if is_academic_pdf and "\n" not in (translated_text or "") and "\r" not in (translated_text or ""):
            collapse_pdf_prose_paragraph_runs(paragraph)

    def _get_run_format_key(run):
        from docx.oxml.ns import qn as _qn
        try:
            from lxml import etree
            rPr = run._element.find(_qn('w:rPr'))
            if rPr is not None:
                return etree.tostring(rPr)
            return b''
        except Exception:
            return b''

    def _group_runs_by_format(runs):
        groups = []
        for i, run in enumerate(runs):
            text = run.text or ""
            fmt = _get_run_format_key(run)

            if text.strip() and _is_structural_text(text):
                groups.append((b'__structural__' + fmt, [i]))
                continue

            if not text.strip():
                if groups:
                    groups[-1][1].append(i)
                else:
                    groups.append((fmt, [i]))
                continue

            if groups and groups[-1][0] == fmt:
                groups[-1][1].append(i)
            else:
                groups.append((fmt, [i]))
        return groups

    def _write_group_text(runs, indices, original_texts, new_text):
        written = False
        for i in indices:
            run = runs[i]
            if not written and (original_texts[i] or "").strip():
                run.text = new_text or ""
                written = True
            else:
                run.text = ""
        if not written and indices:
            runs[indices[0]].text = new_text or ""
            for i in indices[1:]:
                runs[i].text = ""

    def _split_translated_by_source_weights(translated_text, source_chunks):
        chunks = list(source_chunks or [])
        n = len(chunks)
        txt = translated_text or ""
        if n <= 0:
            return []
        if n == 1:
            return [txt]
        if not txt:
            return [""] * n

        weights = [max(1, len(c or "")) for c in chunks]
        total_weight = max(1, sum(weights))
        txt_len = len(txt)
        boundaries = []
        acc = 0
        for i in range(1, n):
            acc += weights[i - 1]
            target = int(round((acc / total_weight) * txt_len))
            target = max(1, min(txt_len - 1, target))

            left = max(1, target - 24)
            right = min(txt_len - 1, target + 24)
            cut = target

            found = None
            p = target
            while p >= left:
                if txt[p - 1].isspace():
                    found = p
                    break
                p -= 1
            if found is None:
                p = target + 1
                while p <= right:
                    if txt[p - 1].isspace():
                        found = p
                        break
                    p += 1
            if found is not None:
                cut = found

            if boundaries and cut <= boundaries[-1]:
                cut = min(txt_len - 1, boundaries[-1] + 1)
            boundaries.append(cut)

        out = []
        prev = 0
        for b in boundaries:
            out.append(txt[prev:b])
            prev = b
        out.append(txt[prev:])

        if len(out) < n:
            out.extend([""] * (n - len(out)))
        elif len(out) > n:
            out = out[:n - 1] + ["".join(out[n - 1:])]
        return out

    def _translate_format_groups(paragraph, translate_fn):
        from docx.oxml.ns import qn as _qn
        runs = list(paragraph.runs)
        if not runs:
            return

        original_texts = [(r.text or "") for r in runs]
        paragraph_text = join_docx_run_texts(original_texts)
        if not paragraph_text.strip():
            return

        has_form_leaders = bool(re.search(r"(\.{3,}|_{3,}|-{3,}|…+|\t+)", paragraph_text or ""))
        if (
            is_academic_pdf
            and "\n" not in paragraph_text
            and not has_form_leaders
            and len(runs) > 2
        ):
            try:
                translated = translate_fn(paragraph_text)
                _apply_translated_text(paragraph, translated)
                return
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Academic prose translation failed, fallback group mode: {e}")

        has_structural_runs = any(
            (original_texts[i] or "").strip() and _is_structural_text(original_texts[i])
            for i in range(len(runs))
        )

        groups = _group_runs_by_format(runs)

        if len(groups) <= 1 and not has_structural_runs:
            translated = translate_fn(paragraph_text)
            _apply_translated_text(paragraph, translated)
            return

        if len(groups) <= 1 and has_structural_runs:
            translated = translate_fn(paragraph_text)
            _apply_translated_text(paragraph, translated)
            return

        # Semantic-quality path for regular prose:
        # translate the whole paragraph once, then redistribute into style groups.
        # This avoids meaning drift caused by translating tiny style fragments.
        if not has_form_leaders:
            non_struct_groups = []
            for _, indices in groups:
                group_text = "".join(original_texts[i] for i in indices)
                if not group_text.strip() or _is_structural_text(group_text):
                    continue
                non_struct_groups.append((indices, group_text))

            if non_struct_groups:
                try:
                    translated_para = translate_fn(paragraph_text)
                    translated_chunks = _split_translated_by_source_weights(
                        translated_para,
                        [g[1] for g in non_struct_groups],
                    )
                    for (indices, _src), t_chunk in zip(non_struct_groups, translated_chunks):
                        _write_group_text(runs, indices, original_texts, t_chunk)
                    return
                except ProviderRateLimitError:
                    raise
                except Exception as e:
                    print(f"Paragraph-level redistribution failed, fallback group mode: {e}")

        for fmt_key, indices in groups:
            group_text = "".join(original_texts[i] for i in indices)
            if not group_text.strip():
                continue

            if _is_structural_text(group_text):
                continue

            try:
                translated_group = translate_fn(group_text)
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Format-group translation failed: {e}")
                if api_only:
                    raise
                translated_group = group_text
            _write_group_text(runs, indices, original_texts, translated_group)

    def _insert_paragraph_after(ref_para, text, italic=True, *, clear_first_line=False):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy

        def _to_twips(v):
            if v is None:
                return None
            try:
                return int(v.twips)
            except Exception:
                try:
                    return int(v)
                except Exception:
                    return None

        new_p = OxmlElement('w:p')

        try:
            pPr_src = ref_para._element.find(_qn('w:pPr'))
            if pPr_src is not None:
                new_pPr = _copy.deepcopy(pPr_src)

                try:
                    if new_pPr.find(_qn('w:pStyle')) is None:
                        style_id = str(getattr(getattr(ref_para, 'style', None), 'style_id', '') or '').strip()
                        if style_id:
                            p_style = OxmlElement('w:pStyle')
                            p_style.set(_qn('w:val'), style_id)
                            new_pPr.insert(0, p_style)
                except Exception:
                    pass

                numPr = new_pPr.find(_qn('w:numPr'))
                if numPr is not None:
                    src_left = None
                    src_hanging = None
                    src_firstLine = None
                    try:
                        fmt = ref_para.paragraph_format
                        src_left = _to_twips(fmt.left_indent)
                        src_firstLine = _to_twips(fmt.first_line_indent)
                    except Exception:
                        pass
                    if src_left is None:
                        try:
                            orig_ind = pPr_src.find(_qn('w:ind'))
                            if orig_ind is not None:
                                l_val = orig_ind.get(_qn('w:left')) or orig_ind.get(_qn('w:start'))
                                if l_val:
                                    src_left = int(l_val)
                                h_val = orig_ind.get(_qn('w:hanging'))
                                if h_val:
                                    src_hanging = int(h_val)
                                fl_val = orig_ind.get(_qn('w:firstLine'))
                                if fl_val:
                                    src_firstLine = int(fl_val)
                        except Exception:
                            pass

                    new_pPr.remove(numPr)

                    ind = new_pPr.find(_qn('w:ind'))
                    if ind is None and src_left is not None:
                        ind = OxmlElement('w:ind')
                        ind.set(_qn('w:left'), str(src_left))
                        if src_hanging is not None:
                            ind.set(_qn('w:hanging'), str(src_hanging))
                        elif src_firstLine is not None:
                            if src_firstLine < 0:
                                ind.set(_qn('w:hanging'), str(abs(src_firstLine)))
                            elif src_firstLine > 0:
                                ind.set(_qn('w:firstLine'), str(src_firstLine))
                        new_pPr.append(ind)

                jc = new_pPr.find(_qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(_qn('w:val'), '')
                    if jc_val in ('distribute', 'thai-distribute'):
                        new_pPr.remove(jc)

                new_p.insert(0, new_pPr)
            else:
                try:
                    style_id = str(getattr(getattr(ref_para, 'style', None), 'style_id', '') or '').strip()
                    if style_id:
                        new_pPr = OxmlElement('w:pPr')
                        p_style = OxmlElement('w:pStyle')
                        p_style.set(_qn('w:val'), style_id)
                        new_pPr.append(p_style)
                        new_p.insert(0, new_pPr)
                except Exception:
                    pass
        except Exception:
            pass

        def _run_is_superscript(run):
            try:
                if bool(getattr(getattr(run, 'font', None), 'superscript', False)):
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    va = rpr.find(_qn('w:vertAlign'))
                    if va is not None:
                        v = str(va.get(_qn('w:val'), '') or '').strip().lower()
                        if v in ('superscript', 'subscript'):
                            return True
            except Exception:
                pass
            return False

        def _run_is_bold(run):
            try:
                if run.bold is True:
                    return True
            except Exception:
                pass
            try:
                if getattr(getattr(run, 'font', None), 'bold', None) is True:
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(_qn('w:b'))
                    if b is not None:
                        bv = str(b.get(_qn('w:val'), '') or '').strip().lower()
                        if bv not in ('0', 'false', 'off'):
                            return True
            except Exception:
                pass
            return False

        def _clone_rpr(src_rpr, *, force_non_bold=False):
            if src_rpr is not None:
                rpr = _copy.deepcopy(src_rpr)
            else:
                rpr = OxmlElement('w:rPr')
            try:
                sp = rpr.find(_qn('w:spacing'))
                if sp is not None:
                    rpr.remove(sp)
            except Exception:
                pass
            try:
                for tag in ('w:caps', 'w:smallCaps'):
                    el = rpr.find(_qn(tag))
                    if el is not None:
                        rpr.remove(el)
            except Exception:
                pass
            if force_non_bold:
                try:
                    for tag in ('w:b', 'w:bCs'):
                        b_el = rpr.find(_qn(tag))
                        if b_el is not None:
                            rpr.remove(b_el)
                        off_el = OxmlElement(tag)
                        off_el.set(_qn('w:val'), '0')
                        rpr.append(off_el)
                except Exception:
                    pass
            return rpr

        def _append_run_with_rpr(p_el, txt, src_rpr, *, force_non_bold=False):
            if txt is None or txt == '':
                return
            chunks = str(txt).replace('\r\n', '\n').replace('\r', '\n').split('\n')
            for idx, chunk in enumerate(chunks):
                if idx > 0:
                    br_el = OxmlElement('w:r')
                    br_rpr = _clone_rpr(src_rpr, force_non_bold=force_non_bold)
                    if italic:
                        try:
                            if br_rpr.find(_qn('w:i')) is None:
                                br_rpr.append(OxmlElement('w:i'))
                        except Exception:
                            pass
                    br_el.insert(0, br_rpr)
                    br_el.append(OxmlElement('w:br'))
                    p_el.append(br_el)
                if chunk == '':
                    continue
                r_el = OxmlElement('w:r')
                rpr = _clone_rpr(src_rpr, force_non_bold=force_non_bold)
                if italic:
                    try:
                        if rpr.find(_qn('w:i')) is None:
                            rpr.append(OxmlElement('w:i'))
                    except Exception:
                        pass
                r_el.insert(0, rpr)
                t_el = OxmlElement('w:t')
                t_el.set(_qn('xml:space'), 'preserve')
                t_el.text = chunk
                r_el.append(t_el)
                p_el.append(r_el)

        def _append_with_super_markers(p_el, txt, base_rpr, sup_rpr, markers, *, force_non_bold_base=False):
            if txt is None or txt == '':
                return
            marker_values = [m for m in (markers or []) if m]
            if sup_rpr is None or not marker_values:
                _append_run_with_rpr(p_el, txt, base_rpr, force_non_bold=force_non_bold_base)
                return
            marker_values = sorted(set(marker_values), key=len, reverse=True)
            pattern = r'(?<!\d)(' + '|'.join(re.escape(m) for m in marker_values) + r')(?!\d)'
            try:
                parts = re.split(pattern, txt)
            except Exception:
                _append_run_with_rpr(p_el, txt, base_rpr, force_non_bold=force_non_bold_base)
                return
            if not parts:
                _append_run_with_rpr(p_el, txt, base_rpr, force_non_bold=force_non_bold_base)
                return
            for idx, part in enumerate(parts):
                if part is None or part == '':
                    continue
                if idx % 2 == 1:
                    _append_run_with_rpr(p_el, part, sup_rpr)
                else:
                    _append_run_with_rpr(p_el, part, base_rpr, force_non_bold=force_non_bold_base)

        body_rpr = None
        bold_rpr = None
        super_rpr = None
        super_digit_markers = []
        has_mixed_bold_prefix = False
        has_leading_sup_marker = False
        body_should_be_non_bold = False

        def _is_heading_like(p):
            try:
                sn = str(getattr(getattr(p, 'style', None), 'name', '') or '').strip().lower()
                if sn.startswith('heading') or sn in ('title', 'subtitle'):
                    return True
                if 'title' in sn or 'tiêu đề' in sn or 'tieu de' in sn:
                    return True
                return False
            except Exception:
                return False

        def _is_centered_like(p):
            try:
                if getattr(p, 'alignment', None) == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    return True
            except Exception:
                pass
            try:
                ppr = p._element.find(_qn('w:pPr'))
                if ppr is not None:
                    jc = ppr.find(_qn('w:jc'))
                    if jc is not None:
                        return str(jc.get(_qn('w:val'), '') or '').strip().lower() == 'center'
            except Exception:
                pass
            return False

        def _is_marker_run_text(t):
            s = (t or '').strip()
            if not s:
                return False
            return bool(re.match(r'^\d+(?:[\.)]\d+)*[\.)]?$', s))

        def _extract_sup_markers_from_text(t):
            s = (t or '').strip()
            if not s:
                return []
            # Capture short affiliation/reference markers (1, 2, 10, ...)
            # and ignore long numbers like years/ORCID chunks.
            return re.findall(r'(?<!\d)(\d{1,2})(?=(?:[\)\]\.,;\s]|$|\[))', s)

        def _rpr_size(rpr):
            if rpr is None:
                return -1
            try:
                sz = rpr.find(_qn('w:sz'))
                if sz is not None:
                    raw = str(sz.get(_qn('w:val'), '') or '').strip()
                    if raw.isdigit():
                        return int(raw)
            except Exception:
                pass
            return -1

        def _best_rpr_from_runs(runs):
            best = None
            best_sz = -1
            for rr in runs or []:
                try:
                    if _run_is_superscript(rr):
                        continue
                    rpr = rr._element.find(_qn('w:rPr'))
                    if rpr is None:
                        continue
                    sz = _rpr_size(rpr)
                    if best is None or sz > best_sz:
                        best = rpr
                        best_sz = sz
                except Exception:
                    continue
            return best

        try:
            src_runs = [r for r in list(ref_para.runs) if (r.text or '').strip()]
            if src_runs:
                non_sup_runs = [r for r in src_runs if not _run_is_superscript(r)]
                content_runs = [r for r in non_sup_runs if not _is_marker_run_text(r.text or '')]
                para_text = join_docx_run_texts([(r.text or '') for r in src_runs])

                total_chars = sum(len((r.text or '').strip()) for r in content_runs)
                bold_chars = sum(len((r.text or '').strip()) for r in content_runs if _run_is_bold(r))
                bold_ratio = (float(bold_chars) / float(max(1, total_chars))) if total_chars > 0 else 0.0
                line_count = len([ln for ln in (para_text or '').splitlines() if ln.strip()])
                centered_title_like = _is_centered_like(ref_para) and total_chars <= 220 and line_count <= 3
                prefer_bold_body = (
                    _is_heading_like(ref_para)
                    or (centered_title_like and bold_ratio >= 0.35)
                    or (bold_chars >= 4 and bold_ratio >= 0.60)
                )

                body_should_be_non_bold = (not prefer_bold_body) and any(
                    (not _run_is_bold(r)) for r in content_runs
                )
                # Body style: prefer non-bold + non-superscript to avoid whole-line bold/superscript drift.
                chosen_run = None
                if prefer_bold_body:
                    for r in content_runs:
                        if _run_is_bold(r):
                            chosen_run = r
                            break
                    if chosen_run is None:
                        for r in non_sup_runs:
                            if not _is_marker_run_text(r.text or ''):
                                chosen_run = r
                                break
                else:
                    for r in content_runs:
                        if not _run_is_bold(r):
                            chosen_run = r
                            break
                    if chosen_run is None:
                        for r in non_sup_runs:
                            if _run_is_bold(r):
                                continue
                            chosen_run = r
                            break
                if chosen_run is None:
                    chosen_run = (content_runs[0] if content_runs else (non_sup_runs[0] if non_sup_runs else None))
                if chosen_run is None:
                    chosen_run = src_runs[0]

                body_rpr = chosen_run._element.find(_qn('w:rPr'))
                if body_rpr is None:
                    body_rpr = _best_rpr_from_runs(content_runs if content_runs else non_sup_runs)

                for r in src_runs:
                    if _run_is_bold(r) and not _run_is_superscript(r):
                        bold_rpr = r._element.find(_qn('w:rPr'))
                        if bold_rpr is not None:
                            break
                for r in src_runs:
                    if _run_is_superscript(r):
                        super_rpr = r._element.find(_qn('w:rPr'))
                        if super_rpr is not None:
                            break
                for r in src_runs:
                    if not _run_is_superscript(r):
                        continue
                    token = (r.text or '').strip()
                    if token and len(token) <= 64 and (not re.search(r'\s', token)) and any(ch.isdigit() for ch in token):
                        if token not in super_digit_markers:
                            super_digit_markers.append(token)
                    for marker in _extract_sup_markers_from_text(token):
                        if marker not in super_digit_markers:
                            super_digit_markers.append(marker)

                first_run = src_runs[0]
                first_text = (first_run.text or '').strip()
                has_leading_sup_marker = bool(
                    _run_is_superscript(first_run)
                    and re.match(r'^\d{1,2}(?:[\)\.]?)?(?:$|[\s\[,;])', first_text)
                )

                prefix = ''
                saw_after = False
                for r in src_runs:
                    rt = r.text or ''
                    if not rt.strip() or _run_is_superscript(r):
                        continue
                    if _run_is_bold(r) and not saw_after:
                        prefix += rt
                        continue
                    if prefix.strip():
                        saw_after = True
                    break
                if prefix.strip() and saw_after and len(prefix.strip()) <= 48:
                    has_mixed_bold_prefix = bool(re.search(r'[:：\.]\s*$', prefix.strip()))

            if body_rpr is None:
                body_rpr = _best_rpr_from_runs(src_runs if 'src_runs' in locals() else [])
            if not italic:
                try:
                    dom_rpr = _best_rpr_from_runs(
                        content_runs if content_runs else (non_sup_runs if 'non_sup_runs' in locals() else [])
                    )
                    if dom_rpr is not None:
                        body_rpr = dom_rpr
                except Exception:
                    pass
        except Exception:
            try:
                body_rpr = _best_rpr_from_runs(list(ref_para.runs))
            except Exception:
                body_rpr = None

        remaining = '' if text is None else str(text)

        if has_leading_sup_marker and super_rpr is not None:
            m_ref = re.match(r'^(\s*\d{1,2}(?:[\)\.]?)(?:\s+|(?=[\[,;])))(.*)$', remaining, flags=re.DOTALL)
            if m_ref:
                _append_run_with_rpr(new_p, m_ref.group(1), super_rpr)
                remaining = m_ref.group(2)

        wrote_split = False
        if has_mixed_bold_prefix and bold_rpr is not None:
            m_label = re.match(r'^(\s*[^\n]{1,48}?(?:[:：\.])\s*)(.*)$', remaining, flags=re.DOTALL)
            if m_label and (m_label.group(2) or '').strip():
                _append_run_with_rpr(new_p, m_label.group(1), bold_rpr)
                _append_with_super_markers(
                    new_p,
                    m_label.group(2),
                    body_rpr,
                    super_rpr,
                    super_digit_markers,
                    force_non_bold_base=body_should_be_non_bold,
                )
                wrote_split = True

        if not wrote_split:
            _append_with_super_markers(
                new_p,
                remaining,
                body_rpr,
                super_rpr,
                super_digit_markers,
                force_non_bold_base=body_should_be_non_bold,
            )
        try:
            ref_para._element.addnext(new_p)
            if clear_first_line:
                ppr = new_p.find(_qn('w:pPr'))
                if ppr is not None:
                    ind = ppr.find(_qn('w:ind'))
                    if ind is not None:
                        for key in (_qn('w:firstLine'), _qn('w:hanging')):
                            ind.attrib.pop(key, None)
            return new_p
        except Exception:
            return None

    def _append_translation_linebreak(paragraph, text, italic=True):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as _qn
        import copy as _copy

        txt = (text or '').strip()
        if not txt:
            return False

        def _run_is_superscript(run):
            try:
                if bool(getattr(getattr(run, 'font', None), 'superscript', False)):
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    va = rpr.find(_qn('w:vertAlign'))
                    if va is not None:
                        v = str(va.get(_qn('w:val'), '') or '').strip().lower()
                        if v in ('superscript', 'subscript'):
                            return True
            except Exception:
                pass
            return False

        def _run_is_bold(run):
            try:
                if run.bold is True:
                    return True
            except Exception:
                pass
            try:
                if getattr(getattr(run, 'font', None), 'bold', None) is True:
                    return True
            except Exception:
                pass
            try:
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(_qn('w:b'))
                    if b is not None:
                        bv = str(b.get(_qn('w:val'), '') or '').strip().lower()
                        if bv not in ('0', 'false', 'off'):
                            return True
            except Exception:
                pass
            return False

        def _apply_rpr_to_run(dst_run, src_rpr, *, force_non_bold=False):
            if src_rpr is None:
                return
            try:
                new_rpr = _copy.deepcopy(src_rpr)
                sp = new_rpr.find(_qn('w:spacing'))
                if sp is not None:
                    new_rpr.remove(sp)
                for tag in ('w:caps', 'w:smallCaps'):
                    el = new_rpr.find(_qn(tag))
                    if el is not None:
                        new_rpr.remove(el)
                if force_non_bold:
                    for tag in ('w:b', 'w:bCs'):
                        b_el = new_rpr.find(_qn(tag))
                        if b_el is not None:
                            new_rpr.remove(b_el)
                        off_el = OxmlElement(tag)
                        off_el.set(_qn('w:val'), '0')
                        new_rpr.append(off_el)
                if italic and new_rpr.find(_qn('w:i')) is None:
                    new_rpr.append(OxmlElement('w:i'))
                old_rpr = dst_run._element.find(_qn('w:rPr'))
                if old_rpr is not None:
                    dst_run._element.remove(old_rpr)
                dst_run._element.insert(0, new_rpr)
            except Exception:
                pass

        def _append_styled_text_with_breaks(paragraph_obj, raw_text, src_rpr, *, force_non_bold=False):
            chunks = str(raw_text or '').replace('\r\n', '\n').replace('\r', '\n').split('\n')
            wrote_any = False
            for idx, chunk in enumerate(chunks):
                if idx > 0:
                    br = paragraph_obj.add_run('')
                    _apply_rpr_to_run(br, src_rpr, force_non_bold=force_non_bold)
                    br.add_break()
                    wrote_any = True
                if chunk == '':
                    continue
                rr = paragraph_obj.add_run(chunk)
                _apply_rpr_to_run(rr, src_rpr, force_non_bold=force_non_bold)
                wrote_any = True
            return wrote_any

        try:
            paragraph.add_run('').add_break()

            src_runs = [r for r in list(paragraph.runs) if (r.text or '').strip()]
            body_rpr = None
            bold_rpr = None
            has_mixed_bold_prefix = False
            body_should_be_non_bold = False

            def _is_heading_like(p):
                try:
                    sn = str(getattr(getattr(p, 'style', None), 'name', '') or '').strip().lower()
                    if sn.startswith('heading') or sn in ('title', 'subtitle'):
                        return True
                    if 'title' in sn or 'tiêu đề' in sn or 'tieu de' in sn:
                        return True
                    return False
                except Exception:
                    return False

            def _is_centered_like(p):
                try:
                    if getattr(p, 'alignment', None) == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        return True
                except Exception:
                    pass
                try:
                    ppr = p._element.find(_qn('w:pPr'))
                    if ppr is not None:
                        jc = ppr.find(_qn('w:jc'))
                        if jc is not None:
                            return str(jc.get(_qn('w:val'), '') or '').strip().lower() == 'center'
                except Exception:
                    pass
                return False

            def _is_marker_run_text(t):
                s = (t or '').strip()
                if not s:
                    return False
                return bool(re.match(r'^\d+(?:[\.)]\d+)*[\.)]?$', s))

            def _rpr_size(rpr):
                if rpr is None:
                    return -1
                try:
                    sz = rpr.find(_qn('w:sz'))
                    if sz is not None:
                        raw = str(sz.get(_qn('w:val'), '') or '').strip()
                        if raw.isdigit():
                            return int(raw)
                except Exception:
                    pass
                return -1

            def _best_rpr_from_runs(runs):
                best = None
                best_sz = -1
                for rr in runs or []:
                    try:
                        if _run_is_superscript(rr):
                            continue
                        rpr = rr._element.find(_qn('w:rPr'))
                        if rpr is None:
                            continue
                        sz = _rpr_size(rpr)
                        if best is None or sz > best_sz:
                            best = rpr
                            best_sz = sz
                    except Exception:
                        continue
                return best

            try:
                non_sup_runs = [r for r in src_runs if not _run_is_superscript(r)]
                content_runs = [r for r in non_sup_runs if not _is_marker_run_text(r.text or '')]
                para_text = join_docx_run_texts([(r.text or '') for r in src_runs])

                total_chars = sum(len((r.text or '').strip()) for r in content_runs)
                bold_chars = sum(len((r.text or '').strip()) for r in content_runs if _run_is_bold(r))
                bold_ratio = (float(bold_chars) / float(max(1, total_chars))) if total_chars > 0 else 0.0
                line_count = len([ln for ln in (para_text or '').splitlines() if ln.strip()])
                centered_title_like = _is_centered_like(paragraph) and total_chars <= 220 and line_count <= 3
                prefer_bold_body = (
                    _is_heading_like(paragraph)
                    or (centered_title_like and bold_ratio >= 0.35)
                    or (bold_chars >= 4 and bold_ratio >= 0.60)
                )

                body_should_be_non_bold = (not prefer_bold_body) and any(
                    (not _run_is_bold(r)) for r in content_runs
                )

                chosen = None
                if prefer_bold_body:
                    for r in content_runs:
                        if _run_is_bold(r):
                            chosen = r
                            break
                    if chosen is None:
                        for r in non_sup_runs:
                            if not _is_marker_run_text(r.text or ''):
                                chosen = r
                                break
                else:
                    for r in content_runs:
                        if not _run_is_bold(r):
                            chosen = r
                            break
                    if chosen is None:
                        for r in non_sup_runs:
                            if _run_is_bold(r):
                                continue
                            chosen = r
                            break
                if chosen is None:
                    chosen = (content_runs[0] if content_runs else (non_sup_runs[0] if non_sup_runs else None))
                if chosen is None and src_runs:
                    chosen = src_runs[0]
                if chosen is not None:
                    body_rpr = chosen._element.find(_qn('w:rPr'))
                if body_rpr is None:
                    body_rpr = _best_rpr_from_runs(content_runs if content_runs else non_sup_runs)
                if body_rpr is None:
                    body_rpr = _best_rpr_from_runs(src_runs)

                for r in src_runs:
                    if _run_is_bold(r) and not _run_is_superscript(r):
                        bold_rpr = r._element.find(_qn('w:rPr'))
                        if bold_rpr is not None:
                            break

                prefix = ''
                saw_after = False
                for r in src_runs:
                    rt = r.text or ''
                    if not rt.strip() or _run_is_superscript(r):
                        continue
                    if _run_is_bold(r) and not saw_after:
                        prefix += rt
                        continue
                    if prefix.strip():
                        saw_after = True
                    break
                if prefix.strip() and saw_after and len(prefix.strip()) <= 48:
                    has_mixed_bold_prefix = bool(re.search(r'[:：\.]\s*$', prefix.strip()))
            except Exception:
                pass

            m_label = None
            if has_mixed_bold_prefix and bold_rpr is not None:
                m_label = re.match(r'^(\s*[^\n]{1,48}?(?:[:：\.])\s*)(.*)$', txt, flags=re.DOTALL)

            if m_label and (m_label.group(2) or '').strip():
                _append_styled_text_with_breaks(paragraph, m_label.group(1), bold_rpr)
                _append_styled_text_with_breaks(
                    paragraph,
                    m_label.group(2),
                    body_rpr,
                    force_non_bold=body_should_be_non_bold,
                )
            else:
                _append_styled_text_with_breaks(
                    paragraph,
                    txt,
                    body_rpr,
                    force_non_bold=body_should_be_non_bold,
                )
            return True
        except Exception:
            return False

    leader_re = re.compile(r"(\.{3,}|_{3,}|-{3,}|…+|\t+)")

    def _is_structural_text(text):
        t = (text or "").strip()
        if not t:
            return True
        return not re.search(r'[\w\u00C0-\u1EF9]', t, flags=re.UNICODE)

    def _translate_preserve_form_leaders(text):
        raw = text or ""
        if not raw.strip():
            return raw
        if _is_structural_text(raw):
            return raw
        if _guard_enabled and _should_skip_translation(raw):
            return raw

        # Check batch pre-translation cache for speed
        _cache_key = raw.strip()
        if _cache_key in _translation_cache:
            return _translation_cache[_cache_key]

        if not leader_re.search(raw):
            masked, placeholders = _guard_mask_tokens(raw) if _guard_enabled else (raw, {})
            result = _cleanup_translated_text(
                service._translate_with_retry(masked, target_lang, context='document_docx_line')
            )
            result = _guard_restore_tokens(result, placeholders)
            _translation_cache[_cache_key] = result
            return result

        parts = leader_re.split(raw)
        out_parts = []
        for i, part in enumerate(parts):
            if i % 2 == 1:
                out_parts.append(part)
                continue

            seg = part or ""
            if not seg.strip():
                out_parts.append(seg)
                continue
            if _is_structural_text(seg):
                out_parts.append(seg)
                continue

            try:
                _seg_key = seg.strip()
                if _guard_enabled and _should_skip_translation(seg):
                    out_parts.append(seg)
                elif _seg_key in _translation_cache:
                    out_parts.append(_translation_cache[_seg_key])
                else:
                    masked, placeholders = _guard_mask_tokens(seg) if _guard_enabled else (seg, {})
                    translated = _cleanup_translated_text(
                        service._translate_with_retry(masked, target_lang, context='document_docx_line')
                    )
                    translated = _guard_restore_tokens(translated, placeholders)
                    _translation_cache[_seg_key] = translated
                    out_parts.append(translated)
            except ProviderRateLimitError:
                raise
            except Exception:
                if api_only:
                    raise
                out_parts.append(seg)

        return _cleanup_translated_text("".join(out_parts))

    _pdf_list_line_re = re.compile(r"^\s*(?:[-*]|\(?\d{1,3}[.)]|[A-Za-z][.)])\s+\S")

    def _normalize_from_pdf_soft_breaks(text: str) -> str:
        """Collapse pdf2docx soft wraps into normal prose, keep list/reference structures."""
        raw = "" if text is None else str(text)
        if preserve_pdf_lines:
            return raw
        if not bool(from_pdf) or not raw.strip():
            return raw

        normalized = raw.replace("\r\n", "\n").replace("\r", "\n")
        if "\n" not in normalized:
            return raw

        if _paragraph_is_multi_reference_list(normalized):
            return raw
        if leader_re.search(normalized):
            return raw
        if _looks_like_formula(normalized) or _looks_like_code(normalized):
            return raw

        lines = [ln.strip() for ln in normalized.split("\n") if ln.strip()]
        if len(lines) < 2:
            return raw
        if any(_pdf_list_line_re.match(ln) for ln in lines):
            return raw

        compact = re.sub(r"(?<=[A-Za-z\u00C0-\u1EF9])-\n(?=[A-Za-z\u00C0-\u1EF9])", "", normalized)
        compact = re.sub(r"(?<!\n)\n(?!\n)", " ", compact)
        compact = re.sub(r"[ \t]+", " ", compact)
        compact = re.sub(r" *\n *", "\n", compact)
        compact = re.sub(r"\n{3,}", "\n\n", compact)
        compact = compact.strip()
        return compact or raw

    if bool(from_pdf):
        try:
            from app.services.document_v2.pdf_docx_pipeline.layout_recovery import (
                is_pdf_artifact_text,
                is_running_header_text,
            )
        except Exception:
            is_running_header_text = lambda _t: False  # noqa: E731
            is_pdf_artifact_text = lambda _t: False  # noqa: E731

        def _is_pdf_noise_text(text: str) -> bool:
            return is_running_header_text(text) or is_pdf_artifact_text(text)
    else:
        def _is_pdf_noise_text(text: str) -> bool:
            return False

    def _translate_body_paragraph_text(text):
        if _is_pdf_noise_text((text or "").strip()):
            return text or ""
        if preserve_pdf_lines:
            return _translate_preserve_exact_lines(text)
        prepared = _normalize_from_pdf_soft_breaks(text)
        if bool(from_pdf):
            return _translate_preserve_form_leaders(prepared)
        return _translate_preserve_exact_lines(prepared)

    def _is_toc_paragraph(paragraph):
        try:
            style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
            if "toc" in style_name:
                return True
        except Exception:
            pass

        p_el = paragraph._element
        try:
            for node in p_el.xpath('.//*[local-name()="instrText"]'):
                txt = "".join(node.itertext())
                if "toc" in (txt or "").lower():
                    return True
        except Exception:
            pass

        try:
            for node in p_el.xpath('.//*[local-name()="fldSimple"]'):
                for k, v in (node.attrib or {}).items():
                    if str(k).endswith("}instr") and "toc" in str(v or "").lower():
                        return True
        except Exception:
            pass

        return False

    def _flatten_hyperlinks_in_paragraph(paragraph):
        changed = False
        p_el = paragraph._element
        while True:
            links = list(p_el.xpath('./*[local-name()="hyperlink"]'))
            if not links:
                break
            for link in links:
                parent = link.getparent()
                if parent is None:
                    continue
                idx = parent.index(link)
                for child in list(link):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(link)
                changed = True
        return changed

    def _normalize_toc_run_appearance(paragraph):
        return

    def _normalize_toc_hyperlinks(document):
        touched = 0
        for para in iter_all_paragraphs(document):
            if not _is_toc_paragraph(para):
                continue
            if _flatten_hyperlinks_in_paragraph(para):
                touched += 1
            _normalize_toc_run_appearance(para)
        return touched

    def _normalize_generic_run_appearance(paragraph):
        for run in paragraph.runs:
            try:
                run.underline = False
            except Exception:
                pass
            try:
                run.font.color.theme_color = None
            except Exception:
                pass
            try:
                run.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass

    def _strip_all_hyperlinks(document):
        touched = 0
        for para in iter_all_paragraphs(document):
            if _flatten_hyperlinks_in_paragraph(para):
                touched += 1
            _normalize_generic_run_appearance(para)
        return touched

    def _paragraph_has_drawing(paragraph):
        try:
            return bool(paragraph._element.xpath('.//*[local-name()="drawing" or local-name()="pict"]'))
        except Exception:
            return False

    def _set_paragraph_text_preserve_runs(paragraph, new_text):
        runs = list(paragraph.runs)
        if not runs:
            paragraph.add_run(new_text or "")
            return

        if _paragraph_has_drawing(paragraph):
            non_drawing_runs = []
            for r in runs:
                try:
                    has_draw = bool(r._element.xpath('.//*[local-name()="drawing" or local-name()="pict"]'))
                except Exception:
                    has_draw = False
                if not has_draw:
                    non_drawing_runs.append(r)

            if not non_drawing_runs:
                return

            target = None
            for r in non_drawing_runs:
                if (r.text or "").strip():
                    target = r
                    break
            if target is None:
                target = non_drawing_runs[0]

            target.text = new_text or ""
            for r in non_drawing_runs:
                if r is not target:
                    r.text = ""
            return

        target = None
        for r in runs:
            if (r.text or "").strip():
                target = r
                break
        if target is None:
            target = runs[0]
        target.text = new_text or ""
        for r in runs:
            if r is not target:
                r.text = ""

    def _is_in_table_cell(paragraph):
        try:
            parent = paragraph._element.getparent()
            return bool(parent is not None and (parent.tag or '').endswith('}tc'))
        except Exception:
            return False

    def _is_heading_paragraph(paragraph):
        try:
            style_name = str(getattr(getattr(paragraph, 'style', None), 'name', '') or '').lower()
            style_name = style_name.strip()
            if style_name.startswith('heading') or style_name in ('title', 'subtitle'):
                return True
            if 'title' in style_name or 'tiêu đề' in style_name or 'tieu de' in style_name:
                return True
            return False
        except Exception:
            return False

    def _append_inline_bilingual(paragraph, translated_text, delimiter):
        t = (translated_text or '').strip()
        if not t:
            return False

        apply_docx_paragraph_spacing(paragraph)
        runs = list(paragraph.runs)

        if not runs:
            paragraph.add_run(t)
            return True

        def _run_is_superscript(run):
            try:
                if bool(getattr(getattr(run, 'font', None), 'superscript', False)):
                    return True
            except Exception:
                pass
            try:
                from docx.oxml.ns import qn as _qn
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    va = rpr.find(_qn('w:vertAlign'))
                    if va is not None:
                        v = str(va.get(_qn('w:val'), '') or '').strip().lower()
                        if v in ('superscript', 'subscript'):
                            return True
            except Exception:
                pass
            return False

        def _run_is_bold(run):
            try:
                if run.bold is True:
                    return True
            except Exception:
                pass
            try:
                if getattr(getattr(run, 'font', None), 'bold', None) is True:
                    return True
            except Exception:
                pass
            try:
                from docx.oxml.ns import qn as _qn
                rpr = run._element.find(_qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(_qn('w:b'))
                    if b is not None:
                        bv = str(b.get(_qn('w:val'), '') or '').strip().lower()
                        if bv not in ('0', 'false', 'off'):
                            return True
            except Exception:
                pass
            return False

        def _apply_run_style(dst_run, src_run):
            if src_run is None:
                return
            try:
                import copy as _copy
                from docx.oxml.ns import qn as _qn
                src_rpr = src_run._element.find(_qn('w:rPr'))
                if src_rpr is None:
                    return
                new_rpr = _copy.deepcopy(src_rpr)
                sp = new_rpr.find(_qn('w:spacing'))
                if sp is not None:
                    new_rpr.remove(sp)
                old_rpr = dst_run._element.find(_qn('w:rPr'))
                if old_rpr is not None:
                    dst_run._element.remove(old_rpr)
                dst_run._element.insert(0, new_rpr)
            except Exception:
                pass

        src_runs = [r for r in runs if (r.text or '').strip()]
        non_sup_runs = [r for r in src_runs if not _run_is_superscript(r)]

        bold_template = None
        for r in non_sup_runs:
            if _run_is_bold(r):
                bold_template = r
                break

        body_template = None
        for r in non_sup_runs:
            if not _run_is_bold(r):
                body_template = r
                break
        if body_template is None and non_sup_runs:
            body_template = non_sup_runs[0]
        if body_template is None and src_runs:
            body_template = src_runs[0]

        has_mixed_bold_prefix = False
        try:
            prefix = ''
            saw_after = False
            for r in src_runs:
                rt = r.text or ''
                if not rt.strip() or _run_is_superscript(r):
                    continue
                if _run_is_bold(r) and not saw_after:
                    prefix += rt
                    continue
                if prefix.strip():
                    saw_after = True
                break
            if prefix.strip() and saw_after and len(prefix.strip()) <= 48:
                has_mixed_bold_prefix = bool(re.search(r'[:：\.]\s*$', prefix.strip()))
        except Exception:
            pass

        last_run = None
        for r in reversed(runs):
            if (r.text or '').strip():
                last_run = r
                break
        if last_run is None:
            last_run = runs[-1]

        spacer = "" if (last_run.text or "").endswith((" ", "\t")) else " "
        last_run.text = f"{last_run.text or ''}{spacer}{delimiter} "

        if has_mixed_bold_prefix and bold_template is not None:
            m_label = re.match(r'^(\s*[^\n]{1,48}?(?:[:：\.])\s*)(.*)$', t, flags=re.DOTALL)
            if m_label and (m_label.group(2) or '').strip():
                r1 = paragraph.add_run(m_label.group(1))
                _apply_run_style(r1, bold_template)
                r2 = paragraph.add_run(m_label.group(2))
                _apply_run_style(r2, body_template)
                return True

        tr = paragraph.add_run(t)
        _apply_run_style(tr, body_template)
        return True

    def _paragraph_uses_numbering(paragraph):
        try:
            ppr = paragraph._element.find('.//*[local-name()="pPr"]')
            if ppr is not None:
                num_pr = ppr.find('.//*[local-name()="numPr"]')
                if num_pr is not None:
                    return True
        except Exception:
            pass
        try:
            txt = ''.join((r.text or '') for r in paragraph.runs).strip()
        except Exception:
            txt = ''
        if not txt:
            return False
        return bool(re.match(r'^\s*\d+(?:[\.)]|(?:\.\d+)+(?:[\.)])?)\s+\S', txt))

    def _prefer_linebreak_for_newline(paragraph):
        return _is_heading_paragraph(paragraph) or _paragraph_uses_numbering(paragraph)

    def _normalize_heading_case(document):
        return 0

    def _normalize_table_header_text(text: str) -> str:
        t = (text or '').strip()
        norm = re.sub(r'\s+', ' ', t).lower()
        mapping = {
            'user id': 'User ID',
            'userid': 'User ID',
            'data type': 'Data Types',
            'data types': 'Data Types',
            'description': 'Description',
            'constraints': 'Constraints',
            'constraint': 'Constraints',
            'not nul': 'Not null',
        }
        return mapping.get(norm, t)

    def _normalize_table_layout_and_text(document):
        touched = 0
        term_map = {
            'school': 'Field',
            'mô tả': 'Description',
            'mo ta': 'Description',
            'ràng buộc': 'Constraints',
            'rang buoc': 'Constraints',
            'data types': 'Data Types',
            'data type': 'Data Types',
        }

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        raw = ''.join((rr.text or '') for rr in para.runs)
                        if not raw.strip():
                            continue

                        fixed = _cleanup_translated_text(raw)
                        lowered = fixed.strip().lower()
                        if lowered in term_map:
                            fixed = term_map[lowered]

                        if str(target_lang).strip().lower().startswith('en') and re.search(r'[à-ỹđ]', fixed, flags=re.IGNORECASE):
                            try:
                                fixed = _cleanup_translated_text(_translate_preserve_exact_lines(fixed))
                            except Exception:
                                pass

                        if fixed != raw:
                            _set_paragraph_text_preserve_runs(para, fixed)
                            touched += 1
        return touched

    def _normalize_profile_tab_leaders(document):
        touched = 0
        key_re = re.compile(r'^(\s*)(student\s*id|email|class)\s*[:\-]?\s*(.*)$', flags=re.IGNORECASE)
        for para in iter_all_paragraphs(document):
            raw = ''.join((r.text or '') for r in para.runs)
            if not raw.strip():
                continue
            low = raw.lower()
            if 'email' in low and 'class' in low:
                continue
            if '\t' in raw and not re.search(r'\.{3,}', raw):
                continue

            m = key_re.match(raw.strip())
            if not m:
                continue

            label = m.group(2)
            rest = m.group(3) or ''
            value = re.sub(r'^[\.\-_:\s]+', '', rest).strip()
            if not value:
                continue

            new_text = f"{label.title()}:\t{value}"
            if new_text != raw:
                _set_paragraph_text_preserve_runs(para, new_text)
                touched += 1

            try:
                para.paragraph_format.tab_stops.add_tab_stop(Inches(5.6), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
            except Exception:
                pass
        return touched

    def _shrink_table_cell_fonts(document):
        touched = 0
        max_len = int(os.getenv('DOCX_TABLE_SHRINK_LEN', '140'))
        min_pt = float(os.getenv('DOCX_TABLE_SHRINK_MIN_PT', '8'))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = "".join(p.text or "" for p in cell.paragraphs).strip()
                    if not cell_text:
                        continue
                    if len(cell_text) < max_len:
                        continue
                    scale = max(0.7, min(1.0, float(max_len) / float(max(1, len(cell_text)))))
                    adjusted = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            try:
                                if run.font.size is None:
                                    continue
                                new_pt = max(min_pt, float(run.font.size.pt) * scale)
                                run.font.size = Pt(new_pt)
                                adjusted = True
                            except Exception:
                                continue
                    if adjusted:
                        touched += 1
        return touched

    def _relax_table_row_heights(document):
        """Allow rows to grow for bilingual text instead of clipping fixed-height cells."""
        relaxed = 0
        for table in document.tables:
            try:
                table.autofit = True
            except Exception:
                pass
            for row in table.rows:
                changed = False
                try:
                    if row.height is not None:
                        row.height = None
                        changed = True
                except Exception:
                    pass
                try:
                    if row.height_rule != WD_ROW_HEIGHT_RULE.AT_LEAST:
                        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                        changed = True
                except Exception:
                    pass
                try:
                    tr_pr = row._tr.get_or_add_trPr()
                    for tr_h in list(tr_pr.xpath('./*[local-name()="trHeight"]')):
                        tr_pr.remove(tr_h)
                        changed = True
                except Exception:
                    pass
                if changed:
                    relaxed += 1
        return relaxed

    def _center_inline_images(document):
        centered = 0
        for para in iter_all_paragraphs(document):
            has_drawing = False
            try:
                for run in para.runs:
                    dr = run._element.xpath('.//*[local-name()="drawing"]')
                    if dr:
                        has_drawing = True
                        break
            except Exception:
                has_drawing = False
            if not has_drawing:
                continue

            # Only center image-only paragraphs; avoid centering mixed text paragraphs.
            try:
                para_text = merged_paragraph_plain(para).strip()
            except Exception:
                para_text = ""
            if para_text:
                continue

            try:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                centered += 1
            except Exception:
                pass
        return centered

    def _force_remaining_phrase_fixes(document):
        touched = 0
        replacements = [
            (r"\bKHOA\s+CÔNG\s+NGHỆ\s+THÔNG\s+TIN\b", "FACULTY OF INFORMATION TECHNOLOGY"),
            (r"\bKHOA\s+CONG\s+NGHE\s+THONG\s+TIN\b", "FACULTY OF INFORMATION TECHNOLOGY"),
            (r"Gửi\s+lại\s+phiếu\s+đăng\s+ký\s+qua\s+Email\s*:", "Resubmit the registration form via Email:"),
            (r"Gui\s+lai\s+phieu\s+dang\s+ky\s+qua\s+Email\s*:", "Resubmit the registration form via Email:"),
        ]

        for para in iter_all_paragraphs(document):
            if _paragraph_has_drawing(para):
                continue
            raw = ''.join((r.text or '') for r in para.runs)
            if not raw.strip():
                continue
            fixed = raw
            for pat, rep in replacements:
                fixed = re.sub(pat, rep, fixed, flags=re.IGNORECASE)
            fixed = _cleanup_translated_text(fixed)
            if fixed != raw:
                _set_paragraph_text_preserve_runs(para, fixed)
                touched += 1
        return touched

    def translate_paragraph_runs(paragraph, idx=None, total=None):
        runs = list(paragraph.runs)
        if not runs:
            return
        if _paragraph_has_drawing(paragraph):
            return

        if bool(from_pdf) or bi_mode in ('inline', 'newline'):
            apply_docx_paragraph_spacing(paragraph)
            runs = list(paragraph.runs)
        paragraph_text = _paragraph_source_text(paragraph)
        if not paragraph_text.strip():
            return

        if bi_mode in ("inline", "newline") and _paragraph_is_multi_reference_list(paragraph_text):
            _handle_multi_reference_bilingual(paragraph, paragraph_text)
            return

        if bi_mode == 'newline':
            if _is_structural_text(paragraph_text):
                return
            try:
                translated_para = _translate_body_paragraph_text(paragraph_text)
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Translator failed for paragraph: {e}")
                if api_only:
                    raise
                translated_para = paragraph_text
            if (translated_para or '').strip() and (translated_para or '').strip() != paragraph_text.strip():
                if _prefer_linebreak_for_newline(paragraph):
                    _append_translation_linebreak(paragraph, translated_para, italic=False)
                else:
                    new_p = _insert_paragraph_after(
                        paragraph,
                        translated_para,
                        italic=False,
                        clear_first_line=False,
                    )
                    if new_p is None:
                        _append_translation_linebreak(paragraph, translated_para, italic=False)
                    else:
                        try:
                            _seen_para_elems.add(id(new_p))
                        except Exception:
                            pass
        elif bi_mode == 'inline':
            try:
                translated_para = _translate_body_paragraph_text(paragraph_text)
            except ProviderRateLimitError:
                raise
            except Exception as e:
                print(f"Translator failed for paragraph: {e}")
                if api_only:
                    raise
                translated_para = paragraph_text
            t = (translated_para or '').strip()
            if t and t != paragraph_text.strip():
                d = service._normalize_bilingual_delimiter(bilingual_delimiter)
                _append_inline_bilingual(paragraph, t, d)
        else:
            _translate_format_groups(paragraph, _translate_body_paragraph_text)

        if progress_callback and idx is not None and total is not None:
            progress_callback(10 + int((idx / total) * 70), f"Translating paragraph {idx+1}/{total}")

    _seen_para_elems = set()

    def _seen_or_mark(paragraph):
        try:
            key = id(paragraph._element)
        except Exception:
            key = id(paragraph)
        if key in _seen_para_elems:
            return True
        _seen_para_elems.add(key)
        return False

    def _paragraph_text(paragraph):
        try:
            return _paragraph_source_text(paragraph)
        except Exception:
            return ''

    def _norm_ws(s):
        return re.sub(r'\s+', ' ', (s or '').strip())

    def _looks_source_for_target(raw_text):
        t_root = str(target_lang).strip().lower()
        txt = (raw_text or '').strip()
        if not txt:
            return False
        if t_root.startswith('en'):
            try:
                return bool(service._looks_vietnamese_like_text(txt))
            except Exception:
                return bool(re.search(r'[à-ỹđ]', txt, flags=re.IGNORECASE))
        return True

    def _rescue_table_cells_all_modes():
        touched = 0
        d = service._normalize_bilingual_delimiter(bilingual_delimiter)
        target_root = str(target_lang).strip().lower()

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paras = list(cell.paragraphs)
                    i = 0
                    while i < len(paras):
                        p = paras[i]
                        raw = _paragraph_text(p)
                        core = (raw or '').strip()
                        if not core or _paragraph_has_drawing(p) or _is_structural_text(core):
                            i += 1
                            continue
                        should_retry = _looks_source_for_target(core)
                        # Fallback detector for EN target: keep retrying table text that still looks Vietnamese.
                        if (not should_retry) and target_root.startswith('en'):
                            should_retry = bool(re.search(r'[à-ỹđ]', core, flags=re.IGNORECASE))
                        if not should_retry:
                            i += 1
                            continue

                        try:
                            translated = _translate_preserve_exact_lines(core)
                        except ProviderRateLimitError:
                            raise
                        except Exception:
                            i += 1
                            continue

                        t = (translated or '').strip()
                        if not t or _norm_ws(t) == _norm_ws(core):
                            i += 1
                            continue

                        if bi_mode == 'inline':
                            if d in core:
                                i += 1
                                continue
                            _append_inline_bilingual(p, t, d)
                            touched += 1
                            i += 1
                            continue

                        if bi_mode == 'newline':
                            # newline mode: avoid duplicate insertion if the next paragraph is already the translation.
                            has_next_same = False
                            if i + 1 < len(paras):
                                nxt = (_paragraph_text(paras[i + 1]) or '').strip()
                                if _norm_ws(nxt) == _norm_ws(t):
                                    has_next_same = True
                                elif nxt and (not _is_structural_text(nxt)) and (not _looks_source_for_target(nxt)):
                                    # Next line already looks like translated target text.
                                    has_next_same = True

                            if not has_next_same:
                                if _prefer_linebreak_for_newline(p):
                                    _append_translation_linebreak(p, t, italic=False)
                                else:
                                    new_p = _insert_paragraph_after(
                                        p,
                                        t,
                                        italic=False,
                                        clear_first_line=False,
                                    )
                                    if new_p is None:
                                        _append_translation_linebreak(p, t, italic=False)
                                touched += 1
                                paras = list(cell.paragraphs)
                                i += 2
                                continue

                            i += 1
                            continue

                        # normal mode: force-replace untranslated source text inside table cells.
                        _set_paragraph_text_preserve_runs(p, translated)
                        touched += 1
                        i += 1

        return touched

    def _sanitize_from_pdf_paragraphs(document):
        """Normalize tab-driven spacing artifacts commonly emitted by pdf2docx."""
        fixed_count = 0
        tab_paragraphs = 0

        try:
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WPA
            justify_values = {
                int(v)
                for v in (
                    getattr(_WPA, 'JUSTIFY', None),
                    getattr(_WPA, 'DISTRIBUTE', None),
                    getattr(_WPA, 'THAI_JUSTIFY', None),
                )
                if v is not None
            }
            center_like_values = {
                int(v)
                for v in (
                    getattr(_WPA, 'CENTER', None),
                    getattr(_WPA, 'RIGHT', None),
                )
                if v is not None
            }
            left_value = getattr(_WPA, 'LEFT', 0)
        except Exception:
            justify_values = {3, 4, 7, 8, 9}
            center_like_values = {1, 2}
            left_value = 0

        try:
            from docx.oxml.ns import qn as _qn
        except Exception:
            _qn = None

        for para in iter_all_paragraphs(document):
            try:
                runs = list(para.runs)
            except Exception:
                continue
            if not runs:
                continue

            local_changed = False
            local_has_tab = False

            for run in runs:
                try:
                    raw = run.text or ''
                except Exception:
                    continue
                if not raw:
                    continue
                cleaned = raw.replace('\u00A0', ' ').replace('\t', ' ').replace('\u00AD', '')
                if '\t' in raw:
                    local_has_tab = True
                if cleaned != raw:
                    run.text = cleaned
                    local_changed = True

            if local_changed:
                apply_docx_paragraph_spacing(para)
                fixed_count += 1

            if local_has_tab:
                tab_paragraphs += 1
                if _qn is not None:
                    try:
                        ppr = para._element.find(_qn('w:pPr'))
                        if ppr is not None:
                            tabs = ppr.find(_qn('w:tabs'))
                            if tabs is not None:
                                ppr.remove(tabs)
                    except Exception:
                        pass

            if str(os.getenv("PDF_DOCX_SANITIZE_FORCE_LEFT", "0")).strip().lower() in (
                "1",
                "true",
                "yes",
                "on",
            ):
                try:
                    a = para.alignment
                    a_val = int(a) if a is not None else None
                    para_text = merged_paragraph_plain(para).strip()
                    word_count = len(re.findall(r'\w+', para_text, flags=re.UNICODE))
                    looks_center_artifact = bool(para_text) and (
                        len(para_text) >= 60
                        or word_count >= 10
                        or (
                            len(para_text) >= 30
                            and para_text.endswith(('.', '!', '?', ':', ';'))
                        )
                    )

                    should_force_left = (
                        a_val in justify_values
                        or (
                            a_val in center_like_values
                            and looks_center_artifact
                            and not _is_heading_paragraph(para)
                            and not _is_in_table_cell(para)
                            and not _paragraph_has_drawing(para)
                        )
                    )

                    if should_force_left:
                        para.alignment = left_value
                        fixed_count += 1
                except Exception:
                    pass

        return fixed_count, tab_paragraphs

    if bool(from_pdf) and not preserve_pdf_lines:
        try:
            fixed_count, tab_paragraphs = _sanitize_from_pdf_paragraphs(doc)
            if progress_callback and (fixed_count > 0 or tab_paragraphs > 0):
                progress_callback(
                    9,
                    f"DOCX from_pdf cleanup: fixed {fixed_count} paragraphs, tab-paragraphs {tab_paragraphs}",
                )
        except Exception:
            pass

    try:
        _normalize_toc_hyperlinks(doc)
    except Exception:
        pass

    paragraphs = [p for p in doc.paragraphs]
    from concurrent.futures import as_completed

    total_work = 0
    completed = 0

    body_paras = []
    for para in paragraphs:
        if _seen_or_mark(para):
            continue
        paragraph_text = _paragraph_source_text(para)
        if not paragraph_text.strip():
            continue
        body_paras.append(para)
    total_work = max(1, len(body_paras))

    # === BATCH PRE-TRANSLATION for speed optimization ===
    # Collect all unique translatable texts from body + tables + headers/footers,
    # batch them with <<<S>>> separator into fewer API calls, cache results.
    _translation_cache = {}
    _all_texts = []
    _seen_texts = set()

    def _queue_text(para):
        """Collect paragraph text for batch pre-translation."""
        if _paragraph_has_drawing(para):
            return
        text = _normalize_from_pdf_soft_breaks(_paragraph_source_text(para)).strip()
        if text and not _is_structural_text(text) and text not in _seen_texts:
            if _should_skip_translation(text) or _is_pdf_noise_text(text):
                return
            _seen_texts.add(text)
            masked, placeholders = _guard_mask_tokens(text) if _guard_enabled else (text, {})
            _all_texts.append({"raw": text, "masked": masked, "placeholders": placeholders})

    for para in body_paras:
        _queue_text(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _queue_text(p)
    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                _queue_text(p)
            for p in section.footer.paragraphs:
                _queue_text(p)
    except Exception:
        pass

    if _all_texts:
        _BATCH_CHARS = int(os.getenv('DOCX_BATCH_CHAR_LIMIT', '3000'))
        _BATCH_COUNT = int(os.getenv('DOCX_BATCH_COUNT_LIMIT', '25'))
        _SEP = '\n<<<S>>>\n'
        _i = 0
        _total_texts = len(_all_texts)
        while _i < _total_texts:
            _batch_entries = []
            _batch_masked = []
            _chars = 0
            while _i < _total_texts and len(_batch_entries) < _BATCH_COUNT and _chars < _BATCH_CHARS:
                entry = _all_texts[_i]
                _batch_entries.append(entry)
                _batch_masked.append(entry.get("masked") or "")
                _chars += len(entry.get("masked") or "")
                _i += 1
            if not _batch_entries:
                break
            try:
                if len(_batch_entries) == 1:
                    entry = _batch_entries[0]
                    _r = service._translate_with_retry(entry.get("masked") or "", target_lang, context='document_docx_batch')
                    restored = _guard_restore_tokens(_cleanup_translated_text(_r), entry.get("placeholders") or {})
                    _translation_cache[entry.get("raw")] = restored
                else:
                    _joined = _SEP.join(_batch_masked)
                    _r = service._translate_with_retry(_joined, target_lang, context='document_docx_batch')
                    _parts = [p.strip() for p in _r.split('<<<S>>>')]
                    if len(_parts) == len(_batch_entries):
                        for entry, _dst in zip(_batch_entries, _parts):
                            restored = _guard_restore_tokens(_cleanup_translated_text(_dst), entry.get("placeholders") or {})
                            _translation_cache[entry.get("raw")] = restored
                    else:
                        # Separator count mismatch — fallback to individual translation
                        for entry in _batch_entries:
                            try:
                                _r2 = service._translate_with_retry(entry.get("masked") or "", target_lang, context='document_docx_batch')
                                restored = _guard_restore_tokens(
                                    _cleanup_translated_text(_r2),
                                    entry.get("placeholders") or {},
                                )
                                _translation_cache[entry.get("raw")] = restored
                            except ProviderRateLimitError:
                                raise
                            except Exception:
                                pass
            except ProviderRateLimitError:
                raise
            except Exception:
                # Batch failed — fallback to individual translation
                for entry in _batch_entries:
                    try:
                        _r2 = service._translate_with_retry(entry.get("masked") or "", target_lang, context='document_docx_batch')
                        restored = _guard_restore_tokens(
                            _cleanup_translated_text(_r2),
                            entry.get("placeholders") or {},
                        )
                        _translation_cache[entry.get("raw")] = restored
                    except ProviderRateLimitError:
                        raise
                    except Exception:
                        pass
            if progress_callback:
                progress_callback(
                    10 + int((min(_i, _total_texts) / max(1, _total_texts)) * 70),
                    f"Batch translating {min(_i, _total_texts)}/{_total_texts}",
                )

    for para in body_paras:
        try:
            if _paragraph_has_drawing(para):
                continue
            if bool(from_pdf) or bi_mode in ('inline', 'newline'):
                apply_docx_paragraph_spacing(para)
            paragraph_text_src = _paragraph_source_text(para)

            if _is_pdf_noise_text(paragraph_text_src.strip()):
                completed += 1
                if progress_callback:
                    progress_callback(
                        10 + int((completed / total_work) * 70),
                        f"Translating paragraph {completed}/{total_work}",
                    )
                continue

            if bi_mode in ("inline", "newline") and _paragraph_is_multi_reference_list(paragraph_text_src):
                _handle_multi_reference_bilingual(para, paragraph_text_src)
                completed += 1
                if progress_callback:
                    progress_callback(
                        10 + int((completed / total_work) * 70),
                        f"Translating paragraph {completed}/{total_work}",
                    )
                continue

            if bi_mode == 'inline':
                translated = _translate_body_paragraph_text(paragraph_text_src)
                t = (translated or '').strip()
                if t and t != paragraph_text_src.strip():
                    d = service._normalize_bilingual_delimiter(bilingual_delimiter)
                    _append_inline_bilingual(para, t, d)
            elif bi_mode == 'newline':
                if _is_structural_text(paragraph_text_src):
                    pass
                else:
                    translated = _translate_body_paragraph_text(paragraph_text_src)
                    if (translated or '').strip() and (translated or '').strip() != paragraph_text_src.strip():
                        if _prefer_linebreak_for_newline(para):
                            _append_translation_linebreak(para, translated, italic=False)
                        else:
                            new_p = _insert_paragraph_after(
                                para,
                                translated,
                                italic=False,
                                clear_first_line=False,
                            )
                            if new_p is None:
                                _append_translation_linebreak(para, translated, italic=False)
                            else:
                                try:
                                    _seen_para_elems.add(id(new_p))
                                except Exception:
                                    pass
            else:
                _translate_format_groups(para, _translate_body_paragraph_text)
        except ProviderRateLimitError:
            print("Provider rate limit detected during paragraph processing, aborting job.")
            raise
        except Exception as e:
            print(f"Paragraph translation failed: {e}")
            if api_only:
                raise
        completed += 1
        if progress_callback:
            progress_callback(
                10 + int((completed / total_work) * 70),
                f"Translating paragraph {completed}/{total_work}",
            )

    for table in doc.tables:
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                cell = table.rows[r].cells[c]
                for p_idx, p in enumerate(cell.paragraphs):
                    if _seen_or_mark(p):
                        continue
                    translate_paragraph_runs(p, p_idx, len(cell.paragraphs))

    try:
        for section in doc.sections:
            header = section.header
            for p_idx, p in enumerate(header.paragraphs):
                if _seen_or_mark(p):
                    continue
                translate_paragraph_runs(p, p_idx, len(header.paragraphs))
            footer = section.footer
            for p_idx, p in enumerate(footer.paragraphs):
                if _seen_or_mark(p):
                    continue
                translate_paragraph_runs(p, p_idx, len(footer.paragraphs))
    except Exception:
        pass

    try:
        rescued_table_cells = _rescue_table_cells_all_modes()
        if progress_callback and rescued_table_cells > 0:
            progress_callback(81, f"DOCX table rescue: {rescued_table_cells} cells")
    except ProviderRateLimitError:
        raise
    except Exception:
        pass

    try:
        if bool(from_pdf) or bi_mode in ('inline', 'newline'):
            relaxed_rows = _relax_table_row_heights(doc)
            if progress_callback and relaxed_rows > 0:
                progress_callback(82, f"DOCX table layout: relaxed {relaxed_rows} row heights")
    except Exception:
        pass

    if ocr_images and service.ocr_translate_overlay:
        if progress_callback:
            progress_callback(82, "OCR images in DOCX...")

        protected_image_partnames = _collect_header_footer_image_partnames(doc)

        paras_to_scan = iter_all_paragraphs(doc)
        total_paras = len(paras_to_scan) or 1
        images_found = 0
        ocr_attempted = 0
        ocr_success = 0
        ocr_disabled = False

        image_replacements = {}
        ocr_export_entries = []
        text_insert_entries = []
        text_replace_entries = []

        def _is_probably_logo_or_nontext(ocr_text: str) -> bool:
            raw = (ocr_text or '').strip()
            if not raw:
                return True
            words = re.findall(r'\w+', raw, flags=re.UNICODE)
            if len(words) <= 2 and len(raw) < 24:
                return True
            return False

        for idx, para in enumerate(paras_to_scan):
            if ocr_disabled:
                break
            rids = paragraph_image_rids(para)
            if not rids:
                continue
            for rid in rids:
                img_part = rid_to_image_part(para, rid)
                if not img_part:
                    continue
                partname = str(getattr(img_part, 'partname', '') or '').lstrip('/')
                if partname and partname in protected_image_partnames:
                    continue
                try:
                    blob = getattr(img_part, 'blob', None)
                    if not blob:
                        continue

                    images_found += 1

                    ext = image_part_ext(img_part)
                    tmp_name = f"docx_img_{uuid.uuid4().hex}{ext}"
                    tmp_path = os.path.join(service.upload_folder, tmp_name)
                    with open(tmp_path, 'wb') as f:
                        f.write(blob)

                    ocr_attempted += 1
                    try:
                        ocr_text, translated_text, png_bytes, ai_recommended_mode = service.ocr_translate_overlay(
                            tmp_path,
                            'auto',
                            target_lang,
                            ocr_langs,
                        )
                    finally:
                        try:
                            os.remove(tmp_path)
                        except Exception:
                            pass

                    if not ocr_text or not str(ocr_text).strip():
                        continue

                    if _is_probably_logo_or_nontext(ocr_text):
                        continue

                    per_mode = mode
                    if mode == 'auto':
                        per_mode = _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode)

                    print(f"  [IMAGE #{images_found}] OCR={len(ocr_text)}chars, AI_class={ai_recommended_mode}, per_mode={per_mode}")

                    try:
                        if per_mode in ('text', 'both'):
                            partname = str(getattr(img_part, 'partname', '') or '').lstrip('/')
                            ocr_export_entries.append({
                                'image': partname or '(embedded image)',
                                'ocr_text': (ocr_text or '').strip(),
                                'translated_text': (translated_text or '').strip(),
                            })
                            normalized_translated = _normalize_ocr_text_for_docx((translated_text or '').strip())

                            if per_mode == 'text' and mode == 'text':
                                text_replace_entries.append((para, rid, normalized_translated))
                            else:
                                text_insert_entries.append((para, normalized_translated))
                    except Exception:
                        pass

                    if per_mode in ('image', 'both') and png_bytes and len(png_bytes) > 100:
                        try:
                            if partname:
                                new_bytes = _overlay_bytes_to_original_format(png_bytes, ext)
                                image_replacements[partname] = new_bytes
                                ocr_success += 1
                        except Exception:
                            continue
                except Exception as e:
                    msg = str(e).lower()
                    if ('tesseract' in msg and ('not installed' in msg or 'path' in msg)) or \
                       ('ocr unavailable' in msg):
                        ocr_disabled = True
                        if progress_callback:
                            progress_callback(85, "Skipping DOCX image OCR (OCR not available)")
                        break
                    if 'ai provider' in msg and ('not configured' in msg or 'rate' in msg):
                        ocr_disabled = True
                        if progress_callback:
                            progress_callback(85, f"Skipping DOCX image OCR: {e}")
                        break
                    print(f"DOCX image OCR error (continuing): {e}")
                    continue

            if progress_callback and (idx % 10 == 0):
                progress_callback(82 + int((idx / total_paras) * 10), f"OCR scanning {idx+1}/{total_paras}")

        try:
            pkg_scan_enabled = str(os.getenv('DOCX_OCR_PACKAGE_SCAN', '0')).strip().lower() in ('1', 'true', 'yes', 'on')
            if pkg_scan_enabled and (not ocr_disabled) and mode in ('image', 'both', 'auto'):
                pkg = getattr(getattr(doc, 'part', None), 'package', None)
                parts = list(getattr(pkg, 'parts', []) or [])

                extra_attempted = 0
                extra_replaced = 0
                for part in parts:
                    try:
                        ct = str(getattr(part, 'content_type', '') or '').lower()
                        if not ct.startswith('image/'):
                            continue
                        partname = str(getattr(part, 'partname', '') or '').lstrip('/')
                        if not partname:
                            continue
                        if partname in protected_image_partnames:
                            continue
                        if partname in image_replacements:
                            continue
                        blob = getattr(part, 'blob', None)
                        if not blob:
                            continue

                        ext = image_part_ext(part)
                        tmp_name = f"docx_img_pkg_{uuid.uuid4().hex}{ext}"
                        tmp_path = os.path.join(service.upload_folder, tmp_name)
                        with open(tmp_path, 'wb') as f:
                            f.write(blob)

                        extra_attempted += 1
                        try:
                            ocr_text, translated_text, png_bytes, ai_recommended_mode = service.ocr_translate_overlay(
                                tmp_path,
                                'auto',
                                target_lang,
                                ocr_langs,
                            )
                        finally:
                            try:
                                os.remove(tmp_path)
                            except Exception:
                                pass

                        if not ocr_text or not str(ocr_text).strip():
                            continue

                        if _is_probably_logo_or_nontext(ocr_text):
                            continue

                        per_mode = mode
                        if mode == 'auto':
                            per_mode = _auto_pick_mode(ocr_text, translated_text, ai_recommended_mode)

                        if per_mode in ('image', 'both') and png_bytes and len(png_bytes) > 100:
                            try:
                                new_bytes = _overlay_bytes_to_original_format(png_bytes, ext)
                                image_replacements[partname] = new_bytes
                                extra_replaced += 1
                            except Exception:
                                continue
                    except Exception as e:
                        msg = str(e).lower()
                        if ('tesseract' in msg and ('not installed' in msg or 'path' in msg)) or \
                           ('ocr unavailable' in msg):
                            ocr_disabled = True
                            if progress_callback:
                                progress_callback(85, "Skipping DOCX image OCR (OCR not available)")
                            break
                        if 'ai provider' in msg and ('not configured' in msg or 'rate' in msg):
                            ocr_disabled = True
                            if progress_callback:
                                progress_callback(85, f"Skipping DOCX image OCR: {e}")
                            break
                        continue

                if progress_callback and (extra_attempted or extra_replaced):
                    progress_callback(
                        92,
                        f"DOCX OCR (package scan): attempted={extra_attempted}, replaced={extra_replaced}",
                    )
        except Exception:
            pass

        if progress_callback:
            if images_found <= 0:
                progress_callback(92, "DOCX OCR: no embedded images found")
            else:
                progress_callback(
                    92,
                    f"DOCX OCR: found={images_found}, attempted={ocr_attempted}, replaced={ocr_success}",
                )

        try:
            if text_replace_entries:
                for para, rid, trans_text in text_replace_entries:
                    if not trans_text or not trans_text.strip():
                        continue
                    try:
                        replace_image_with_text(para, rid, trans_text)
                    except Exception:
                        continue
        except Exception:
            pass

    try:
        if ocr_images and 'text_insert_entries' in locals() and text_insert_entries:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn

            for para, trans_text in reversed(text_insert_entries):
                if not trans_text or not trans_text.strip():
                    continue
                try:
                    text_paragraphs = [p.strip() for p in trans_text.split('\n') if p.strip()]
                    if not text_paragraphs:
                        text_paragraphs = [trans_text.strip()]

                    for t_idx, t_para in enumerate(reversed(text_paragraphs)):
                        new_p = OxmlElement('w:p')

                        run = OxmlElement('w:r')

                        t_el = OxmlElement('w:t')
                        t_el.set(_qn('xml:space'), 'preserve')
                        t_el.text = t_para
                        run.append(t_el)
                        new_p.append(run)

                        para._element.addnext(new_p)
                except Exception:
                    continue
    except Exception:
        pass

    try:
        if bi_mode in ('inline', 'newline'):
            link_count = 0
            leader_count = 0
            table_count = 0
            img_count = 0
            forced_count = 0
            if progress_callback:
                progress_callback(96, "DOCX targeted fixes skipped in bilingual mode")
        else:
            link_count = _strip_all_hyperlinks(doc)
            leader_count = _normalize_profile_tab_leaders(doc)
            table_count = _normalize_table_layout_and_text(doc)
            img_count = _center_inline_images(doc)
            forced_count = _force_remaining_phrase_fixes(doc)
            if progress_callback:
                progress_callback(
                    96,
                    (
                        f"DOCX targeted fixes: links={link_count}, "
                        f"leaders={leader_count}, table={table_count}, images={img_count}, forced={forced_count}"
                    ),
                )
    except Exception as e:
        if progress_callback:
            progress_callback(96, f"DOCX targeted fixes skipped: {e}")

    try:
        if str(os.getenv('DOCX_TABLE_SHRINK', '0')).strip().lower() in ('1', 'true', 'yes', 'on'):
            shrink_count = _shrink_table_cell_fonts(doc)
            if progress_callback:
                progress_callback(97, f"DOCX layout recovery: shrunk {shrink_count} table cells")
    except Exception:
        pass

    output_filename = f"translated_{os.path.basename(file_path)}"
    if not output_filename.lower().endswith('.docx'):
        output_filename += '.docx'
    output_path = os.path.join(service.download_folder, output_filename)

    doc.save(output_path)

    try:
        if ocr_images and mode in ('image', 'both', 'auto') and 'image_replacements' in locals() and image_replacements:
            if progress_callback:
                progress_callback(93, "Applying translated overlays to DOCX images...")
            tmp_out = output_path + ".tmp"
            with zipfile.ZipFile(output_path, 'r') as zin, zipfile.ZipFile(tmp_out, 'w') as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    repl = image_replacements.get(item.filename)
                    if repl is not None:
                        data = repl
                    zout.writestr(item, data)
            try:
                os.replace(tmp_out, output_path)
            except Exception:
                try:
                    os.remove(output_path)
                except Exception:
                    pass
                os.rename(tmp_out, output_path)
    except Exception as e:
        if progress_callback:
            progress_callback(94, f"DOCX image overlay patch failed: {e}")

    try:
        docx.Document(output_path)
    except Exception as e:
        if progress_callback:
            progress_callback(95, "DOCX validation failed, writing fallback text file")
        fallback_filename = output_filename
        if not fallback_filename.lower().endswith('.txt'):
            fallback_filename += '.txt'
        fallback_path = os.path.join(service.download_folder, fallback_filename)
        lines = []
        for p in doc.paragraphs:
            lines.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    lines.append(cell.text)
        with open(fallback_path, 'w', encoding='utf-8') as f:
            f.write("NOTE: DOCX creation failed on server. Showing plain text fallback below.\n\n")
            f.write('\n'.join(lines))
        output_path = fallback_path

    if progress_callback:
        progress_callback(100, "Completed")
    return output_path
