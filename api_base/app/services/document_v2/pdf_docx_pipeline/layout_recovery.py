"""Recover DOCX layout after PDF->DOCX translation.

Keeps original spacing/indent/alignment as closely as possible:
- Title block: copy from source (centered author/affiliation)
- Abstract + Keywords: inset column, justify, merge pdf2docx line fragments
- Section 1 body: body template from section 2 reference
- Section 2+: preserve full paragraph properties from source (spacing, firstLine)
"""

from __future__ import annotations

import copy
import os
import re
from typing import Any, Dict, Iterable, Iterator, List, Optional, Tuple

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

_ABSTRACT_EXTRA_INDENT_TWIPS = 566

_ABSTRACT_START_RE = re.compile(
    r"^(abstract|abstrakt|abstrak|opsomming|p[eë]rmbledhje|"
    r"t[oóô]m\s*t[aắ]t|t[oóô]mt[aắ]t|zusammenfassung|r[eé]sum[eé])\b",
    re.IGNORECASE,
)
_KEYWORDS_LINE_RE = re.compile(
    r"^(keywords?|keyword|t[uừ]\s*kh[oó]a|t[uừ]kh[oó]a|"
    r"sleutelwoorde?|fjal[eë]t\s*ky[çc]e|schl[uü]sselw[oö]rter)\b",
    re.IGNORECASE,
)
_SECTION1_HEAD_RE = re.compile(
    r"^1\s+(introduction|inleiding|hyrje|gi[oớ]i\s*th[iệ]u|gi[oớ]ith[iệ]u|einleitung)\b",
    re.IGNORECASE,
)
_FALSE_SECTION2_RE = re.compile(r"^2\s+F\.\s", re.IGNORECASE)
_RUNNING_HEADER_RE = re.compile(
    r"^\d+\s+[A-ZÀ-Ỹ]\.\s+(?:Author|Skrywer)\b",
    re.IGNORECASE,
)
_PAGE_NO_ONLY_RE = re.compile(r"^\d{1,4}$")
_PDF_ARTIFACT_RE = re.compile(
    r"^(?:__?\s*[\wÀ-Ỹ]{2,10}\s*[_\s]*\d+\s*__?|[\s_,.|/-]{5,})$",
    re.IGNORECASE,
)
_AFFILIATION_LINE_RE = re.compile(
    r"^\d+\s+(faculty|department|university|springer|institute|fakulteit|khoa|tr[uườ]ng)\b",
    re.IGNORECASE,
)
_SECTION2_HEAD_RE = re.compile(r"^2\s+\S", re.IGNORECASE)
_ASSIGNMENT_HEAD_RE = re.compile(r"^BÀI\s*\d+\s*[:：]", re.IGNORECASE)


def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in ("1", "true", "yes", "on")


def normalize_bilingual_mode(mode: Optional[str]) -> str:
    """Map frontend/API bilingual_mode values to none|inline|newline."""
    bi = (str(mode or "").strip().lower() or "none")
    if bi in ("preserve_layout", "inline", "lien_ke", "adjacent", "side_by_side"):
        return "inline"
    if bi in ("line_by_line", "newline", "xuong_dong", "stacked"):
        return "newline"
    if bi in ("none", "off", "0", "false"):
        return "none"
    return "none"


def resolve_pdf_layout_mode(
    analysis: Optional[Dict[str, Any]] = None,
    paras: Optional[List] = None,
) -> str:
    """Return academic | conservative.

    auto (default): detect journal/paper structure vs general PDF.
    """
    raw = (os.getenv("PDF_DOCX_LAYOUT_MODE") or "auto").strip().lower()
    if raw in ("academic", "journal", "paper"):
        return "academic"
    if raw in ("conservative", "general", "regular", "plain"):
        return "conservative"

    analysis = analysis or {}
    score = 0
    if analysis.get("has_references"):
        score += 2
    if analysis.get("has_formulas"):
        score += 1
    if analysis.get("has_multiple_columns"):
        score += 1
    if int(analysis.get("pages") or 0) >= 5:
        score += 1
    if int(analysis.get("text_chars") or 0) >= 6000:
        score += 1

    if paras:
        abs_i = _find_abstract_start(paras)
        sec1 = _find_section_one_start(paras)
        sec2 = _find_section_two_start(paras)
        if abs_i is not None:
            score += 3
        if sec1 is not None:
            score += 2
        if sec2 is not None:
            score += 2

    return "academic" if score >= 4 else "conservative"


def uses_regional_layout(layout_mode: str) -> bool:
    """Regional title/abstract/body profiles are for academic PDFs only."""
    if layout_mode == "conservative":
        return False
    if layout_mode == "academic":
        return _env_bool("PDF_DOCX_REGIONAL_LAYOUT", True)
    return _env_bool("PDF_DOCX_REGIONAL_LAYOUT", True)


def should_preserve_pdf_lines(layout_mode: Optional[str] = None) -> bool:
    """Keep each PDF/DOCX line separate; do not collapse soft breaks or merge paragraphs."""
    if str(os.getenv("PDF_DOCX_PRESERVE_LINES", "1")).strip().lower() in ("0", "false", "no", "off"):
        return False
    if str(os.getenv("PDF_DOCX_COLLAPSE_SOFT_BREAKS", "0")).strip().lower() in ("1", "true", "yes", "on"):
        return False
    if layout_mode == "academic" and _env_bool("PDF_DOCX_ACADEMIC_MERGE_LINES", True):
        return False
    return True


def should_merge_pdf_fragments(layout_mode: str) -> bool:
    """Only merge pdf2docx line fragments for academic papers when explicitly allowed."""
    if not should_preserve_pdf_lines(layout_mode):
        return layout_mode == "academic" and _env_bool("PDF_DOCX_SANITIZE_MERGE_FRAGMENTS", True)
    return False


# SymbolMT / Wingdings private-use glyphs from pdf2docx
_SYMBOL_PUA_MAP = {
    "\uf020": " ",
    "\uf02a": "*",
    "\uf02b": "+",
    "\uf076": "➤",
    "\uf0a7": "▪",
    "\uf0b7": "•",
    "\uf0d8": "◆",
}


_LEADING_MARKER_RE = re.compile(r"^(\s*(?:[\uf02b\uf0b7\uf076\uf0a7+\-•]|o)\s*)")


def should_preserve_symbol_glyphs() -> bool:
    return _env_bool("PDF_DOCX_PRESERVE_SYMBOL_GLYPHS", False)


def _is_pua_symbol_char(ch: str) -> bool:
    return len(ch) == 1 and 0xF020 <= ord(ch) <= 0xF0FF


def _text_has_pua_symbol(text: str) -> bool:
    return any(_is_pua_symbol_char(c) for c in (text or ""))


def _is_symbol_glyph_text(text: str) -> bool:
    raw = text or ""
    if _text_has_pua_symbol(raw):
        without_pua = "".join(c for c in raw if not _is_pua_symbol_char(c))
        if without_pua.strip() and re.search(r"[\w\u00C0-\u1EF9]", without_pua, flags=re.UNICODE):
            return False
    stripped = raw.strip()
    if not stripped:
        return False
    if len(stripped) <= 2 and not re.search(r"[\w\u00C0-\u1EF9]", stripped, flags=re.UNICODE):
        return True
    return False


def _split_leading_glyph_prefix(text: str) -> tuple[str, str]:
    raw = text or ""
    m = _LEADING_MARKER_RE.match(raw)
    if m:
        return m.group(1), raw[m.end() :]
    if raw and _is_pua_symbol_char(raw[0]):
        idx = 1
        while idx < len(raw) and raw[idx] in " \t":
            idx += 1
        if idx < len(raw):
            return raw[:idx], raw[idx:]
    return "", raw


def normalize_pdf_symbol_chars(text: str, *, for_match: bool = False) -> str:
    """Map PUA glyphs to portable Unicode (optional; off when preserving original glyphs)."""
    if not text:
        return text
    if not for_match and should_preserve_symbol_glyphs():
        return text.replace("\uf020", " ")
    out = text
    for src, dst in _SYMBOL_PUA_MAP.items():
        out = out.replace(src, dst)
    return out


def normalize_pdf_symbol_chars_in_doc(doc: docx.Document) -> int:
    if should_preserve_symbol_glyphs():
        return ensure_symbol_font_runs_in_doc(doc)
    changed = 0
    for para in doc.paragraphs:
        if _is_in_table_cell(para):
            continue
        for run in para.runs:
            raw = run.text or ""
            if not raw:
                continue
            fixed = normalize_pdf_symbol_chars(raw)
            if fixed != raw:
                run.text = fixed
                changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        raw = run.text or ""
                        if not raw:
                            continue
                        fixed = normalize_pdf_symbol_chars(raw)
                        if fixed != raw:
                            run.text = fixed
                            changed += 1
    return changed


def _set_run_symbol_font(run) -> None:
    _set_run_font_name(run, "Symbol")


def repair_latin_runs_in_symbol_font(doc: docx.Document) -> int:
    """Symbol font on Latin letters renders as Greek glyphs — force body font."""
    changed = 0
    body_font = _detect_document_body_font(doc) or "Times New Roman"
    for para in iter_all_paragraphs(doc):
        for run in para.runs:
            raw = run.text or ""
            if not raw.strip():
                continue
            if not re.search(r"[A-Za-z\u00C0-\u1EF9]", raw, flags=re.UNICODE):
                continue
            fname = _read_run_font_name(run)
            if _is_auxiliary_pdf_font(fname) or fname == "Symbol":
                _set_run_font_name(run, body_font)
                changed += 1
    return changed


def _sanitize_rpr_body_font(rpr, para, *, body_font: Optional[str] = None):
    if rpr is None:
        try:
            return _fallback_body_run_rpr(para) if para is not None else None
        except Exception:
            return None
    cloned = copy.deepcopy(rpr)
    if body_font is None:
        try:
            doc = para.part.document if para is not None else None
        except Exception:
            doc = None
        body_font = (_detect_document_body_font(doc) if doc else None) or "Times New Roman"
    if _is_auxiliary_pdf_font(_rpr_font_name(cloned)):
        rf = cloned.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            cloned.insert(0, rf)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rf.set(qn(f"w:{attr}"), body_font)
    return cloned


def _portable_marker_text(marker: str) -> str:
    mapped = normalize_pdf_symbol_chars(marker or "")
    token = (mapped or "+").strip()
    return token if token else "+"


def strip_leading_markers_from_text(text: str) -> str:
    """Remove leading list markers when a separate marker run already exists."""
    out = text or ""
    return re.sub(r"^(?:\s*[+\-•\uf02b\uf0b7]\s*)+", "", out)


def _dedupe_leading_markers_in_paragraph(paragraph) -> int:
    """Collapse ++ / PUA++ duplicate bullets at paragraph start."""
    runs = [r for r in paragraph.runs if (r.text or "")]
    if not runs:
        return 0
    changed = 0
    marker_pat = re.compile(r"^[\uf02b\uf0b7\uf076\uf0a7+\-•o\s]*$", re.IGNORECASE)
    while len(runs) >= 2 and marker_pat.match((runs[0].text or "").strip()):
        nxt = (runs[1].text or "").lstrip()
        if nxt.startswith(("+", "-", "•", "\uf02b", "\uf0b7")):
            runs[1].text = re.sub(r"^(?:\s*[+\-•\uf02b\uf0b7]\s*)+", "", runs[1].text or "")
            changed += 1
        if marker_pat.match((runs[0].text or "").strip()) and marker_pat.match((runs[1].text or "").strip()):
            try:
                runs[1]._element.getparent().remove(runs[1]._element)
                runs.pop(1)
                changed += 1
                continue
            except Exception:
                break
        break
    if len(runs) >= 2 and marker_pat.match((runs[0].text or "").strip()):
        lead = (runs[1].text or "")
        if re.match(r"^\s*[+\-•\uf02b\uf0b7]", lead):
            runs[1].text = re.sub(r"^(?:\s*[+\-•\uf02b\uf0b7]\s*)+", "", lead)
            changed += 1
    return changed


def ensure_symbol_font_runs_in_doc(doc: docx.Document) -> int:
    """Keep PDF symbol glyphs in Symbol font; fix Latin text wrongly placed in Symbol runs."""
    changed = 0
    for para in iter_all_paragraphs(doc):
        for run in para.runs:
            raw = run.text or ""
            if not raw:
                continue
            fname = _read_run_font_name(run)
            has_word = bool(re.search(r"[\w\u00C0-\u1EF9]", raw, flags=re.UNICODE))
            if (_is_auxiliary_pdf_font(fname) or fname == "Symbol") and has_word and not _text_has_pua_symbol(raw):
                body = _detect_document_body_font(doc) or "Times New Roman"
                _set_run_font_name(run, body)
                changed += 1
                continue
            is_glyph = _text_has_pua_symbol(raw) or (
                _is_auxiliary_pdf_font(fname)
                and len(raw.strip()) <= 2
                and not has_word
            )
            if not is_glyph:
                continue
            cleaned = raw.replace("\uf020", " ")
            if cleaned != raw:
                run.text = cleaned
                changed += 1
            if _read_run_font_name(run) != "Symbol":
                _set_run_symbol_font(run)
                changed += 1
    return changed


def _split_pdf_soft_lines(text: str) -> List[str]:
    """Split pdf2docx merged paragraph text back into logical lines."""
    raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not raw.strip():
        return []
    lines: List[str] = []
    for chunk in raw.split("\n"):
        chunk = chunk.strip()
        if not chunk:
            continue
        parts = re.split(
            r"\t(?=[+\-*•➕▪])|\t(?=[\uF020-\uF0FF])|(?<=\.)\s+(?=\+)",
            chunk,
        )
        for part in parts:
            part = part.strip()
            if not should_preserve_symbol_glyphs():
                part = normalize_pdf_symbol_chars(part)
            if part:
                lines.append(part)
    return lines


def _clone_element(elem):
    if elem is None:
        return None
    return copy.deepcopy(elem)


def _first_run_rpr(p_el):
    for r_el in p_el.findall(qn("w:r")):
        rpr = r_el.find(qn("w:rPr"))
        if rpr is not None:
            return _clone_element(rpr)
    return None


def _build_paragraph_element(line_text: str, ppr, rpr) -> OxmlElement:
    new_p = OxmlElement("w:p")
    if ppr is not None:
        new_p.append(_clone_element(ppr))
    r_el = OxmlElement("w:r")
    if rpr is not None:
        r_el.append(_clone_element(rpr))
    t_el = OxmlElement("w:t")
    if line_text and (line_text[0].isspace() or line_text[-1].isspace()):
        t_el.set(qn("xml:space"), "preserve")
    t_el.text = line_text
    r_el.append(t_el)
    new_p.append(r_el)
    return new_p


def _rewrite_paragraph_single_line(para, line_text: str, *, rpr=None) -> None:
    p_el = para._element
    if rpr is None:
        rpr = _first_run_rpr(p_el)
    for child in list(p_el):
        if child.tag.endswith("}pPr"):
            continue
        p_el.remove(child)
    r_el = OxmlElement("w:r")
    if rpr is not None:
        r_el.append(_clone_element(rpr))
    t_el = OxmlElement("w:t")
    if line_text and (line_text[0].isspace() or line_text[-1].isspace()):
        t_el.set(qn("xml:space"), "preserve")
    t_el.text = line_text
    r_el.append(t_el)
    p_el.append(r_el)


def _rpr_font_name(rpr) -> str:
    if rpr is None:
        return ""
    try:
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            return ""
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            val = rf.get(qn(f"w:{attr}"))
            if val:
                return str(val)
    except Exception:
        pass
    return ""


def _fallback_body_run_rpr(para):
    try:
        doc = para.part.document
    except Exception:
        doc = None
    font_name = (_detect_document_body_font(doc) if doc else None) or "Times New Roman"
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn(f"w:{attr}"), font_name)
    rpr.append(rf)
    sz = _rpr_font_size(_first_run_rpr(para._element)) or 26
    for tag in ("sz", "szCs"):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), str(int(sz)))
        rpr.append(el)
    return rpr


def _resolve_body_run_rpr(para, base_rpr=None):
    for fn in (_dominant_body_run_rpr, _prose_run_rpr):
        try:
            rpr = fn(para)
        except Exception:
            rpr = None
        if rpr is not None and not _is_auxiliary_pdf_font(_rpr_font_name(rpr)):
            return rpr
    fb = _fallback_body_run_rpr(para)
    if fb is not None:
        return fb
    return base_rpr


def _rewrite_paragraph_run_segments(para, segments: List[tuple], *, base_rpr=None) -> None:
    """Rebuild paragraph runs from (text, bold) segments."""
    p_el = para._element
    if base_rpr is None:
        base_rpr = _first_run_rpr(p_el)
    body_rpr = _resolve_body_run_rpr(para, base_rpr)
    for child in list(p_el):
        if child.tag.endswith("}pPr"):
            continue
        p_el.remove(child)
    for text, bold in segments:
        if not text:
            continue
        parts = []
        prefix, rest = _split_leading_glyph_prefix(text)
        if prefix and rest:
            parts = [(prefix, bold), (rest, bold)]
        else:
            parts = [(text, bold)]
        for chunk, chunk_bold in parts:
            if not chunk:
                continue
            r_el = OxmlElement("w:r")
            use_rpr = body_rpr
            if should_preserve_symbol_glyphs() and _is_symbol_glyph_text(chunk):
                use_rpr = base_rpr
            if use_rpr is not None:
                rpr = _clone_element(use_rpr)
                for tag in ("w:b", "w:bCs"):
                    old = rpr.find(qn(tag))
                    if old is not None:
                        rpr.remove(old)
                if chunk_bold:
                    b = OxmlElement("w:b")
                    b.set(qn("w:val"), "1")
                    rpr.append(b)
                r_el.append(rpr)
            elif chunk_bold:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                b.set(qn("w:val"), "1")
                rpr.append(b)
                r_el.append(rpr)
            t_el = OxmlElement("w:t")
            if chunk and (chunk[0].isspace() or chunk[-1].isspace()):
                t_el.set(qn("xml:space"), "preserve")
            t_el.text = chunk
            r_el.append(t_el)
            p_el.append(r_el)


def _paragraph_literal(paragraph) -> str:
    try:
        return "".join(r.text or "" for r in paragraph.runs)
    except Exception:
        return ""


def _loosely_match_pdf_text(pdf_text: str, doc_text: str) -> bool:
    if not (pdf_text or "").strip() or not (doc_text or "").strip():
        return False
    a = _normalize_match_text(pdf_text)
    b = _normalize_match_text(doc_text)
    if a == b:
        return True
    if a in b or b in a:
        return True
    aw = set(a.split())
    bw = set(b.split())
    if not aw or not bw:
        return False
    return len(aw & bw) / max(len(aw), len(bw)) >= 0.55


def _apply_pdf_span_segments_to_paragraph(paragraph, span_specs: List[Dict[str, Any]]) -> bool:
    """Rebuild runs so bold matches PDF span boundaries; keep DOCX text intact."""
    if not span_specs:
        return False
    doc_text = _paragraph_literal(paragraph)
    if not doc_text.strip():
        return False

    pdf_text = "".join(str(s.get("text") or "") for s in span_specs)
    norm_pdf = _normalize_match_text(normalize_pdf_symbol_chars(pdf_text, for_match=True))
    norm_doc = _normalize_match_text(doc_text)
    if not norm_pdf or not norm_doc:
        return False
    if norm_pdf != norm_doc and norm_pdf not in norm_doc and norm_doc not in norm_pdf:
        overlap = min(len(norm_pdf), len(norm_doc)) / max(len(norm_pdf), len(norm_doc))
        if overlap < 0.82:
            return False

    weights = [len(str(s.get("text") or "")) for s in span_specs]
    if sum(weights) <= 0:
        return False

    bold_flags = [bool(s.get("bold")) for s in span_specs]
    chunks = _split_text_by_weights(doc_text, weights)
    segments = [(chunk, bold_flags[i]) for i, chunk in enumerate(chunks) if chunk]
    if not segments:
        return False

    actual = [
        ((r.text or ""), _run_is_bold(r))
        for r in paragraph.runs
        if (r.text or "")
    ]
    if len(actual) == len(segments):
        if all(at == et and ab == eb for (at, ab), (et, eb) in zip(actual, segments)):
            return False

    base_rpr = _first_run_rpr(paragraph._element)
    _rewrite_paragraph_run_segments(paragraph, segments, base_rpr=base_rpr)
    return True


def _strip_unmatched_pdf2docx_bold(paragraph, pdf_lines: List[Dict[str, Any]]) -> int:
    """Drop pdf2docx false-bold when no PDF line matches this paragraph."""
    doc_text = _paragraph_literal(paragraph)
    if not doc_text.strip():
        return 0
    if not any(_run_is_bold(r) for r in paragraph.runs if (r.text or "").strip()):
        return 0
    for pl in pdf_lines:
        if _loosely_match_pdf_text(pl.get("text") or "", doc_text):
            return 0
    changed = 0
    for run in paragraph.runs:
        if (run.text or "").strip() and _run_is_bold(run):
            _set_run_bold(run, False)
            changed += 1
    return changed


def sync_bold_spans_from_pdf_to_doc(doc: docx.Document, pdf_path: str) -> int:
    """Align DOCX run-level bold with the PDF text layer (ground truth)."""
    if not pdf_path or not os.path.isfile(pdf_path):
        return 0
    if not _env_bool("PDF_DOCX_PDF_BOLD_SYNC", True):
        return 0

    pdf_lines = extract_pdf_paragraph_formats(pdf_path)
    if not pdf_lines:
        return 0

    paras = list(iter_all_paragraphs(doc))
    para_texts = [_paragraph_plain(p) for p in paras]
    hints = _match_pdf_formats_to_paragraphs(para_texts, pdf_lines)

    changed = 0
    for para, hint in zip(paras, hints):
        spans = (hint or {}).get("spans")
        if spans:
            if _apply_pdf_span_segments_to_paragraph(para, spans):
                changed += 1
        else:
            changed += _strip_unmatched_pdf2docx_bold(para, pdf_lines)
    return changed


def _split_text_by_weights(text: str, weights: List[int]) -> List[str]:
    n = len(weights)
    txt = text or ""
    if n <= 0:
        return []
    if n == 1:
        return [txt]
    if not txt:
        return [""] * n

    total_weight = max(1, sum(max(1, int(w)) for w in weights))
    txt_len = len(txt)
    boundaries: List[int] = []
    acc = 0
    for i in range(1, n):
        acc += max(1, int(weights[i - 1]))
        target = int(round((acc / total_weight) * txt_len))
        target = max(1, min(txt_len - 1, target))
        left = max(1, target - 24)
        right = min(txt_len - 1, target + 24)
        cut = target
        found = None
        p = target
        while p >= left:
            if txt[p - 1].isspace():
                found = p
                break
            p -= 1
        if found is None:
            p = target + 1
            while p <= right:
                if txt[p - 1].isspace():
                    found = p
                    break
                p += 1
        if found is not None:
            cut = found
        if boundaries and cut <= boundaries[-1]:
            cut = min(txt_len - 1, boundaries[-1] + 1)
        boundaries.append(cut)

    out: List[str] = []
    prev = 0
    for b in boundaries:
        out.append(txt[prev:b])
        prev = b
    out.append(txt[prev:])
    if len(out) < n:
        out.extend([""] * (n - len(out)))
    elif len(out) > n:
        out = out[: n - 1] + ["".join(out[n - 1 :])]
    return out


def _collect_run_style_regions(runs: List) -> List[Dict[str, Any]]:
    regions: List[Dict[str, Any]] = []
    for run in runs:
        chunk = run.text or ""
        if not chunk:
            continue
        rpr = run._element.find(qn("w:rPr"))
        regions.append(
            {
                "bold": _run_is_bold(run),
                "rpr": copy.deepcopy(rpr) if rpr is not None else None,
                "weight": max(1, len(chunk)),
            }
        )
    return regions


def _rebuild_para_runs_from_style_regions(paragraph, text: str, regions: List[Dict[str, Any]]) -> int:
    if not text or not regions:
        return 0
    chunks = _split_text_by_weights(text, [r["weight"] for r in regions])
    segments = [(t, regions[i]["bold"]) for i, t in enumerate(chunks) if t]
    if not segments:
        return 0
    base = regions[0].get("rpr")
    base_rpr = base if base is not None else _first_run_rpr(paragraph._element)
    _rewrite_paragraph_run_segments(paragraph, segments, base_rpr=base_rpr)
    return len(segments)


def split_merged_pdf_paragraphs(doc: docx.Document) -> int:
    """Split pdf2docx paragraphs that contain \\n/\\t merged lines into one paragraph per line."""
    split_count = 0
    body = doc.element.body
    p_elements = [el for el in body if el.tag == qn("w:p")]

    for p_el in reversed(p_elements):
        para = docx.text.paragraph.Paragraph(p_el, doc)
        if _is_in_table_cell(para):
            continue
        orig_runs = list(para.runs)
        raw = "".join(r.text or "" for r in orig_runs)
        lines = _split_pdf_soft_lines(raw)
        if len(lines) <= 1:
            continue

        ppr = p_el.find(qn("w:pPr"))
        ppr_copy = _clone_element(ppr)

        _rewrite_paragraph_single_line(
            para,
            lines[0],
            rpr=_line_rpr_for_text(orig_runs, lines[0]),
        )

        insert_after = p_el
        for line in lines[1:]:
            new_p = _build_paragraph_element(
                line,
                ppr_copy,
                _line_rpr_for_text(orig_runs, line),
            )
            insert_after.addnext(new_p)
            insert_after = new_p
            split_count += 1

    return split_count


_PDF_CONTINUATION_START_RE = re.compile(r"^[a-z(]|^COD\b|^OnlinePayment\b")


def _is_pdf_continuation_paragraph(prev_text: str, cur_text: str) -> bool:
    prev = (prev_text or "").strip()
    cur = (cur_text or "").strip()
    if not prev or not cur:
        return False
    if re.match(r"^[+\-•\uf02b\uf0b7o]\s*[A-ZÀ-Ỹ]", cur):
        return False
    if re.match(r"^\d+\.\s*\S", cur):
        return False
    if prev.count("(") > prev.count(")"):
        return True
    if prev.rstrip().endswith((",", "(", "–", "—", "/", "như", "such as")):
        return True
    if _PDF_CONTINUATION_START_RE.match(cur):
        return True
    return False


def merge_pdf_continuation_paragraphs(doc: docx.Document) -> int:
    """Join pdf2docx hard line-break fragments (e.g. 'OnlinePayment,' + 'COD).')."""
    merged = 0
    paras = list(doc.paragraphs)
    idx = 1
    while idx < len(paras):
        prev = paras[idx - 1]
        cur = paras[idx]
        if _is_in_table_cell(prev) or _is_in_table_cell(cur):
            idx += 1
            continue
        prev_text = _paragraph_plain(prev)
        cur_text = _paragraph_plain(cur)
        if not _is_pdf_continuation_paragraph(prev_text, cur_text):
            idx += 1
            continue
        joiner = "" if prev_text.rstrip().endswith(("(", "/")) else " "
        combined = f"{prev_text.rstrip()}{joiner}{cur_text.lstrip()}"
        _rewrite_paragraph_single_line(prev, combined, rpr=_first_run_rpr(prev._element))
        try:
            cur._element.getparent().remove(cur._element)
            paras.pop(idx)
            merged += 1
            continue
        except Exception:
            pass
        idx += 1
    return merged


def _src_has_auxiliary_font_runs(src_runs: List) -> bool:
    for run in src_runs:
        if _text_has_pua_symbol(run.text or ""):
            return True
        if _is_auxiliary_pdf_font(_read_run_font_name(run)):
            return True
    return False


def _paragraph_starts_with_marker(text: str, marker: str) -> bool:
    t = (text or "").lstrip()
    m = (marker or "").strip()
    if not m:
        return True
    if t.startswith(m):
        return True
    mapped = normalize_pdf_symbol_chars(m, for_match=True).strip()
    if mapped and t.startswith(mapped):
        return True
    return False


def _prepend_cloned_run(paragraph, src_run) -> None:
    p_el = paragraph._element
    r_el = copy.deepcopy(src_run._element)
    first_r = p_el.find(qn("w:r"))
    if first_r is not None:
        first_r.addprevious(r_el)
    else:
        p_el.append(r_el)


def _preserve_leading_marker_from_source(src_para, dst_para) -> int:
    src_plain = _paragraph_literal(src_para)
    marker = _LEADING_MARKER_RE.match(src_plain or "")
    if not marker:
        return 0
    _dedupe_leading_markers_in_paragraph(dst_para)
    marker_text = marker.group(1)
    dst_plain = _paragraph_literal(dst_para)
    if _paragraph_starts_with_marker(dst_plain, marker_text):
        return 0
    portable = _portable_marker_text(marker_text)
    body_font = _detect_document_body_font(dst_para.part.document) or "Times New Roman"
    p_el = dst_para._element
    r_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn(f"w:{attr}"), body_font)
    rpr.append(rf)
    r_el.append(rpr)
    t_el = OxmlElement("w:t")
    t_el.text = portable
    r_el.append(t_el)
    first_r = p_el.find(qn("w:r"))
    if first_r is not None:
        first_r.addprevious(r_el)
    else:
        p_el.append(r_el)
    return 1


def _normalize_font_key(name: str) -> str:
    return re.sub(r"\s+", " ", (name or "").strip().lower())


def _is_auxiliary_pdf_font(name: str) -> bool:
    n = _normalize_font_key(name)
    if not n:
        return False
    if any(k in n for k in ("courier", "symbol", "wingding", "monospace", "consolas")):
        return True
    return n in {"couriernewpsmt", "symbolmt", "zapfdingbats"}


def _canonical_font_display_name(name: str) -> str:
    n = _normalize_font_key(name)
    if "times" in n:
        return "Times New Roman"
    if "arial" in n:
        return "Arial"
    if "calibri" in n:
        return "Calibri"
    if "noto" in n:
        return "Noto Sans"
    cleaned = re.sub(r"(psmt|ps-boldmt|ps-italicmt|mt)$", "", n, flags=re.IGNORECASE).strip()
    if not cleaned:
        return (name or "Times New Roman").strip()
    return cleaned.title()


def _read_run_font_name(run) -> str:
    try:
        if run.font.name:
            return str(run.font.name)
    except Exception:
        pass
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return ""
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            return ""
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            val = rf.get(qn(f"w:{attr}"))
            if val:
                return str(val)
    except Exception:
        pass
    return ""


def _set_run_font_name(run, font_name: str) -> None:
    if not font_name:
        return
    try:
        r_el = run._element
        rpr = r_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r_el.insert(0, rpr)
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rf.set(qn(f"w:{attr}"), font_name)
        try:
            run.font.name = font_name
        except Exception:
            pass
    except Exception:
        pass


def _detect_document_body_font(doc: docx.Document) -> Optional[str]:
    counts: Dict[str, int] = {}
    for para in iter_all_paragraphs(doc):
        for run in para.runs:
            text = (run.text or "").strip()
            if len(text) < 2:
                continue
            fname = _read_run_font_name(run)
            if not fname or _is_auxiliary_pdf_font(fname):
                continue
            key = _canonical_font_display_name(fname)
            counts[key] = counts.get(key, 0) + len(text)
    if not counts:
        return "Times New Roman"
    return max(counts, key=counts.get)


def unify_pdf_docx_fonts(doc: docx.Document) -> int:
    """Map Courier/Symbol pdf2docx fonts to the document body font (usually Times New Roman)."""
    if not _env_bool("PDF_DOCX_UNIFY_FONTS", True):
        return 0
    body_font = _detect_document_body_font(doc)
    if not body_font:
        body_font = "Times New Roman"
    changed = 0
    for para in iter_all_paragraphs(doc):
        for run in para.runs:
            raw = run.text or ""
            if not raw:
                continue
            if _text_has_pua_symbol(raw):
                _set_run_symbol_font(run)
                continue
            fname = _read_run_font_name(run)
            if _is_auxiliary_pdf_font(fname) and len(raw.strip()) <= 2:
                if not re.search(r"[\w\u00C0-\u1EF9]", raw, flags=re.UNICODE):
                    _set_run_symbol_font(run)
                    continue
            canon = _canonical_font_display_name(fname) if fname else body_font
            needs_unify = (
                not fname
                or _is_auxiliary_pdf_font(fname)
                or canon != body_font
                or fname != body_font
            )
            if not needs_unify:
                continue
            _set_run_font_name(run, body_font)
            changed += 1
    return changed


def _detect_body_font_size_halfpts(doc: docx.Document) -> int:
    sizes: List[int] = []
    for para in iter_all_paragraphs(doc):
        if _paragraph_should_be_bold(para):
            continue
        for run in para.runs:
            if not (run.text or "").strip():
                continue
            if _run_is_bold(run):
                continue
            sz = _rpr_font_size(run._element.find(qn("w:rPr")))
            if sz > 0:
                sizes.append(sz)
    if not sizes:
        return 26
    sizes.sort()
    return sizes[len(sizes) // 2]


def _normalize_body_run_sizes(doc: docx.Document) -> int:
    """Raise tiny Courier/pdf2docx sizes to document body size."""
    if not _env_bool("PDF_DOCX_UNIFY_FONTS", True):
        return 0
    target = _detect_body_font_size_halfpts(doc)
    if target <= 0:
        return 0
    changed = 0
    for para in iter_all_paragraphs(doc):
        for run in para.runs:
            if not (run.text or "").strip():
                continue
            rpr = run._element.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run._element.insert(0, rpr)
            sz = _rpr_font_size(rpr)
            if sz >= target:
                continue
            if _paragraph_should_be_bold(para) and len((run.text or "").strip()) > 2:
                continue
            for tag in ("w:sz", "w:szCs"):
                old = rpr.find(qn(tag))
                if old is not None:
                    rpr.remove(old)
                el = OxmlElement(tag)
                el.set(qn("w:val"), str(int(target)))
                rpr.append(el)
            changed += 1
    return changed


def _paragraph_plain(paragraph) -> str:
    try:
        return (paragraph.text or "").replace("\r", "").replace("\n", " ").strip()
    except Exception:
        return ""


def _paragraph_word_count(paragraph) -> int:
    return len(_paragraph_plain(paragraph).split())


def _looks_like_body_prose(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if len(text) < 40:
        return False
    if _is_section_heading_line(paragraph):
        return False
    return _paragraph_word_count(paragraph) >= 8


def _is_short_label_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    return 0 < len(text) <= 90 and _paragraph_word_count(paragraph) <= 14


def _is_section_heading_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if not text:
        return False
    if _is_running_header_line(paragraph):
        return False
    if _SECTION1_HEAD_RE.match(text):
        return True
    if _SECTION2_HEAD_RE.match(text) and not _FALSE_SECTION2_RE.match(text):
        return True
    if re.match(r"^\d+\s+[A-ZÀ-Ỹ]", text) and _paragraph_word_count(paragraph) <= 8:
        return True
    if _ASSIGNMENT_HEAD_RE.match(text) and _paragraph_word_count(paragraph) <= 16:
        return True
    return False


def _is_running_header_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if not text:
        return False
    if _RUNNING_HEADER_RE.match(text):
        return True
    if _FALSE_SECTION2_RE.match(text):
        return True
    return False


def _is_in_table_cell(paragraph) -> bool:
    try:
        el = paragraph._element
        while el is not None:
            tag = getattr(el, "tag", "") or ""
            if tag.endswith("}tc"):
                return True
            el = el.getparent()
    except Exception:
        pass
    return False


def iter_all_paragraphs(doc: docx.Document) -> Iterator:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _dominant_font_pt(paragraph) -> Optional[float]:
    sizes: List[float] = []
    for run in paragraph.runs:
        try:
            if run.font.size is not None:
                sizes.append(float(run.font.size.pt))
        except Exception:
            pass
    if not sizes:
        return None
    sizes.sort()
    return sizes[len(sizes) // 2]


def _find_abstract_start(paras: List) -> Optional[int]:
    for i, p in enumerate(paras):
        if _ABSTRACT_START_RE.match(_paragraph_plain(p)):
            return i
    return None


def _find_keywords_paragraph_index(paras: List, start: int = 0) -> Optional[int]:
    for i in range(start, len(paras)):
        if _KEYWORDS_LINE_RE.match(_paragraph_plain(paras[i])):
            return i
    return None


def _find_section_one_start(paras: List) -> Optional[int]:
    for i, p in enumerate(paras):
        text = _paragraph_plain(p)
        if _SECTION1_HEAD_RE.match(text):
            return i
        if re.match(r"^1\s+[A-ZÀ-Ỹ]", text) and _paragraph_word_count(p) <= 8:
            if not _AFFILIATION_LINE_RE.match(text):
                return i
    return None


def _find_section_two_start(paras: List, *, min_index: int = 0) -> Optional[int]:
    sec1 = _find_section_one_start(paras)
    start_at = max(int(min_index or 0), (sec1 + 1) if sec1 is not None else 0)
    for i, p in enumerate(paras):
        if i < start_at:
            continue
        text = _paragraph_plain(p)
        if _FALSE_SECTION2_RE.match(text):
            continue
        if _is_running_header_line(p):
            continue
        if _AFFILIATION_LINE_RE.match(text):
            continue
        if re.search(r"\bskrywer\b", text, re.IGNORECASE) and _paragraph_word_count(p) <= 12:
            continue
        if _SECTION2_HEAD_RE.match(text) and _paragraph_word_count(p) <= 8:
            return i
    return None


def is_running_header_text(text: str) -> bool:
    t = (text or "").replace("\r", " ").replace("\n", " ").strip()
    if not t:
        return False
    if _PAGE_NO_ONLY_RE.match(t):
        return True
    if _RUNNING_HEADER_RE.match(t):
        return True
    if _FALSE_SECTION2_RE.match(t):
        return True
    if re.match(r"^\d+\s+[A-ZÀ-Ỹ]\.\s+\S+\s+(?:en|and)\s+[A-ZÀ-Ỹ]\.\s+(?:Author|Skrywer)\b", t, re.IGNORECASE):
        return True
    return False


def is_pdf_artifact_text(text: str) -> bool:
    t = (text or "").replace("\r", " ").replace("\n", " ").strip()
    if not t:
        return False
    if _PDF_ARTIFACT_RE.match(t):
        return True
    if re.search(r"__\s*[\wÀ-Ỹ]{2,10}\s*_\s*\d+\s*__", t, re.IGNORECASE):
        return True
    if re.match(r"^__?[\wÀ-Ỹ]{2,10}_?\d*__?\s*[,.\s]*$", t, re.IGNORECASE):
        return True
    if re.search(r"__?\s*HOU[\s_\d]*__?", t, re.IGNORECASE):
        return True
    if re.match(r"^[,.\s|_\-]+$", t):
        return True
    letters = len(re.findall(r"[A-Za-zÀ-ỹ]", t))
    if len(t) <= 48 and letters <= 2 and re.search(r"[_|,.\s-]{3,}", t):
        return True
    return False


def _is_noise_paragraph(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    return is_running_header_text(text) or is_pdf_artifact_text(text)


def strip_noise_paragraphs(doc: docx.Document) -> Dict[str, int]:
    """Remove running headers/footers and pdf2docx artifacts injected into body flow."""
    stats = {"noise_removed": 0, "inline_artifacts_stripped": 0}
    if not _env_bool("PDF_DOCX_STRIP_NOISE", True):
        return stats
    stats["inline_artifacts_stripped"] = _strip_inline_pdf_artifacts(doc)
    paras = list(doc.paragraphs)
    idx = 0
    while idx < len(paras):
        if _is_noise_paragraph(paras[idx]):
            try:
                paras[idx]._element.getparent().remove(paras[idx]._element)
                paras.pop(idx)
                stats["noise_removed"] += 1
                continue
            except Exception:
                pass
        idx += 1
    return stats


_INLINE_ARTIFACT_RE = re.compile(
    r"__?\s*[\wÀ-Ỹ]{2,10}\s*_\s*\d+\s*__?\s*[,.\s]*",
    re.IGNORECASE,
)


def _strip_inline_pdf_artifacts(doc: docx.Document) -> int:
    """Remove pdf2docx placeholder tokens left inside paragraph runs."""
    stripped = 0
    for para in doc.paragraphs:
        changed = False
        for run in para.runs:
            raw = run.text or ""
            if not raw:
                continue
            if not raw.strip():
                continue
            cleaned = _INLINE_ARTIFACT_RE.sub("", raw)
            cleaned = re.sub(r"^\s*[,.\s]+\s*", "", cleaned)
            if cleaned != raw:
                run.text = cleaned
                changed = True
        if changed:
            stripped += 1
    return stripped


def _paragraph_jc_val(paragraph) -> Optional[str]:
    ppr = paragraph._element.find(qn("w:pPr"))
    return _read_ppr_jc_val(ppr)


def _detect_title_block_end(paras: List) -> int:
    """Detect title/author block; body formatting starts after this index."""
    abs_i = _find_abstract_start(paras)
    if abs_i is not None:
        return abs_i
    sec1 = _find_section_one_start(paras)
    if sec1 is not None:
        return sec1
    end = 0
    for i, p in enumerate(paras[:24]):
        text = _paragraph_plain(p)
        if not text:
            end = max(end, i + 1)
            continue
        jc = _paragraph_jc_val(p) or ""
        if jc == "center" and len(text) <= 140 and _paragraph_word_count(p) <= 22:
            end = i + 1
            continue
        if i < 8 and len(text) <= 90:
            end = i + 1
            continue
        break
    return end


def _paragraph_is_centered(paragraph) -> bool:
    jc = _paragraph_jc_val(paragraph)
    if jc == "center":
        return True
    try:
        return paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        return False


def _detect_layout_regions(paras: List) -> Tuple[int, int, int, int, Optional[int]]:
    """Return title_end, abstract_start, abstract_end, body_start, sec2_start."""
    title_end = _detect_title_block_end(paras)
    abs_start = _find_abstract_start(paras)
    if abs_start is None:
        abs_start = title_end
    kw_idx = _find_keywords_paragraph_index(paras, title_end)
    sec1 = _find_section_one_start(paras)
    sec2 = _find_section_two_start(paras, min_index=(sec1 + 1) if sec1 is not None else 0)

    if kw_idx is not None:
        abstract_end = kw_idx + 1
    elif sec1 is not None:
        abstract_end = sec1
    elif sec2 is not None:
        abstract_end = sec2
    else:
        abstract_end = min(abs_start + 24, len(paras))

    body_start = sec1 if sec1 is not None else abstract_end
    return title_end, abs_start, abstract_end, body_start, sec2


def _emu_to_twips(value_emu: int) -> int:
    return int(int(value_emu) * 1440 / 914400)


def _get_content_width_twips(doc: docx.Document) -> int:
    try:
        sec = doc.sections[0]
        content_emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
        return max(3600, _emu_to_twips(content_emu))
    except Exception:
        return 9024


def _get_fallback_body_profile(doc: docx.Document) -> Dict[str, int]:
    left = _env_int("PDF_DOCX_BODY_INDENT_LEFT", 1054)
    right = _env_int("PDF_DOCX_BODY_INDENT_RIGHT", 1008)
    try:
        sec = doc.sections[0]
        lm = _emu_to_twips(int(sec.left_margin))
        rm = _emu_to_twips(int(sec.right_margin))
        if 400 <= lm <= 1800:
            left = lm
        if 400 <= rm <= 1800:
            right = rm
    except Exception:
        pass
    return {"left": left, "right": right}


def _is_abstract_style_indent(left: int, right: int) -> bool:
    return left >= 1350 and right >= 1350


def _is_narrow_column_indent(left: int, right: int, content_twips: int) -> bool:
    total = int(left) + int(right)
    if total <= 0:
        return False
    if _is_abstract_style_indent(left, right):
        return True
    return total > max(2200, int(content_twips * 0.28))


def _derive_layout_profiles_from_pdf(
    pdf_path: str,
    doc: docx.Document,
) -> Optional[Tuple[Dict[str, int], Dict[str, int]]]:
    if not pdf_path or not os.path.isfile(pdf_path):
        return None
    if not _env_bool("PDF_DOCX_PDF_GEOMETRY_PROFILE", True):
        return None

    formats = extract_pdf_paragraph_formats(pdf_path)
    if not formats:
        return None

    content_twips = _get_content_width_twips(doc)
    try:
        sec = doc.sections[0]
        lm_ratio = float(int(sec.left_margin)) / float(int(sec.page_width))
        rm_ratio = float(int(sec.right_margin)) / float(int(sec.page_width))
    except Exception:
        lm_ratio, rm_ratio = 0.121, 0.121
    content_ratio = max(0.45, 1.0 - lm_ratio - rm_ratio)

    body_indents: List[Tuple[int, int, float]] = []
    abstract_indents: List[Tuple[int, int, float]] = []

    for fmt in formats:
        pw = float(fmt.get("page_width") or 0)
        x0 = fmt.get("x0")
        x1 = fmt.get("x1")
        if pw <= 0 or x0 is None or x1 is None:
            continue

        left_pt = max(0.0, float(x0) - pw * lm_ratio)
        right_pt = max(0.0, pw * (1.0 - rm_ratio) - float(x1))
        width_pt = max(0.0, float(x1) - float(x0))
        width_ratio = width_pt / max(1.0, pw * content_ratio)

        left_twips = int(max(0.0, left_pt / max(1.0, pw * content_ratio)) * content_twips)
        right_twips = int(max(0.0, right_pt / max(1.0, pw * content_ratio)) * content_twips)

        if width_ratio >= 0.74 and left_twips <= 1800:
            body_indents.append((left_twips, right_twips, width_ratio))
        elif 0.52 <= width_ratio <= 0.78 and abs(left_twips - right_twips) <= 400:
            abstract_indents.append((left_twips, right_twips, width_ratio))

    if len(body_indents) < 3:
        return None

    body_indents.sort(key=lambda item: item[0] + item[1])
    wider = body_indents[: max(3, len(body_indents) // 2)]
    lefts = [item[0] for item in wider]
    rights = [item[1] for item in wider]
    body_profile = {
        "left": _median_int(lefts, 1054),
        "right": _median_int(rights, 1008),
    }

    if abstract_indents:
        abstract_indents.sort(key=lambda item: item[0] + item[1])
        mid = abstract_indents[len(abstract_indents) // 2]
        abs_left = int(mid[0])
        abs_right = int(mid[1])
        inset = max(abs_left, abs_right, int(body_profile["left"]) + _ABSTRACT_EXTRA_INDENT_TWIPS)
        abstract_profile = {"left": inset, "right": inset}
    else:
        abstract_profile = {
            "left": body_profile["left"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
            "right": body_profile["right"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        }
    return body_profile, abstract_profile


def _compute_indent_profile(
    paras: List,
    start: int,
    end: int,
    *,
    doc: Optional[docx.Document] = None,
    default_left: int = 1054,
    default_right: int = 1008,
    for_body: bool = True,
) -> Dict[str, int]:
    candidates: List[Tuple[int, int, int]] = []
    for p in paras[start:end]:
        text = _paragraph_plain(p)
        if len(text) < 25:
            continue
        if _is_section_heading_line(p) or _is_running_header_line(p):
            continue
        if _paragraph_is_centered(p):
            continue
        ppr = p._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))
        right = int(ind.get("right", 0))
        if left <= 0 and right <= 0:
            continue
        if for_body and _is_abstract_style_indent(left, right):
            continue
        if left + right > 4200:
            continue
        candidates.append((left + right, left, right))

    if not candidates:
        if doc is not None:
            return _get_fallback_body_profile(doc)
        return {"left": default_left, "right": default_right}

    candidates.sort(key=lambda item: item[0])
    wider = candidates[: max(1, len(candidates) // 2 + 1)]
    lefts = [item[1] for item in wider]
    rights = [item[2] for item in wider]
    return {
        "left": _median_int(lefts, default_left),
        "right": _median_int(rights, default_right),
    }


def _paragraph_needs_indent_fix(
    paragraph,
    *,
    region: str,
    body_profile: Dict[str, int],
    abstract_profile: Dict[str, int],
    content_twips: int,
) -> bool:
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    left = int(ind.get("left", 0))
    right = int(ind.get("right", 0))
    text = _paragraph_plain(paragraph)
    if len(text) < 15:
        return False

    if _paragraph_is_centered(paragraph) and len(text) >= 15:
        return True

    if region == "abstract":
        target_l = int(abstract_profile.get("left", 0))
        target_r = int(abstract_profile.get("right", 0))
        if left + right == 0 and len(text) >= 20:
            return True
        if abs(left - target_l) > 700 or abs(right - target_r) > 700:
            return left + right > 0
        return False

    if region in ("body", "section1"):
        if _is_abstract_style_indent(left, right):
            return True
        if _is_narrow_column_indent(left, right, content_twips):
            return True
        if left + right == 0 and len(text) >= 20 and not _is_section_heading_line(paragraph):
            return True
        if _is_indent_outlier(left, body_profile):
            return True
        if right == 0 and left > int(body_profile.get("left", 0)) + 900:
            return True
    return False


def _resolve_layout_profiles(
    doc: docx.Document,
    paras: List,
    *,
    pdf_path: Optional[str] = None,
) -> Tuple[Dict[str, int], Dict[str, int]]:
    pdf_profiles = _derive_layout_profiles_from_pdf(pdf_path, doc) if pdf_path else None
    _, abs_start, abstract_end, body_start, sec2 = _detect_layout_regions(paras)
    body_end = sec2 if sec2 is not None else len(paras)
    first_line_profile = _compute_first_line_profile(paras, body_start, body_end)

    if pdf_profiles:
        pdf_body, pdf_abstract = pdf_profiles
        if int(pdf_body.get("left", 0)) + int(pdf_body.get("right", 0)) <= 2400:
            pdf_body = dict(pdf_body)
            pdf_body["firstLine"] = first_line_profile
            return pdf_body, pdf_abstract

    body_profile = _compute_indent_profile(
        paras,
        body_start,
        body_end,
        doc=doc,
        for_body=True,
    )
    abstract_profile = _compute_indent_profile(
        paras,
        abs_start,
        abstract_end,
        doc=doc,
        default_left=body_profile["left"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        default_right=body_profile["right"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        for_body=False,
    )
    if abstract_profile["left"] <= body_profile["left"]:
        abstract_profile = {
            "left": body_profile["left"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
            "right": body_profile["right"] + _ABSTRACT_EXTRA_INDENT_TWIPS,
        }
    body_profile["firstLine"] = first_line_profile
    return body_profile, abstract_profile


def _read_first_line_twips(paragraph) -> int:
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    return int(ind.get("firstLine", 0))


def _paragraph_has_hanging_indent(paragraph) -> bool:
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    return int(ind.get("hanging", 0)) > 0


def _is_first_paragraph_after_heading(paras: List, idx: int, body_start: int) -> bool:
    if idx <= body_start:
        return True
    for j in range(idx - 1, max(body_start - 1, -1), -1):
        prev = paras[j]
        text = _paragraph_plain(prev)
        if not text:
            continue
        if _is_section_heading_line(prev) and _paragraph_word_count(prev) <= 10:
            return True
        return False
    return True


def _compute_first_line_profile(paras: List, start: int, end: int) -> int:
    values: List[int] = []
    first_after_heading = True
    for p in paras[start:end]:
        text = _paragraph_plain(p)
        if not text:
            continue
        if _is_section_heading_line(p) and _paragraph_word_count(p) <= 10:
            first_after_heading = True
            continue
        if first_after_heading:
            first_after_heading = False
            continue
        fl = _read_first_line_twips(p)
        if fl > 0:
            values.append(fl)
    return _median_int(values, 0)


def _resolve_body_first_line_twips(
    dst_para,
    src_para,
    body_profile: Dict[str, int],
    *,
    is_first_after_heading: bool,
) -> Optional[int]:
    """Resolve firstLine twips; None means keep existing hanging/firstLine untouched."""
    del body_profile
    if is_first_after_heading:
        return 0
    for para in (dst_para, src_para):
        if para is None:
            continue
        if _paragraph_has_hanging_indent(para):
            return None
        fl = _read_first_line_twips(para)
        if fl > 0:
            return fl
    return 0


def _apply_regional_layout_profiles(
    doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
    src_doc: Optional[docx.Document] = None,
) -> Dict[str, int]:
    """Apply title/abstract/body layout profiles to every paragraph (global PDF fix)."""
    stats = {"regional_layout_applied": 0}
    if not _env_bool("PDF_DOCX_REGIONAL_LAYOUT", True):
        return stats

    paras = list(doc.paragraphs)
    if not paras:
        return stats

    title_end, abs_start, abstract_end, body_start, _sec2 = _detect_layout_regions(paras)
    body_profile, abstract_profile = _resolve_layout_profiles(doc, paras, pdf_path=pdf_path)
    src_paras = list(src_doc.paragraphs) if src_doc is not None else []

    for i, p in enumerate(paras):
        if _is_in_table_cell(p):
            continue
        text = _paragraph_plain(p)
        if not text:
            continue
        src_p = src_paras[i] if i < len(src_paras) else None

        if i < title_end:
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.CENTER)
            _set_paragraph_indents(p, left=0, right=0, first_line=0)
            stats["regional_layout_applied"] += 1
            continue

        if _is_running_header_line(p):
            _compact_running_header(p, profile=body_profile)
            stats["regional_layout_applied"] += 1
            continue

        if _is_section_heading_line(p) and _paragraph_word_count(p) <= 10:
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.LEFT)
            _set_paragraph_indents(
                p,
                left=int(body_profile["left"]),
                right=0,
                first_line=0,
            )
            stats["regional_layout_applied"] += 1
            continue

        if abs_start <= i < abstract_end or (title_end <= i < body_start):
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            _set_paragraph_indents(
                p,
                left=int(abstract_profile["left"]),
                right=int(abstract_profile["right"]),
                first_line=0,
            )
            _clear_first_line_indent(p)
            stats["regional_layout_applied"] += 1
            continue

        if i >= body_start:
            _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            first_line = _resolve_body_first_line_twips(
                p,
                src_p,
                body_profile,
                is_first_after_heading=_is_first_paragraph_after_heading(paras, i, body_start),
            )
            indent_kwargs: Dict[str, int] = {
                "left": int(body_profile["left"]),
                "right": int(body_profile["right"]),
            }
            if first_line is not None:
                indent_kwargs["first_line"] = int(first_line)
            _set_paragraph_indents(p, **indent_kwargs)
            stats["regional_layout_applied"] += 1

    return stats


def normalize_converted_docx_layout_in_doc(
    doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
    src_doc: Optional[docx.Document] = None,
    layout_mode: str = "academic",
) -> Dict[str, int]:
    """Fix pdf2docx center alignment and inconsistent indents before/after translation."""
    stats: Dict[str, int] = {"alignment_normalized": 0, "indents_normalized": 0}
    if layout_mode == "conservative":
        return stats
    if not _env_bool("PDF_DOCX_NORMALIZE_ALIGN", True):
        return stats

    paras = list(doc.paragraphs)
    if not paras:
        return stats

    if uses_regional_layout(layout_mode):
        regional = _apply_regional_layout_profiles(
            doc,
            pdf_path=pdf_path,
            src_doc=src_doc or doc,
        )
        stats["regional_layout_applied"] = int(regional.get("regional_layout_applied", 0))
        stats["indents_normalized"] = stats["regional_layout_applied"]
        stats["alignment_normalized"] = stats["regional_layout_applied"]
        return stats

    title_end, abs_start, abstract_end, body_start, sec2 = _detect_layout_regions(paras)
    force_justify = _env_bool("PDF_DOCX_FORCE_BODY_JUSTIFY", True)
    content_twips = _get_content_width_twips(doc)
    body_profile, abstract_profile = _resolve_layout_profiles(doc, paras, pdf_path=pdf_path)

    body_end = sec2 if sec2 is not None else len(paras)
    narrow_body_count = 0
    body_para_count = 0
    for i in range(body_start, body_end):
        text = _paragraph_plain(paras[i])
        if len(text) < 20:
            continue
        body_para_count += 1
        ppr = paras[i]._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))
        right = int(ind.get("right", 0))
        if _paragraph_is_centered(paras[i]) or _is_narrow_column_indent(left, right, content_twips):
            narrow_body_count += 1
    force_body_width = (
        body_para_count >= 2
        and narrow_body_count / max(1, body_para_count) >= 0.35
    )

    for i, p in enumerate(paras):
        if i < title_end:
            continue
        text = _paragraph_plain(p)
        if len(text) < 15:
            continue
        if _is_running_header_line(p):
            continue
        if _is_section_heading_line(p) and _paragraph_word_count(p) <= 10:
            continue

        if abs_start <= i < abstract_end:
            region = "abstract"
        elif i >= body_start:
            region = "body"
        else:
            region = "abstract"

        centered = _paragraph_is_centered(p)
        jc = _paragraph_jc_val(p) or ""
        if centered or (force_justify and jc in ("left", "start", "") and len(text) >= 30):
            if not (_is_section_heading_line(p) and _paragraph_word_count(p) <= 10):
                _set_paragraph_alignment(p, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
                stats["alignment_normalized"] += 1

        if not _env_bool("PDF_DOCX_NORMALIZE_INDENTS", True):
            continue

        needs_fix = _paragraph_needs_indent_fix(
            p,
            region=region,
            body_profile=body_profile,
            abstract_profile=abstract_profile,
            content_twips=content_twips,
        )
        if force_body_width and region == "body" and len(text) >= 20:
            if not _is_section_heading_line(p) or _paragraph_word_count(p) > 10:
                needs_fix = True

        if not needs_fix:
            continue

        target = abstract_profile if region == "abstract" else body_profile
        ppr = p._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        first_line = 0 if region == "abstract" else int(ind.get("firstLine", 0))
        _set_paragraph_indents(
            p,
            left=int(target["left"]),
            right=int(target["right"]),
            first_line=first_line,
        )
        if region == "abstract":
            _clear_first_line_indent(p)
        stats["indents_normalized"] += 1

    return stats


def sanitize_converted_docx(
    docx_path: str,
    *,
    pdf_path: Optional[str] = None,
    layout_mode: str = "academic",
) -> Dict[str, int]:
    """Clean pdf2docx noise + normalize body alignment/indents before translation."""
    if not os.path.isfile(docx_path):
        return {"noise_removed": 0, "alignment_normalized": 0, "indents_normalized": 0}
    doc = docx.Document(docx_path)
    stats: Dict[str, int] = {
        "noise_removed": 0,
        "alignment_normalized": 0,
        "indents_normalized": 0,
    }
    if should_preserve_pdf_lines(layout_mode):
        stats["symbols_normalized"] = normalize_pdf_symbol_chars_in_doc(doc)
        stats["lines_split"] = split_merged_pdf_paragraphs(doc)
        stats["lines_merged"] = merge_pdf_continuation_paragraphs(doc)
        stats["fonts_unified"] = unify_pdf_docx_fonts(doc)
        stats["sizes_normalized"] = _normalize_body_run_sizes(doc)
        if pdf_path and os.path.isfile(pdf_path):
            stats["pdf_bold_synced"] = sync_bold_spans_from_pdf_to_doc(doc, pdf_path)
        if _env_bool("PDF_DOCX_STRIP_NOISE", True):
            noise = strip_noise_paragraphs(doc)
            stats["noise_removed"] = int(noise.get("noise_removed", 0))
            stats["inline_artifacts_stripped"] = int(noise.get("inline_artifacts_stripped", 0))
        try:
            from app.services.docx_service import apply_docx_paragraph_spacing

            spaced = 0
            for para in doc.paragraphs:
                before = "".join(r.text or "" for r in para.runs)
                apply_docx_paragraph_spacing(para)
                after = "".join(r.text or "" for r in para.runs)
                if before != after:
                    spaced += 1
            stats["spacing_fixed"] = spaced
            stats["latin_font_repaired"] = repair_latin_runs_in_symbol_font(doc)
        except Exception:
            pass
    else:
        stats.update(strip_noise_paragraphs(doc))
    if pdf_path and os.path.isfile(pdf_path):
        stats["pdf_bold_synced"] = stats.get("pdf_bold_synced", 0) or sync_bold_spans_from_pdf_to_doc(doc, pdf_path)
    if should_merge_pdf_fragments(layout_mode):
        paras = list(doc.paragraphs)
        sec2 = _find_section_two_start(paras) or len(paras)
        stats["fragments_merged"] = _collapse_pre_section2_fragments(doc, paras, sec2)
    if layout_mode == "academic":
        stats.update(
            normalize_converted_docx_layout_in_doc(
                doc,
                pdf_path=pdf_path,
                layout_mode=layout_mode,
            )
        )
    if any(int(v or 0) for v in stats.values()):
        doc.save(docx_path)
    return stats


def _find_title_block_end(paras: List) -> int:
    idx = _find_abstract_start(paras)
    return idx if idx is not None else 0


def _resolve_layout_region(
    idx: int,
    *,
    title_end: int,
    sec1_start: Optional[int],
    sec2_start: Optional[int],
) -> str:
    if idx < title_end:
        return "title"
    if sec2_start is not None and idx >= sec2_start:
        return "preserve"
    if sec1_start is not None and idx >= sec1_start:
        return "section1"
    return "abstract"


def _read_ppr_jc_val(ppr) -> Optional[str]:
    if ppr is None:
        return None
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        return None
    return jc.get(qn("w:val"))


def _read_ppr_ind_twips(ppr) -> Dict[str, int]:
    out: Dict[str, int] = {}
    if ppr is None:
        return out
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return out
    for key in ("left", "right", "firstLine", "hanging"):
        raw = ind.get(qn(f"w:{key}"))
        if raw is not None:
            try:
                out[key] = int(raw)
            except Exception:
                pass
    return out


def _read_paragraph_alignment(paragraph) -> Optional[WD_PARAGRAPH_ALIGNMENT]:
    try:
        return paragraph.alignment
    except Exception:
        return None


def _set_paragraph_alignment(paragraph, align: WD_PARAGRAPH_ALIGNMENT) -> None:
    try:
        paragraph.alignment = align
    except Exception:
        pass
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    val = {
        WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
        WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
        WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
        WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "both",
        WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: "distribute",
    }.get(align, "left")
    jc.set(qn("w:val"), val)


def _copy_paragraph_properties(src_ppr, dst_ppr) -> None:
    if src_ppr is None:
        return
    new_ppr = copy.deepcopy(src_ppr)
    parent = dst_ppr.getparent()
    if parent is None:
        return
    if dst_ppr is not None:
        parent.remove(dst_ppr)
    parent.insert(0, new_ppr)


def _copy_paragraph_layout_template(src_para, dst_para, *, include_spacing: bool = True) -> None:
    src_ppr = src_para._element.find(qn("w:pPr"))
    if src_ppr is None:
        return
    cloned = copy.deepcopy(src_ppr)
    if not include_spacing:
        spacing = cloned.find(qn("w:spacing"))
        if spacing is not None:
            cloned.remove(spacing)
    dst_ppr = dst_para._element.find(qn("w:pPr"))
    _copy_paragraph_properties(cloned, dst_ppr)


def _set_paragraph_indents(
    paragraph,
    *,
    left: Optional[int] = None,
    right: Optional[int] = None,
    first_line: Optional[int] = None,
    hanging: Optional[int] = None,
) -> None:
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if left is not None:
        ind.set(qn("w:left"), str(int(left)))
    if right is not None:
        ind.set(qn("w:right"), str(int(right)))
    if hanging is not None and int(hanging) > 0:
        ind.set(qn("w:hanging"), str(int(hanging)))
        ind.attrib.pop(qn("w:firstLine"), None)
    elif first_line is not None:
        if int(first_line) > 0:
            ind.set(qn("w:firstLine"), str(int(first_line)))
            ind.attrib.pop(qn("w:hanging"), None)
        else:
            ind.attrib.pop(qn("w:firstLine"), None)
            if hanging is None:
                ind.attrib.pop(qn("w:hanging"), None)


def _clear_first_line_indent(paragraph) -> None:
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return
    for key in (qn("w:firstLine"), qn("w:hanging")):
        if ind.get(key) is not None:
            ind.attrib.pop(key, None)


def _apply_font_pt(paragraph, pt: float) -> None:
    if pt <= 0:
        return
    for run in paragraph.runs:
        try:
            run.font.size = Pt(pt)
        except Exception:
            pass


def _score_body_reference_paragraph(paragraph) -> int:
    text = _paragraph_plain(paragraph)
    if not text or len(text) < 30 or _is_in_table_cell(paragraph):
        return -1
    if _is_section_heading_line(paragraph) or _is_running_header_line(paragraph):
        return -1
    ppr = paragraph._element.find(qn("w:pPr"))
    ind = _read_ppr_ind_twips(ppr)
    if ind.get("left", 0) <= 0:
        return -1
    return 1


def _find_body_reference_paragraph(paras: List, sec2_start: int) -> Optional[int]:
    for i in range(sec2_start + 1, min(sec2_start + 12, len(paras))):
        if _score_body_reference_paragraph(paras[i]) > 0:
            return i
    return None


def _find_abstract_reference_paragraph(paras: List, start: int, end: int) -> Optional[int]:
    for i in range(start, end):
        ppr = paras[i]._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        if ind.get("left", 0) >= 1500:
            return i
    return start


def _body_mode_font_pt(paras: List, sec2_start: int) -> Optional[float]:
    ref = _find_body_reference_paragraph(paras, sec2_start)
    if ref is None:
        return None
    return _dominant_font_pt(paras[ref])


def _set_paragraph_spacing(
    paragraph,
    *,
    line: Optional[int] = None,
    line_rule: Optional[str] = None,
    before: Optional[int] = None,
    after: Optional[int] = None,
) -> None:
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph._element.insert(0, ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    if line is not None:
        spacing.set(qn("w:line"), str(int(line)))
    if line_rule is not None:
        spacing.set(qn("w:lineRule"), line_rule)
    if before is not None:
        spacing.set(qn("w:before"), str(int(before)))
    if after is not None:
        spacing.set(qn("w:after"), str(int(after)))


def _read_spacing_twips(paragraph) -> Dict[str, int]:
    out: Dict[str, int] = {}
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is None:
        return out
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        return out
    for key in ("line", "before", "after"):
        raw = spacing.get(qn(f"w:{key}"))
        if raw is not None:
            try:
                out[key] = int(raw)
            except Exception:
                pass
    rule = spacing.get(qn("w:lineRule"))
    if rule:
        out["lineRule"] = rule  # type: ignore[assignment]
    return out


def _env_int(name: str, default: int) -> int:
    raw = os.getenv(name)
    if raw is None:
        return int(default)
    try:
        return int(str(raw).strip())
    except Exception:
        return int(default)


def _median_int(values: List[int], default: int = 0) -> int:
    if not values:
        return default
    ordered = sorted(int(v) for v in values)
    return ordered[len(ordered) // 2]


def _compute_body_layout_profile(paras: List) -> Dict[str, int]:
    """Median indent/spacing from body-like paragraphs (dynamic per document)."""
    lefts: List[int] = []
    rights: List[int] = []
    befores: List[int] = []
    afters: List[int] = []

    for paragraph in paras:
        text = _paragraph_plain(paragraph)
        if len(text) < 20:
            continue
        if _is_section_heading_line(paragraph) and _paragraph_word_count(paragraph) <= 8:
            continue
        ppr = paragraph._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        sp = _read_spacing_twips(paragraph)
        if ind.get("left") is not None:
            lefts.append(int(ind["left"]))
        if ind.get("right") is not None:
            rights.append(int(ind["right"]))
        if sp.get("before") is not None:
            befores.append(int(sp["before"]))
        if sp.get("after") is not None:
            afters.append(int(sp["after"]))

    return {
        "left": _median_int(lefts),
        "right": _median_int(rights),
        "before": _median_int(befores),
        "after": _median_int(afters),
    }


def _is_indent_outlier(value: int, profile: Dict[str, int]) -> bool:
    if value <= 0:
        return False
    hard_max = _env_int("PDF_DOCX_INDENT_OUTLIER_MAX", 2200)
    if value > hard_max:
        return True
    med = int(profile.get("left", 0))
    if med <= 0:
        return value > 1600
    return value > med + max(500, int(med * 0.65))


def _is_confident_academic_structure(paras: List) -> bool:
    """Only apply abstract/section heuristics when structure is clearly academic."""
    if not _env_bool("PDF_DOCX_ACADEMIC_LAYOUT", False):
        return False
    abs_idx = _find_abstract_start(paras)
    sec1 = _find_section_one_start(paras)
    sec2 = _find_section_two_start(paras, min_index=(sec1 + 1) if sec1 is not None else 0)
    if abs_idx is None or sec1 is None or sec2 is None:
        return False
    if not (abs_idx < sec1 < sec2):
        return False
    abs_text = _paragraph_plain(paras[abs_idx]) if abs_idx < len(paras) else ""
    if not _ABSTRACT_START_RE.match(abs_text):
        return False
    body_ref_idx = _find_body_reference_paragraph(paras, sec2)
    if body_ref_idx is None:
        return False
    body_ind = _read_ppr_ind_twips(paras[body_ref_idx]._element.find(qn("w:pPr")))
    if _is_indent_outlier(int(body_ind.get("left", 0)), {"left": 900}):
        return False
    return True


def _apply_abstract_layout(
    paragraph,
    *,
    body_ref,
    abstract_ref,
    spacing_before: Optional[int] = None,
) -> None:
    _copy_paragraph_layout_template(abstract_ref, paragraph, include_spacing=False)

    body_ppr = body_ref._element.find(qn("w:pPr"))
    abs_ppr = abstract_ref._element.find(qn("w:pPr"))
    body_ind = _read_ppr_ind_twips(body_ppr)
    abs_ind = _read_ppr_ind_twips(abs_ppr)

    body_left = int(body_ind.get("left", 0))
    body_right = int(body_ind.get("right", 0))
    abs_left = abs_ind.get("left")
    if abs_left is None:
        abs_left = body_left + _ABSTRACT_EXTRA_INDENT_TWIPS
    else:
        abs_left = int(abs_left)
    max_abs_left = body_left + max(_ABSTRACT_EXTRA_INDENT_TWIPS, 800)
    if abs_left > max_abs_left or _is_indent_outlier(abs_left, {"left": body_left}):
        abs_left = body_left + _ABSTRACT_EXTRA_INDENT_TWIPS
    abs_right = abs_ind.get("right")
    if abs_right is None:
        abs_right = body_right + (_ABSTRACT_EXTRA_INDENT_TWIPS if body_right else 0)
    else:
        abs_right = int(abs_right)

    abs_sp = _read_spacing_twips(abstract_ref)
    _set_paragraph_spacing(
        paragraph,
        line=abs_sp.get("line", 242),
        line_rule=str(abs_sp.get("lineRule", "exact")),
        before=spacing_before if spacing_before is not None else abs_sp.get("before", 0),
        after=abs_sp.get("after", 0),
    )
    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    _set_paragraph_indents(
        paragraph,
        left=abs_left,
        right=abs_right,
        first_line=0,
    )
    _clear_first_line_indent(paragraph)


def _apply_section1_body_layout(paragraph, body_ref, *, first_para: bool = False) -> None:
    _copy_paragraph_layout_template(body_ref, paragraph, include_spacing=False)
    body_ppr = body_ref._element.find(qn("w:pPr"))
    body_ind = _read_ppr_ind_twips(body_ppr)
    body_sp = _read_spacing_twips(body_ref)

    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    _set_paragraph_indents(
        paragraph,
        left=int(body_ind.get("left", 0)),
        right=int(body_ind.get("right", 0)),
        first_line=0 if first_para else int(body_ind.get("firstLine", 0)),
    )
    _set_paragraph_spacing(
        paragraph,
        line=body_sp.get("line", 240),
        line_rule=str(body_sp.get("lineRule", "exact")),
        before=body_sp.get("before", 0) if first_para else min(int(body_sp.get("before", 0) or 0), 120),
        after=0,
    )


def _compact_running_header(paragraph, profile: Optional[Dict[str, int]] = None) -> None:
    profile = profile or {}
    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT)
    _set_paragraph_spacing(paragraph, line=240, line_rule="exact", before=0, after=0)
    _set_paragraph_indents(
        paragraph,
        left=int(profile.get("left", 0)),
        right=int(profile.get("right", 0)),
        first_line=0,
    )


_FLAG_ITALIC = 1 << 1
_FLAG_BOLD = 1 << 4


def _normalize_match_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\r", " ").replace("\n", " ").strip().lower())


def _copy_run_rpr(src_run, dst_run) -> None:
    try:
        src_rpr = src_run._element.find(qn("w:rPr"))
        if src_rpr is None:
            return
        new_rpr = copy.deepcopy(src_rpr)
        dst_rpr = dst_run._element.find(qn("w:rPr"))
        if dst_rpr is not None:
            dst_run._element.remove(dst_rpr)
        dst_run._element.insert(0, new_rpr)
    except Exception:
        pass


def _dominant_source_run(runs: List) -> Optional[Any]:
    if not runs:
        return None
    best = runs[0]
    best_len = len(best.text or "")
    for run in runs[1:]:
        n = len(run.text or "")
        if n > best_len:
            best = run
            best_len = n
    return best


def _preserve_run_formats(src_para, dst_para) -> None:
    src_runs = list(src_para.runs)
    dst_runs = list(dst_para.runs)
    if not src_runs or not dst_runs:
        return
    if len(src_runs) == len(dst_runs):
        for src_run, dst_run in zip(src_runs, dst_runs):
            _copy_run_rpr(src_run, dst_run)
        return
    dominant = _dominant_source_run(src_runs)
    if dominant is None:
        return
    for dst_run in dst_runs:
        _copy_run_rpr(dominant, dst_run)


def _preserve_paragraph_style(src_para, dst_para) -> None:
    try:
        src_style = getattr(getattr(src_para, "style", None), "name", None)
        if src_style and src_style != "Normal":
            dst_para.style = src_style
    except Exception:
        pass


def _apply_pdf_format_hint(paragraph, hint: Dict[str, Any], *, apply_alignment: bool = False) -> None:
    if not hint:
        return
    if apply_alignment:
        align_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        align = align_map.get(str(hint.get("align") or "").lower())
        if align is not None:
            _set_paragraph_alignment(paragraph, align)

    font_size = hint.get("font_size")
    if font_size:
        try:
            _apply_font_pt(paragraph, float(font_size))
        except Exception:
            pass

    bold = hint.get("bold")
    italic = hint.get("italic")
    if bold is None and italic is None:
        return
    if _paragraph_has_mixed_bold(paragraph):
        return
    if _env_bool("PDF_DOCX_PRESERVE_LINES", True) and bold is not None:
        if _paragraph_should_be_bold(paragraph):
            for run in paragraph.runs:
                _set_run_bold(run, bool(bold))
        else:
            for run in paragraph.runs:
                _set_run_bold(run, False)
        return
    for run in paragraph.runs:
        try:
            if bold is not None:
                run.bold = bool(bold)
            if italic is not None:
                run.italic = bool(italic)
        except Exception:
            pass


def extract_pdf_paragraph_formats(pdf_path: str) -> List[Dict[str, Any]]:
    """Extract per-line formatting hints from the PDF text layer."""
    if not pdf_path or not os.path.isfile(pdf_path):
        return []
    try:
        import fitz  # PyMuPDF
    except Exception:
        return []

    formats: List[Dict[str, Any]] = []
    doc = fitz.open(pdf_path)
    try:
        for page in doc:
            try:
                page_width = float(page.rect.width)
            except Exception:
                page_width = 0.0
            try:
                blocks = page.get_text("dict").get("blocks") or []
            except Exception:
                blocks = []

            for block in blocks:
                if int(block.get("type") or 0) != 0:
                    continue
                for line in block.get("lines") or []:
                    spans = line.get("spans") or []
                    if not spans:
                        continue
                    text_raw = "".join(str(sp.get("text") or "") for sp in spans).strip()
                    if not text_raw:
                        continue
                    text = normalize_pdf_symbol_chars(text_raw)

                    sizes = [float(sp.get("size") or 0) for sp in spans if sp.get("text")]
                    font_size = max(sizes) if sizes else None
                    total_chars = sum(len(str(sp.get("text") or "")) for sp in spans)
                    bold_chars = sum(
                        len(str(sp.get("text") or ""))
                        for sp in spans
                        if int(sp.get("flags") or 0) & _FLAG_BOLD
                    )
                    italic_chars = sum(
                        len(str(sp.get("text") or ""))
                        for sp in spans
                        if int(sp.get("flags") or 0) & _FLAG_ITALIC
                    )

                    align = "left"
                    x0: Optional[float] = None
                    x1: Optional[float] = None
                    bbox = line.get("bbox") or block.get("bbox")
                    if bbox and page_width > 0:
                        try:
                            x0, _, x1, _ = [float(v) for v in bbox[:4]]
                        except Exception:
                            x0 = x1 = 0.0
                        width = max(0.0, x1 - x0)
                        left_m = max(0.0, x0)
                        right_m = max(0.0, page_width - x1)
                        center_x = (x0 + x1) * 0.5
                        mid = page_width * 0.5
                        sym_tol = max(8.0, page_width * 0.022)
                        cx_tol = max(6.0, page_width * 0.012)
                        if (
                            width <= page_width * 0.52
                            and abs(left_m - right_m) <= sym_tol
                            and abs(center_x - mid) <= cx_tol
                        ):
                            align = "center"
                        elif x1 >= page_width - max(24.0, page_width * 0.08) * 0.45 and center_x >= mid + page_width * 0.025:
                            align = "right"
                        elif width >= page_width * 0.58 and abs(left_m - right_m) <= max(12.0, page_width * 0.03):
                            align = "justify"

                    span_specs: List[Dict[str, Any]] = []
                    for sp in spans:
                        chunk = str(sp.get("text") or "")
                        if should_preserve_symbol_glyphs():
                            chunk = chunk.replace("\uf020", " ")
                        else:
                            chunk = normalize_pdf_symbol_chars(chunk)
                        if not chunk:
                            continue
                        span_specs.append(
                            {
                                "text": chunk,
                                "bold": bool(int(sp.get("flags") or 0) & _FLAG_BOLD),
                                "italic": bool(int(sp.get("flags") or 0) & _FLAG_ITALIC),
                                "size": float(sp.get("size") or 0) or None,
                            }
                        )

                    formats.append(
                        {
                            "text": text,
                            "text_norm": _normalize_match_text(
                                normalize_pdf_symbol_chars(text_raw, for_match=True)
                            ),
                            "font_size": font_size,
                            "bold": bool(total_chars and bold_chars * 100 >= total_chars * 60),
                            "italic": bool(total_chars and italic_chars * 100 >= total_chars * 50),
                            "align": align,
                            "font": spans[0].get("font") if spans else None,
                            "x0": x0,
                            "x1": x1,
                            "page_width": page_width,
                            "spans": span_specs,
                        }
                    )
    finally:
        doc.close()
    return formats


def _match_pdf_formats_to_paragraphs(
    para_texts: List[str],
    pdf_formats: List[Dict[str, Any]],
) -> List[Optional[Dict[str, Any]]]:
    results: List[Optional[Dict[str, Any]]] = [None] * len(para_texts)
    fmt_idx = 0
    for i, text in enumerate(para_texts):
        norm = _normalize_match_text(text)
        if not norm:
            continue
        best_j: Optional[int] = None
        best_score = 0.0
        search_end = min(len(pdf_formats), fmt_idx + 10)
        for j in range(fmt_idx, search_end):
            fmt_norm = pdf_formats[j].get("text_norm") or ""
            if not fmt_norm:
                continue
            if norm in fmt_norm or fmt_norm in norm:
                score = min(len(norm), len(fmt_norm)) / max(len(norm), len(fmt_norm), 1)
            else:
                a = set(norm.split())
                b = set(fmt_norm.split())
                if not a or not b:
                    continue
                score = len(a & b) / max(len(a), len(b))
            if score > best_score:
                best_score = score
                best_j = j
        if best_j is not None and best_score >= 0.35:
            results[i] = pdf_formats[best_j]
            fmt_idx = best_j + 1
    return results


def sync_formats_from_pdf(
    source_docx: str,
    translated_docx: str,
    pdf_path: str,
    *,
    force: bool = False,
) -> Dict[str, int]:
    stats = {"pdf_formats_applied": 0}
    if not force and not _env_bool("PDF_DOCX_PDF_FORMAT_SYNC", False):
        return stats
    if not os.path.isfile(source_docx) or not os.path.isfile(translated_docx):
        return stats

    pdf_formats = extract_pdf_paragraph_formats(pdf_path)
    if not pdf_formats:
        return stats

    src = docx.Document(source_docx)
    dst = docx.Document(translated_docx)
    src_texts = [_paragraph_plain(p) for p in src.paragraphs]
    hints = _match_pdf_formats_to_paragraphs(src_texts, pdf_formats)

    for i, hint in enumerate(hints):
        if hint is None or i >= len(dst.paragraphs):
            continue
        _apply_pdf_format_hint(
            dst.paragraphs[i],
            hint,
            apply_alignment=_env_bool("PDF_DOCX_PDF_ALIGN_SYNC", False),
        )
        stats["pdf_formats_applied"] += 1

    if stats["pdf_formats_applied"]:
        dst.save(translated_docx)
    return stats


def _sync_paragraph_style_only(src_para, dst_para) -> None:
    """Copy paragraph style only; leave indents/alignment for regional layout pass."""
    _preserve_paragraph_style(src_para, dst_para)


def _sync_paragraph_properties_only(src_para, dst_para) -> None:
    """Copy paragraph-level layout only; keep translated run text untouched."""
    _preserve_paragraph_style(src_para, dst_para)
    src_ppr = src_para._element.find(qn("w:pPr"))
    if src_ppr is None:
        return
    dst_ppr = dst_para._element.find(qn("w:pPr"))
    _copy_paragraph_properties(src_ppr, dst_ppr)


def _preserve_paragraph_layout(src_para, dst_para) -> None:
    _sync_paragraph_properties_only(src_para, dst_para)


def _append_runs_from_paragraph(src_para, dst_para, *, spacer: str = " ") -> None:
    src_runs = list(src_para.runs)
    if not src_runs:
        return
    if _paragraph_plain(dst_para):
        try:
            dst_para.add_run(spacer)
        except Exception:
            pass
    for run in src_runs:
        new_run = dst_para.add_run(run.text or "")
        try:
            if run.bold is not None:
                new_run.bold = run.bold
            if run.italic is not None:
                new_run.italic = run.italic
            if run.font.size is not None:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
        except Exception:
            pass


def _sync_region_by_offset(
    src_paras: List,
    dst_paras: List,
    src_start: int,
    dst_start: int,
) -> int:
    if src_start < 0 or dst_start < 0:
        return 0
    synced = 0
    max_k = min(len(src_paras) - src_start, len(dst_paras) - dst_start)
    for k in range(max(0, max_k)):
        _preserve_paragraph_layout(src_paras[src_start + k], dst_paras[dst_start + k])
        synced += 1
    return synced


def _should_merge_fragment(prev_para, next_para) -> bool:
    t1 = _paragraph_plain(prev_para)
    t2 = _paragraph_plain(next_para)
    if not t1 or not t2:
        return False
    if _is_section_heading_line(prev_para) or _is_section_heading_line(next_para):
        return False
    if _is_running_header_line(prev_para) or _is_running_header_line(next_para):
        return False
    if _SECTION1_HEAD_RE.match(t2):
        return False
    if _SECTION2_HEAD_RE.match(t2) and not _FALSE_SECTION2_RE.match(t2) and _paragraph_word_count(next_para) <= 8:
        return False
    if _KEYWORDS_LINE_RE.match(t2) and not _KEYWORDS_LINE_RE.match(t1):
        return False
    if _ABSTRACT_START_RE.match(t1) and _KEYWORDS_LINE_RE.match(t2):
        return False
    if _KEYWORDS_LINE_RE.match(t1) and (_SECTION1_HEAD_RE.match(t2) or re.match(r"^1\s+\S", t2)):
        return False
    if _paragraph_is_centered(prev_para) or _paragraph_is_centered(next_para):
        if _is_short_label_line(next_para) and re.match(r"^\d+\s+\S", t2):
            return False
        return True
    if not re.search(r"[.!?:;]\s*$", t1) and len(t2.split()) >= 3:
        return True
    return False


def _collapse_paragraph_range(
    doc: docx.Document,
    paras: List,
    start: int,
    end: int,
    *,
    allow_heading_merge: bool = False,
) -> int:
    merged = 0
    i = max(0, start)
    end = min(end, len(paras))
    while i < end - 1 and i < len(paras) - 1:
        cur = paras[i]
        nxt = paras[i + 1]
        if not allow_heading_merge and (
            _is_section_heading_line(cur) or _is_section_heading_line(nxt)
        ):
            i += 1
            continue
        if not _should_merge_fragment(cur, nxt):
            i += 1
            continue
        _append_runs_from_paragraph(nxt, cur)
        try:
            nxt._element.getparent().remove(nxt._element)
        except Exception:
            break
        paras.pop(i + 1)
        end = min(end, len(paras))
        merged += 1
    return merged


def _collapse_pre_section2_fragments(doc: docx.Document, paras: List, sec2_start: int) -> int:
    if sec2_start <= 0:
        return 0

    merged = 0
    title_end = _find_title_block_end(paras)
    if title_end > 1:
        merged += _collapse_title_block_fragments(doc, paras, title_end)
        title_end = _find_title_block_end(paras)
    kw_idx = _find_keywords_paragraph_index(paras, title_end)
    sec1_start = _find_section_one_start(paras)

    abstract_end = kw_idx if kw_idx is not None else (sec1_start if sec1_start is not None else sec2_start)
    if abstract_end > title_end:
        merged += _collapse_paragraph_range(doc, paras, title_end, abstract_end)

    title_end = _find_title_block_end(paras)
    kw_idx = _find_keywords_paragraph_index(paras, title_end)
    sec1_start = _find_section_one_start(paras)
    sec2_start = _find_section_two_start(paras) or sec2_start

    if kw_idx is not None:
        kw_end = sec1_start if sec1_start is not None else sec2_start
        if kw_end > kw_idx:
            merged += _collapse_paragraph_range(doc, paras, kw_idx, kw_end)

    sec1_start = _find_section_one_start(paras)
    sec2_start = _find_section_two_start(paras) or sec2_start

    if sec1_start is not None and sec2_start > sec1_start + 1:
        merged += _collapse_paragraph_range(
            doc,
            paras,
            sec1_start + 1,
            sec2_start,
            allow_heading_merge=False,
        )

    return merged


def _normalize_title_block(
    paras: List,
    src_paras: List,
    *,
    title_end: int,
    src_title_end: int,
) -> int:
    limit = min(title_end, src_title_end, len(paras), len(src_paras))
    for i in range(limit):
        _preserve_paragraph_layout(src_paras[i], paras[i])
    return limit


def _normalize_body_block(
    paras: List,
    src_paras: List,
    *,
    start: int,
    end: int,
    body_ref,
    abstract_ref,
    region: str,
    src_heading_para=None,
    dst_heading_idx: Optional[int] = None,
    src_abs_start: Optional[int] = None,
    src_kw_idx: Optional[int] = None,
    src_sec1_start: Optional[int] = None,
) -> int:
    fixed = 0
    if src_heading_para is not None and dst_heading_idx is not None:
        if 0 <= dst_heading_idx < len(paras):
            _preserve_paragraph_layout(src_heading_para, paras[dst_heading_idx])
            fixed += 1

    first_section1_body = True
    for i in range(start, min(end, len(paras))):
        if dst_heading_idx is not None and i == dst_heading_idx:
            continue
        para = paras[i]
        if _is_running_header_line(para):
            _compact_running_header(para)
            fixed += 1
            continue
        if _is_section_heading_line(para):
            if src_heading_para is not None and i == dst_heading_idx:
                fixed += 1
                continue
            src_i = (
                src_sec1_start
                if region == "section1" and src_sec1_start is not None
                else i
            )
            if src_i < len(src_paras):
                _preserve_paragraph_layout(src_paras[src_i], para)
            fixed += 1
            continue

        if region == "abstract":
            text = _paragraph_plain(para)
            before = None
            if _KEYWORDS_LINE_RE.match(text) and src_kw_idx is not None and src_kw_idx < len(src_paras):
                before = _read_spacing_twips(src_paras[src_kw_idx]).get("before")
            elif src_abs_start is not None and src_abs_start < len(src_paras):
                before = _read_spacing_twips(src_paras[src_abs_start]).get("before")
            _apply_abstract_layout(
                para,
                body_ref=body_ref,
                abstract_ref=abstract_ref,
                spacing_before=before,
            )
        else:
            _apply_section1_body_layout(para, body_ref, first_para=first_section1_body)
            first_section1_body = False
        fixed += 1
    return fixed


def _refresh_layout_after_merge(
    paras: List,
    src_paras: List,
    *,
    title_end: int,
    src_title_end: int,
    sec1_start: Optional[int],
    src_sec1_start: Optional[int],
    sec2_start: Optional[int],
    src_sec2_start: Optional[int],
    body_ref,
    abstract_ref,
) -> int:
    fixed = 0
    sec1_end = sec2_start if sec2_start is not None else len(paras)
    abstract_end = sec1_start if sec1_start is not None else sec1_end

    fixed += _normalize_title_block(
        paras,
        src_paras,
        title_end=title_end,
        src_title_end=src_title_end,
    )
    if body_ref is not None and abstract_ref is not None:
        src_abs = src_title_end
        src_kw = _find_keywords_paragraph_index(src_paras, src_title_end)
        fixed += _normalize_body_block(
            paras,
            src_paras,
            start=title_end,
            end=abstract_end,
            body_ref=body_ref,
            abstract_ref=abstract_ref,
            region="abstract",
            src_abs_start=src_abs,
            src_kw_idx=src_kw,
        )
    if body_ref is not None and sec1_start is not None:
        src_heading = (
            src_paras[src_sec1_start]
            if src_sec1_start is not None and src_sec1_start < len(src_paras)
            else None
        )
        fixed += _normalize_body_block(
            paras,
            src_paras,
            start=sec1_start,
            end=sec1_end,
            body_ref=body_ref,
            abstract_ref=abstract_ref or body_ref,
            region="section1",
            src_heading_para=src_heading,
            dst_heading_idx=sec1_start,
            src_sec1_start=src_sec1_start,
        )
    return fixed


def normalize_document_layout(
    doc: docx.Document,
    src_doc: docx.Document,
) -> Dict[str, int]:
    stats = {
        "alignment_fixed": 0,
        "indents_fixed": 0,
        "fonts_normalized": 0,
        "title_centered": 0,
        "fragments_merged": 0,
    }
    paras = list(doc.paragraphs)
    src_paras = list(src_doc.paragraphs)

    if not _is_confident_academic_structure(src_paras):
        return stats

    src_title_end = _find_title_block_end(src_paras)
    src_sec1_start = _find_section_one_start(src_paras)
    src_sec2_start = _find_section_two_start(src_paras)

    if src_sec2_start is None:
        return stats

    body_ref_idx = _find_body_reference_paragraph(src_paras, src_sec2_start)
    if body_ref_idx is None:
        return stats
    body_ref = src_paras[body_ref_idx]

    abs_start = src_title_end
    abs_end = src_sec1_start if src_sec1_start is not None else src_sec2_start
    abs_ref_idx = _find_abstract_reference_paragraph(src_paras, abs_start, abs_end)
    abstract_ref = src_paras[abs_ref_idx if abs_ref_idx is not None else abs_start]

    dst_sec2_start = _find_section_two_start(paras)
    if _env_bool("PDF_DOCX_MERGE_FRAGMENTS", False):
        stats["fragments_merged"] = _collapse_pre_section2_fragments(
            doc,
            paras,
            dst_sec2_start if dst_sec2_start is not None else len(paras),
        )
        paras = list(doc.paragraphs)

    title_end = _find_title_block_end(paras)
    sec1_start = _find_section_one_start(paras)
    sec2_start = _find_section_two_start(paras)

    fixed = _refresh_layout_after_merge(
        paras,
        src_paras,
        title_end=title_end,
        src_title_end=src_title_end,
        sec1_start=sec1_start,
        src_sec1_start=src_sec1_start,
        sec2_start=sec2_start,
        src_sec2_start=src_sec2_start,
        body_ref=body_ref,
        abstract_ref=abstract_ref,
    )
    stats["alignment_fixed"] = fixed
    stats["indents_fixed"] = fixed
    stats["fonts_normalized"] = fixed
    stats["title_centered"] = min(title_end, len(paras))
    return stats


def _sync_table_layout_from_source(src_doc: docx.Document, dst_doc: docx.Document) -> int:
    synced = 0
    if not _env_bool("PDF_DOCX_TABLE_SYNC", True):
        return 0
    for t_idx, src_table in enumerate(src_doc.tables):
        if t_idx >= len(dst_doc.tables):
            break
        dst_table = dst_doc.tables[t_idx]
        for r_idx, src_row in enumerate(src_table.rows):
            if r_idx >= len(dst_table.rows):
                break
            dst_row = dst_table.rows[r_idx]
            for c_idx, src_cell in enumerate(src_row.cells):
                if c_idx >= len(dst_row.cells):
                    break
                dst_cell = dst_row.cells[c_idx]
                src_cell_paras = list(src_cell.paragraphs)
                dst_cell_paras = list(dst_cell.paragraphs)
                for i in range(min(len(src_cell_paras), len(dst_cell_paras))):
                    _preserve_paragraph_layout(src_cell_paras[i], dst_cell_paras[i])
                    synced += 1
    return synced


def _sync_header_footer_layout(src_doc: docx.Document, dst_doc: docx.Document) -> int:
    synced = 0
    if not _env_bool("PDF_DOCX_HEADER_FOOTER_SYNC", True):
        return 0
    try:
        for i, src_sec in enumerate(src_doc.sections):
            if i >= len(dst_doc.sections):
                break
            dst_sec = dst_doc.sections[i]
            for src_p, dst_p in zip(src_sec.header.paragraphs, dst_sec.header.paragraphs):
                _preserve_paragraph_layout(src_p, dst_p)
                synced += 1
            for src_p, dst_p in zip(src_sec.footer.paragraphs, dst_sec.footer.paragraphs):
                _preserve_paragraph_layout(src_p, dst_p)
                synced += 1
    except Exception:
        pass
    return synced


def _align_paragraph_pairs(src_paras: List, dst_paras: List) -> List[Tuple[int, int]]:
    """Pair src/dst paragraphs sequentially, skipping empty-only rows on either side."""
    pairs: List[Tuple[int, int]] = []
    si = 0
    di = 0
    while si < len(src_paras) and di < len(dst_paras):
        src_t = _paragraph_plain(src_paras[si])
        dst_t = _paragraph_plain(dst_paras[di])
        if not src_t and not dst_t:
            pairs.append((si, di))
            si += 1
            di += 1
            continue
        if not src_t:
            si += 1
            continue
        if not dst_t:
            di += 1
            continue
        pairs.append((si, di))
        si += 1
        di += 1
    return pairs


def mirror_document_layout_from_source(
    source_docx: str,
    translated_docx: str,
) -> Dict[str, int]:
    """Copy paragraph/table/header properties from source DOCX onto translated text."""
    stats = {
        "mirrored": 0,
        "mismatched_paragraphs": 0,
        "table_cells_synced": 0,
        "header_footer_synced": 0,
    }
    if not _env_bool("PDF_DOCX_MIRROR_LAYOUT", True) and not _env_bool("PDF_DOCX_LAYOUT_SYNC", True):
        return stats

    src = docx.Document(source_docx)
    dst = docx.Document(translated_docx)

    pair_count = min(len(src.paragraphs), len(dst.paragraphs))
    stats["mismatched_paragraphs"] = abs(len(src.paragraphs) - len(dst.paragraphs))
    for i in range(pair_count):
        _sync_paragraph_properties_only(src.paragraphs[i], dst.paragraphs[i])
        stats["mirrored"] += 1

    stats["table_cells_synced"] = _sync_table_layout_from_source(src, dst)
    stats["header_footer_synced"] = _sync_header_footer_layout(src, dst)
    stats["mirrored"] += stats["table_cells_synced"] + stats["header_footer_synced"]
    dst.save(translated_docx)
    return stats


def sync_docx_layout_from_source(
    source_docx: str,
    translated_docx: str,
) -> Dict[str, int]:
    """Backward-compatible alias: full mirror from source DOCX."""
    mirror_stats = mirror_document_layout_from_source(source_docx, translated_docx)
    return {
        "paragraphs_synced": int(mirror_stats.get("mirrored", 0)),
        "alignment_fixed": int(mirror_stats.get("mirrored", 0)),
        "mismatched_paragraphs": int(mirror_stats.get("mismatched_paragraphs", 0)),
        "table_cells_synced": int(mirror_stats.get("table_cells_synced", 0)),
        "header_footer_synced": int(mirror_stats.get("header_footer_synced", 0)),
    }


def _fit_paragraph_fonts(paragraph, *, min_scale: float, min_pt: float) -> int:
    changed = 0
    text = _paragraph_plain(paragraph)
    if len(text) < 20:
        return 0
    if _is_in_table_cell(paragraph):
        return 0
    base = _dominant_font_pt(paragraph)
    if base is None or base <= min_pt:
        return 0
    est_chars = max(1, int(base * 2.2))
    if len(text) <= est_chars:
        return 0
    ratio = max(min_scale, est_chars / float(len(text)))
    new_pt = max(min_pt, base * ratio)
    if new_pt < base - 0.2:
        _apply_font_pt(paragraph, new_pt)
        changed += 1
    return changed


def _fit_document_fonts(doc: docx.Document) -> int:
    if not _env_bool("PDF_DOCX_FONT_FIT", False):
        return 0
    min_scale = _env_float("PDF_DOCX_FONT_FIT_MIN_SCALE", 0.62)
    min_pt = _env_float("PDF_DOCX_FONT_FIT_MIN_PT", 7.0)
    fitted = 0
    for paragraph in iter_all_paragraphs(doc):
        fitted += _fit_paragraph_fonts(
            paragraph,
            min_scale=min_scale,
            min_pt=min_pt,
        )
    return fitted


def _recover_tables_and_images(doc: docx.Document) -> Dict[str, int]:
    out = {"table_rows_relaxed": 0, "images_resized": 0}
    try:
        from docx.enum.table import WD_ROW_HEIGHT_RULE
    except Exception:
        WD_ROW_HEIGHT_RULE = None  # type: ignore

    if WD_ROW_HEIGHT_RULE is not None:
        for table in doc.tables:
            for row in table.rows:
                row_changed = False
                try:
                    if row.height is not None:
                        row.height = None
                        row_changed = True
                except Exception:
                    pass
                try:
                    if row.height_rule != WD_ROW_HEIGHT_RULE.AT_LEAST:
                        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                        row_changed = True
                except Exception:
                    pass
                if row_changed:
                    out["table_rows_relaxed"] += 1

    max_inline_width = None
    try:
        if doc.sections:
            sec = doc.sections[0]
            max_inline_width = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin)
    except Exception:
        max_inline_width = None

    if max_inline_width and max_inline_width > 0:
        for shp in getattr(doc, "inline_shapes", []):
            try:
                cur_w = int(shp.width or 0)
                cur_h = int(shp.height or 0)
            except Exception:
                continue
            if cur_w <= 0 or cur_h <= 0 or cur_w <= max_inline_width:
                continue
            try:
                ratio = float(max_inline_width) / float(cur_w)
                shp.width = int(max_inline_width)
                shp.height = max(1, int(cur_h * ratio))
                out["images_resized"] += 1
            except Exception:
                pass
    return out


def _fix_indents_from_pdf_geometry(
    doc: docx.Document,
    src_doc: docx.Document,
    pdf_path: str,
    profile: Dict[str, int],
) -> int:
    """Reset DOCX indents when PDF geometry shows normal left-margin body text."""
    pdf_formats = extract_pdf_paragraph_formats(pdf_path)
    if not pdf_formats:
        return 0

    body_x0: List[float] = []
    for fmt in pdf_formats:
        pw = float(fmt.get("page_width") or 0)
        x0 = fmt.get("x0")
        x1 = fmt.get("x1")
        if pw <= 0 or x0 is None or x1 is None:
            continue
        width = float(x1) - float(x0)
        if width >= pw * 0.42 and float(x0) / pw <= 0.18:
            body_x0.append(float(x0))
    if not body_x0:
        return 0

    body_x0.sort()
    median_x0 = body_x0[len(body_x0) // 2]
    fixed = 0
    src_texts = [_paragraph_plain(p) for p in src_doc.paragraphs]
    hints = _match_pdf_formats_to_paragraphs(src_texts, pdf_formats)

    for i, hint in enumerate(hints):
        if not hint or i >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[i]
        if not _paragraph_plain(para) or _is_in_table_cell(para):
            continue

        pw = float(hint.get("page_width") or 0)
        x0 = hint.get("x0")
        x1 = hint.get("x1")
        if pw <= 0 or x0 is None or x1 is None:
            continue
        width = float(x1) - float(x0)
        near_body = abs(float(x0) - median_x0) <= pw * 0.06
        full_width = width >= pw * 0.38
        if not (near_body and full_width):
            continue

        ppr = para._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))
        right = int(ind.get("right", 0))
        content_twips = _get_content_width_twips(doc)
        wrong_width = (
            _is_indent_outlier(left, profile)
            or _is_abstract_style_indent(left, right)
            or _is_narrow_column_indent(left, right, content_twips)
        )
        if not wrong_width:
            continue

        preserved_first = int(ind.get("firstLine", 0))
        if _paragraph_has_hanging_indent(para):
            _set_paragraph_indents(
                para,
                left=int(profile.get("left", 0)),
                right=int(profile.get("right", 0)),
            )
        else:
            _set_paragraph_indents(
                para,
                left=int(profile.get("left", 0)),
                right=int(profile.get("right", 0)),
                first_line=preserved_first,
            )
        align = str(hint.get("align") or "left").lower()
        if align == "center" and width >= pw * 0.38:
            align = "justify"
        align_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        _set_paragraph_alignment(para, align_map.get(align, WD_PARAGRAPH_ALIGNMENT.JUSTIFY))
        fixed += 1
    return fixed


def _collapse_excessive_blank_paragraphs(doc: docx.Document) -> int:
    """Remove stacked empty paragraphs that create large vertical gaps."""
    removed = 0
    paras = list(doc.paragraphs)
    blank_streak = 0
    idx = 0
    while idx < len(paras):
        if not _paragraph_plain(paras[idx]):
            blank_streak += 1
            if blank_streak > 1:
                try:
                    paras[idx]._element.getparent().remove(paras[idx]._element)
                    paras.pop(idx)
                    removed += 1
                    continue
                except Exception:
                    pass
        else:
            blank_streak = 0
        idx += 1
    return removed


def fix_layout_anomalies(
    doc: docx.Document,
    src_doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
) -> Dict[str, int]:
    """Detect and fix outlier indents/spacing relative to each document's profile."""
    stats = {
        "anomalies_fixed": 0,
        "spacing_fixed": 0,
        "pdf_indent_fixed": 0,
        "empties_collapsed": 0,
    }
    if not _env_bool("PDF_DOCX_FIX_ANOMALIES", True):
        return stats

    src_body = [
        p
        for p in src_doc.paragraphs
        if _paragraph_plain(p) and not _is_in_table_cell(p) and not _paragraph_is_centered(p)
    ]
    profile = _compute_body_layout_profile(src_body)

    for i, para in enumerate(doc.paragraphs):
        if not _paragraph_plain(para) or _is_in_table_cell(para) or _paragraph_is_centered(para):
            continue

        ppr = para._element.find(qn("w:pPr"))
        ind = _read_ppr_ind_twips(ppr)
        left = int(ind.get("left", 0))

        if _is_indent_outlier(left, profile):
            restored = False
            if i < len(src_doc.paragraphs):
                src_para = src_doc.paragraphs[i]
                src_ind = _read_ppr_ind_twips(src_para._element.find(qn("w:pPr")))
                if not _is_indent_outlier(int(src_ind.get("left", 0)), profile):
                    _preserve_paragraph_layout(src_para, para)
                    restored = True
            if not restored:
                _set_paragraph_indents(
                    para,
                    left=int(profile.get("left", 0)),
                    right=int(profile.get("right", 0)),
                    first_line=0,
                )
            stats["anomalies_fixed"] += 1

        sp = _read_spacing_twips(para)
        med_before = int(profile.get("before", 0))
        med_after = int(profile.get("after", 0))
        cap_before = max(280, med_before * 4, 600)
        cap_after = max(280, med_after * 4, 600)
        before = sp.get("before")
        after = sp.get("after")
        new_before, new_after = before, after
        if before is not None and int(before) > cap_before:
            new_before = med_before
        if after is not None and int(after) > cap_after:
            new_after = med_after
        if new_before != before or new_after != after:
            _set_paragraph_spacing(para, before=new_before, after=new_after)
            stats["spacing_fixed"] += 1

    if pdf_path and os.path.isfile(pdf_path):
        stats["pdf_indent_fixed"] = _fix_indents_from_pdf_geometry(doc, src_doc, pdf_path, profile)

    stats["empties_collapsed"] = _collapse_excessive_blank_paragraphs(doc)
    return stats


def _normalize_paragraph_match_key(text: str) -> str:
    t = re.sub(r"\s+", " ", (text or "").replace("\u00ad", "").strip().lower())
    return t


def _paragraph_texts_match(a: str, b: str) -> bool:
    ka = _normalize_paragraph_match_key(a)
    kb = _normalize_paragraph_match_key(b)
    if not ka or not kb:
        return False
    if ka == kb:
        return True
    ka_compact = re.sub(r"\s+", "", ka)
    kb_compact = re.sub(r"\s+", "", kb)
    if ka_compact == kb_compact:
        return True
    shorter, longer = (ka, kb) if len(ka) <= len(kb) else (kb, ka)
    prefix_len = min(len(shorter), 120)
    if prefix_len >= 24 and longer.startswith(shorter[:prefix_len]):
        return True
    return False


def _is_affiliation_line(paragraph) -> bool:
    text = _paragraph_plain(paragraph)
    if not text:
        return False
    if _AFFILIATION_LINE_RE.match(text):
        return True
    if re.match(
        r"^\d+\s*(faculty|department|university|springer|institute|fakulteit|khoa|tr[uườ]ng)\b",
        text,
        re.IGNORECASE,
    ):
        return True
    low = text.lower()
    return bool(
        re.match(r"^\d+\s*", text)
        and any(k in low for k in ("faculty", "department", "university", "institute", "khoa", "trường", "truong"))
    )


def _looks_like_subtitle_line(text: str) -> bool:
    t = (text or "").strip()
    if not t or "," in t or "@" in t:
        return False
    words = [w for w in t.split() if w]
    if not 1 <= len(words) <= 5:
        return False
    return all(w[0].isupper() for w in words if w and w[0].isalpha())


def _should_merge_title_block_fragment(prev_para, next_para) -> bool:
    t1 = _paragraph_plain(prev_para).strip()
    t2 = _paragraph_plain(next_para).strip()
    if not t1 or not t2:
        return False
    if _ABSTRACT_START_RE.match(t2) or _ABSTRACT_START_RE.match(t1):
        return False
    if _KEYWORDS_LINE_RE.match(t2):
        return False

    if _looks_like_subtitle_line(t2) and _paragraph_word_count(prev_para) >= 5:
        return False

    if _is_affiliation_line(prev_para) or re.match(r"^\d+\s*\S", t1):
        low1 = t1.lower()
        if any(k in low1 for k in ("faculty", "department", "university", "institute", "street", "ward", "city")):
            if not re.match(
                r"^\d+\s*(faculty|department|university|institute|khoa|tr[uườ]ng)\b",
                t2,
                re.IGNORECASE,
            ):
                if "@" not in t2:
                    return True

    if "@" in t1:
        return "@" in t2 or bool(re.match(r"^[a-z0-9._-]+@", t2, re.IGNORECASE))
    if "@" in t2:
        return False

    if re.search(r",\s*[\w\-]{1,12}$", t1) and not re.match(r"^\d", t2):
        first = (t2.split() or [""])[0]
        if first and first[0].isupper():
            return True

    if _paragraph_is_centered(prev_para) and _paragraph_is_centered(next_para):
        low1 = t1.lower()
        if " and " in low1 or re.search(r"\d+\*?\s*$", t1):
            return False
        if re.match(r"^\d", t2):
            return False
        if len(t2.split()) >= 2 and not _looks_like_subtitle_line(t2):
            return True

    return False


def _collapse_title_block_fragments(doc: docx.Document, paras: List, title_end: int) -> int:
    """Merge pdf2docx line wraps inside title/author/affiliation/email block."""
    merged = 0
    i = 0
    end = min(max(0, title_end), len(paras))
    while i < end - 1 and i < len(paras) - 1:
        cur = paras[i]
        nxt = paras[i + 1]
        if not _should_merge_title_block_fragment(cur, nxt):
            i += 1
            continue
        _append_runs_from_paragraph(nxt, cur, spacer=" ")
        try:
            nxt._element.getparent().remove(nxt._element)
        except Exception:
            break
        paras.pop(i + 1)
        end = min(end, len(paras))
        merged += 1
    return merged


def _apply_body_indent_from_source(
    paragraph,
    src_para,
    body_profile: Dict[str, int],
    *,
    is_first_after_heading: bool,
) -> None:
    """Apply body left/right + first-line/hanging indent from source paragraph."""
    src_ind = _read_ppr_ind_twips(src_para._element.find(qn("w:pPr")))
    left = int(body_profile["left"])
    right = int(body_profile["right"])
    _set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    if is_first_after_heading:
        _set_paragraph_indents(paragraph, left=left, right=right, first_line=0, hanging=0)
        return

    first_line = int(src_ind.get("firstLine", 0) or 0)
    hanging = int(src_ind.get("hanging", 0) or 0)
    if first_line <= 0 and hanging <= 0:
        profile_fl = int(body_profile.get("firstLine", 0) or 0)
        if profile_fl > 0:
            first_line = profile_fl

    if hanging > 0:
        _set_paragraph_indents(
            paragraph,
            left=max(left, int(src_ind.get("left", left) or left)),
            right=right,
            hanging=hanging,
        )
    else:
        _set_paragraph_indents(
            paragraph,
            left=left,
            right=right,
            first_line=first_line,
        )


def _rpr_font_size(rpr) -> int:
    if rpr is None:
        return -1
    try:
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            raw = str(sz.get(qn("w:val"), "") or "").strip()
            if raw.isdigit():
                return int(raw)
    except Exception:
        pass
    return -1


def _run_is_superscript_run(run) -> bool:
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return False
        va = rpr.find(qn("w:vertAlign"))
        if va is not None:
            v = str(va.get(qn("w:val"), "") or "").strip().lower()
            return v in ("superscript", "subscript")
    except Exception:
        pass
    return False


def _run_is_bold(run) -> bool:
    try:
        if run.bold is True:
            return True
        if getattr(getattr(run, "font", None), "bold", None) is True:
            return True
    except Exception:
        pass
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return False
        b = rpr.find(qn("w:b"))
        if b is None:
            return False
        val = str(b.get(qn("w:val"), "") or "").strip().lower()
        return val not in ("0", "false", "off")
    except Exception:
        return False


def _set_run_bold(run, bold: bool) -> None:
    try:
        run.bold = bool(bold)
    except Exception:
        pass
    try:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            if not bold:
                return
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        for tag in ("w:b", "w:bCs"):
            old = rpr.find(qn(tag))
            if old is not None:
                rpr.remove(old)
            if bold:
                el = OxmlElement(tag)
                el.set(qn("w:val"), "1")
                rpr.append(el)
    except Exception:
        pass


def _copy_run_rpr_style(src_run, dst_run, *, include_fonts: bool = False) -> bool:
    src_rpr = src_run._element.find(qn("w:rPr"))
    if src_rpr is None:
        return False
    try:
        dst_rpr = dst_run._element.find(qn("w:rPr"))
        if dst_rpr is None:
            dst_rpr = OxmlElement("w:rPr")
            dst_run._element.insert(0, dst_rpr)
        tags = ["w:sz", "w:szCs", "w:b", "w:bCs", "w:i", "w:iCs", "w:vertAlign"]
        if include_fonts:
            tags.insert(0, "w:rFonts")
        for tag in tags:
            src_el = src_rpr.find(qn(tag))
            old = dst_rpr.find(qn(tag))
            if old is not None:
                dst_rpr.remove(old)
            if src_el is not None:
                dst_rpr.append(copy.deepcopy(src_el))
        return True
    except Exception:
        return False


def _paragraph_has_mixed_bold(paragraph) -> bool:
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    if len(runs) < 2:
        return False
    states = {_run_is_bold(r) for r in runs}
    return len(states) > 1


def _paragraph_has_heading_typography(paragraph) -> bool:
    """Short line with a visibly larger dominant run (heading size without bold flag)."""
    if _paragraph_word_count(paragraph) > 24:
        return False
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    if not runs:
        return False

    def _is_list_marker_run(run) -> bool:
        text = (run.text or "").strip()
        if len(text) <= 2 and text.lower() in ("o", "•", "-", "+", "*", "·"):
            return True
        return bool(re.match(r"^[\u2022\u25cf\u25cb\u2013\-\+]$", text))

    content_runs = [r for r in runs if not _is_list_marker_run(r)]
    if not content_runs:
        return False

    sizes = [_rpr_font_size(r._element.find(qn("w:rPr"))) for r in content_runs]
    sizes = [s for s in sizes if s > 0]
    if not sizes:
        return False
    max_sz = max(sizes)
    if len(content_runs) == 1 and len((content_runs[0].text or "").strip()) <= 2:
        return False
    non_bold_sizes = [
        _rpr_font_size(r._element.find(qn("w:rPr")))
        for r in content_runs
        if not _run_is_bold(r)
    ]
    non_bold_sizes = [s for s in non_bold_sizes if s > 0]
    ref = min(non_bold_sizes) if non_bold_sizes else min(sizes)
    return max_sz >= ref + 4


def _paragraph_should_be_bold(paragraph) -> bool:
    if _paragraph_has_mixed_bold(paragraph):
        return False
    if _is_section_heading_line(paragraph):
        return True
    if _paragraph_has_heading_typography(paragraph):
        return True
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    if not runs:
        return False
    total = sum(len(r.text or "") for r in runs)
    bold_chars = sum(len(r.text or "") for r in runs if _run_is_bold(r))
    return total > 0 and (bold_chars * 100 >= total * 70)


def _pick_heading_source_run(paragraph):
    best_run = None
    best_key = (-1, False)
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        sz = _rpr_font_size(run._element.find(qn("w:rPr")))
        key = (sz, _run_is_bold(run))
        if key > best_key:
            best_key = key
            best_run = run
    return best_run


def _heading_run_rpr(paragraph):
    run = _pick_heading_source_run(paragraph)
    if run is None:
        return None
    return run._element.find(qn("w:rPr"))


def _apply_template_rpr_to_run(run, template_rpr, para=None) -> None:
    template_rpr = _sanitize_rpr_body_font(template_rpr, para)
    if template_rpr is None:
        return
    try:
        old = run._element.find(qn("w:rPr"))
        if old is not None:
            run._element.remove(old)
        run._element.insert(0, copy.deepcopy(template_rpr))
    except Exception:
        pass


def pick_translation_template_rpr(paragraph):
    """Choose run properties when translated text replaces a whole paragraph."""
    if _paragraph_should_be_bold(paragraph):
        heading = _heading_run_rpr(paragraph)
        if heading is not None:
            return _sanitize_rpr_body_font(heading, paragraph)
    prose = _prose_run_rpr(paragraph)
    if prose is not None:
        cloned = copy.deepcopy(prose)
        if _paragraph_has_mixed_bold(paragraph) or not _paragraph_should_be_bold(paragraph):
            try:
                dst_rpr = cloned
                for tag in ("w:b", "w:bCs"):
                    old = dst_rpr.find(qn(tag))
                    if old is not None:
                        dst_rpr.remove(old)
            except Exception:
                pass
        return _sanitize_rpr_body_font(cloned, paragraph)
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        fname = _read_run_font_name(run)
        if _is_auxiliary_pdf_font(fname) or _text_has_pua_symbol(run.text or ""):
            continue
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            return _sanitize_rpr_body_font(rpr, paragraph)
    return _sanitize_rpr_body_font(_fallback_body_run_rpr(paragraph), paragraph)


def _prose_run_rpr(paragraph):
    """Median-size non-bold run rPr — avoids picking heading sizes for body text."""
    candidates: List[tuple] = []
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        if _run_is_superscript_run(run):
            continue
        fname = _read_run_font_name(run)
        if fname and _is_auxiliary_pdf_font(fname):
            continue
        rpr = run._element.find(qn("w:rPr"))
        sz = _rpr_font_size(rpr)
        if sz <= 0:
            continue
        candidates.append((_run_is_bold(run), sz, rpr))
    if not candidates:
        return _dominant_body_run_rpr(paragraph)
    non_bold = [c for c in candidates if not c[0]]
    pool = non_bold if non_bold else candidates
    pool.sort(key=lambda item: item[1])
    return pool[len(pool) // 2][2]


def _line_rpr_for_text(runs: List, line_text: str):
    needle = (line_text or "").strip()[:24]
    if not needle:
        return None
    cursor = 0
    for run in runs:
        chunk = run.text or ""
        start = cursor
        cursor += len(chunk)
        if needle in chunk or needle in ("".join(r.text or "" for r in runs)[start:cursor]):
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    for run in runs:
        if (run.text or "").strip():
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    return None


_FORM_FIELD_LINE_RE = re.compile(
    r"(?:TV\s*\d|MSSV|Email|Class|Lớp|Họ\s*&?\s*Tên|Tên\s*Đề\s*tài|Development\s+direction)",
    re.IGNORECASE | re.UNICODE,
)


def _is_form_field_paragraph(text: str) -> bool:
    return bool(_FORM_FIELD_LINE_RE.search(text or ""))


def _pick_style_source_run(src_runs: List, index: int) -> object:
    if not src_runs:
        raise ValueError("src_runs empty")
    src_i = min(index, len(src_runs) - 1)
    candidate = src_runs[src_i]
    fname = _read_run_font_name(candidate)
    if _is_auxiliary_pdf_font(fname) or _text_has_pua_symbol(candidate.text or ""):
        for run in src_runs:
            rt = run.text or ""
            if not rt.strip():
                continue
            fn = _read_run_font_name(run)
            if _is_auxiliary_pdf_font(fn) or _text_has_pua_symbol(rt):
                continue
            return run
    return candidate


def _sync_run_styles_from_source(src_para, dst_para) -> int:
    """Restore bold/italic/size on translated runs from pre-translation paragraph."""
    src_runs = [r for r in src_para.runs if (r.text or "")]
    dst_runs = [r for r in dst_para.runs if (r.text or "")]
    if not src_runs or not dst_runs:
        return 0

    synced = 0
    if len(dst_runs) == 1:
        dst_text = dst_runs[0].text or ""
        regions = _collect_run_style_regions(src_runs)
        if (
            len(regions) > 1
            and dst_text
            and not _src_has_auxiliary_font_runs(src_runs)
            and not (len(dst_runs) == 1 and len(regions) >= 4)
        ):
            return _rebuild_para_runs_from_style_regions(dst_para, dst_text, regions)
        dst_run = dst_runs[0]
        src_plain = _paragraph_literal(src_para)
        if _is_form_field_paragraph(src_plain):
            prose = next(
                (
                    r
                    for r in src_runs
                    if not _is_auxiliary_pdf_font(_read_run_font_name(r))
                    and not _text_has_pua_symbol(r.text or "")
                ),
                src_runs[-1],
            )
            _copy_run_rpr_style(prose, dst_run)
            _set_run_bold(dst_run, False)
            synced += 1
            return synced
        if _paragraph_should_be_bold(src_para) and not _paragraph_has_mixed_bold(src_para):
            _copy_run_rpr_style(_pick_heading_source_run(src_para) or src_runs[0], dst_run)
        else:
            prose = next(
                (
                    r
                    for r in src_runs
                    if not _run_is_bold(r)
                    and not _is_auxiliary_pdf_font(_read_run_font_name(r))
                    and not _text_has_pua_symbol(r.text or "")
                ),
                src_runs[-1],
            )
            _copy_run_rpr_style(prose, dst_run)
            if not _paragraph_has_mixed_bold(src_para):
                _set_run_bold(dst_run, False)
        synced += 1
        return synced

    for i, dst_run in enumerate(dst_runs):
        dst_text = dst_run.text or ""
        if _is_symbol_glyph_text(dst_text) or _text_has_pua_symbol(dst_text):
            continue
        if len(src_runs) == 1:
            src_run = src_runs[0]
        else:
            src_run = _pick_style_source_run(src_runs, i)
        if _copy_run_rpr_style(src_run, dst_run):
            synced += 1
    return synced


def _dominant_body_run_rpr(paragraph):
    best_rpr = None
    best_sz = -1
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        if _run_is_superscript_run(run):
            continue
        fname = _read_run_font_name(run)
        if fname and _is_auxiliary_pdf_font(fname):
            continue
        rpr = run._element.find(qn("w:rPr"))
        sz = _rpr_font_size(rpr)
        if sz > best_sz:
            best_sz = sz
            best_rpr = rpr
    if best_rpr is not None:
        return best_rpr
    for run in paragraph.runs:
        if not (run.text or "").strip():
            continue
        if _run_is_superscript_run(run):
            continue
        rpr = run._element.find(qn("w:rPr"))
        sz = _rpr_font_size(rpr)
        if sz > best_sz:
            best_sz = sz
            best_rpr = rpr
    return best_rpr


def _merge_run_size_rpr(dst_run, src_rpr) -> bool:
    if src_rpr is None:
        return False
    try:
        r_el = dst_run._element
        dst_rpr = r_el.find(qn("w:rPr"))
        if dst_rpr is None:
            dst_rpr = OxmlElement("w:rPr")
            r_el.insert(0, dst_rpr)
        for tag in ("w:sz", "w:szCs"):
            src_el = src_rpr.find(qn(tag))
            if src_el is None:
                continue
            old = dst_rpr.find(qn(tag))
            if old is not None:
                dst_rpr.remove(old)
            dst_rpr.append(copy.deepcopy(src_el))
        return True
    except Exception:
        return False


def _merge_run_font_rpr(dst_run, src_rpr) -> bool:
    if src_rpr is None:
        return False
    try:
        r_el = dst_run._element
        dst_rpr = r_el.find(qn("w:rPr"))
        if dst_rpr is None:
            dst_rpr = OxmlElement("w:rPr")
            r_el.insert(0, dst_rpr)
        for tag in ("w:rFonts", "w:sz", "w:szCs"):
            src_el = src_rpr.find(qn(tag))
            if src_el is None:
                continue
            old = dst_rpr.find(qn(tag))
            if old is not None:
                dst_rpr.remove(old)
            dst_rpr.append(copy.deepcopy(src_el))
        for tag in ("w:caps", "w:smallCaps"):
            old = dst_rpr.find(qn(tag))
            if old is not None:
                dst_rpr.remove(old)
        return True
    except Exception:
        return False


def _mirror_paragraph_indents_from_orig(orig_para, trans_para) -> None:
    """Copy alignment + indents from original paragraph to translation line."""
    jc = _paragraph_jc_val(orig_para) or "justify"
    jc_map = {
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "both": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        "distribute": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }
    _set_paragraph_alignment(trans_para, jc_map.get(jc, WD_PARAGRAPH_ALIGNMENT.JUSTIFY))
    orig_ind = _read_ppr_ind_twips(orig_para._element.find(qn("w:pPr")))
    left = int(orig_ind.get("left", 0) or 0)
    right = int(orig_ind.get("right", 0) or 0)
    first_line = int(orig_ind.get("firstLine", 0) or 0)
    hanging = int(orig_ind.get("hanging", 0) or 0)
    if hanging > 0:
        _set_paragraph_indents(trans_para, left=left, right=right, hanging=hanging)
    else:
        _set_paragraph_indents(trans_para, left=left, right=right, first_line=first_line)


def _sync_translation_run_font_from_orig(orig_para, trans_para) -> int:
    """Match translation run bold/size to source paragraph."""
    return _sync_run_styles_from_source(orig_para, trans_para)


def _finalize_newline_translation_pair(orig_para, trans_para) -> None:
    """Ensure translation line mirrors original paragraph layout and font."""
    _mirror_paragraph_indents_from_orig(orig_para, trans_para)
    _sync_translation_run_font_from_orig(orig_para, trans_para)


def _is_pre_body_line(paragraph, src_i: int, abs_start: int) -> bool:
    """Title/author/affiliation block only — strictly before abstract start index."""
    if src_i >= abs_start:
        return False
    text = _paragraph_plain(paragraph)
    if not text:
        return True
    if _is_affiliation_line(paragraph):
        return True
    if len(text) <= 160 and _paragraph_word_count(paragraph) <= 28:
        low = text.lower()
        if any(k in low for k in ("@", "orcid", "author", "corresponding", "email")):
            return True
    return True


def _copy_layout_from_orig_to_translation(
    orig_para,
    trans_para,
    *,
    clear_first_line: bool = False,
) -> None:
    """Mirror orig paragraph layout on translation line; optionally drop first-line indent only."""
    orig_ppr = orig_para._element.find(qn("w:pPr"))
    if orig_ppr is None:
        return
    cloned = copy.deepcopy(orig_ppr)
    spacing = cloned.find(qn("w:spacing"))
    if spacing is not None:
        cloned.remove(spacing)
    if clear_first_line:
        ind = cloned.find(qn("w:ind"))
        if ind is not None:
            ind.attrib.pop(qn("w:firstLine"), None)
    dst_ppr = trans_para._element.find(qn("w:pPr"))
    _copy_paragraph_properties(cloned, dst_ppr)


def _iter_newline_bilingual_pairs(src_doc: docx.Document, dst_doc: docx.Document):
    """Yield (src_index, src_para, orig_dst, trans_dst|None) for stacked bilingual docs."""
    src_items = [
        (i, p)
        for i, p in enumerate(src_doc.paragraphs)
        if _paragraph_plain(p)
    ]
    dst_items = [
        (i, p)
        for i, p in enumerate(dst_doc.paragraphs)
        if _paragraph_plain(p)
    ]

    dst_pos = 0
    for src_pos, (src_i, src_p) in enumerate(src_items):
        src_text = _paragraph_plain(src_p)
        matched = False
        while dst_pos < len(dst_items):
            _, orig_dst = dst_items[dst_pos]
            if _paragraph_texts_match(src_text, _paragraph_plain(orig_dst)):
                matched = True
                break
            dst_pos += 1
        if not matched:
            continue

        next_src_text = None
        if src_pos + 1 < len(src_items):
            next_src_text = _paragraph_plain(src_items[src_pos + 1][1])

        trans_dst = None
        if dst_pos + 1 < len(dst_items):
            _, candidate = dst_items[dst_pos + 1]
            cand_text = _paragraph_plain(candidate)
            if next_src_text and _paragraph_texts_match(cand_text, next_src_text):
                dst_pos += 1
                yield src_i, src_p, orig_dst, None
                continue
            if cand_text.strip() and not _paragraph_texts_match(cand_text, src_text):
                trans_dst = candidate
                dst_pos += 2
                yield src_i, src_p, orig_dst, trans_dst
                continue

        dst_pos += 1
        yield src_i, src_p, orig_dst, trans_dst


def _sync_bilingual_newline_layout(src_doc: docx.Document, dst_doc: docx.Document) -> int:
    """Copy body layout to original+translation paragraph pairs after newline bilingual."""
    synced = 0
    for _src_i, src_p, orig_dst, trans_dst in _iter_newline_bilingual_pairs(src_doc, dst_doc):
        _sync_paragraph_style_only(src_p, orig_dst)
        _preserve_paragraph_layout(src_p, orig_dst)
        if trans_dst is None:
            synced += 1
            continue
        _sync_paragraph_style_only(src_p, trans_dst)
        _copy_layout_from_orig_to_translation(orig_dst, trans_dst, clear_first_line=False)
        _sync_translation_run_font_from_orig(orig_dst, trans_dst)
        synced += 1
    return synced


def _normalize_newline_bilingual_layout_in_doc(
    dst_doc: docx.Document,
    src_doc: docx.Document,
    *,
    pdf_path: Optional[str] = None,
) -> Dict[str, int]:
    """Apply regional left/right/justify to stacked bilingual pairs without touching text."""
    stats: Dict[str, int] = {
        "alignment_normalized": 0,
        "indents_normalized": 0,
        "title_layout_preserved": 0,
    }
    if not _env_bool("PDF_DOCX_NORMALIZE_ALIGN", True):
        return stats

    src_paras = list(src_doc.paragraphs)
    if not src_paras:
        return stats

    _title_end, abs_start, abstract_end, body_start, _sec2 = _detect_layout_regions(src_paras)
    body_profile, abstract_profile = _resolve_layout_profiles(
        src_doc,
        src_paras,
        pdf_path=pdf_path,
    )

    for src_i, src_p, orig_dst, trans_dst in _iter_newline_bilingual_pairs(src_doc, dst_doc):
        if _is_in_table_cell(orig_dst):
            continue
        text = _paragraph_plain(orig_dst)
        if not text.strip():
            continue
        if _is_running_header_line(orig_dst):
            continue

        # Title / author / affiliation: keep source layout (center, hanging indent, etc.)
        if _is_pre_body_line(orig_dst, src_i, abs_start) or _is_pre_body_line(src_p, src_i, abs_start):
            _preserve_paragraph_layout(src_p, orig_dst)
            if trans_dst is not None:
                _copy_layout_from_orig_to_translation(
                    orig_dst,
                    trans_dst,
                    clear_first_line=False,
                )
                _sync_translation_run_font_from_orig(orig_dst, trans_dst)
            stats["title_layout_preserved"] += 1
            continue

        if _is_section_heading_line(orig_dst) and _paragraph_word_count(orig_dst) <= 10:
            _set_paragraph_alignment(orig_dst, WD_PARAGRAPH_ALIGNMENT.LEFT)
            _set_paragraph_indents(
                orig_dst,
                left=int(body_profile["left"]),
                right=0,
                first_line=0,
            )
            stats["indents_normalized"] += 1
            if trans_dst is not None:
                _set_paragraph_alignment(trans_dst, WD_PARAGRAPH_ALIGNMENT.LEFT)
                _set_paragraph_indents(
                    trans_dst,
                    left=int(body_profile["left"]),
                    right=0,
                    first_line=0,
                )
                _sync_translation_run_font_from_orig(orig_dst, trans_dst)
                stats["indents_normalized"] += 1
            continue

        if abs_start <= src_i < abstract_end:
            target = abstract_profile
            is_abstract = True
        elif src_i >= body_start:
            target = body_profile
            is_abstract = False
        else:
            _preserve_paragraph_layout(src_p, orig_dst)
            _set_paragraph_alignment(orig_dst, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            if trans_dst is not None:
                _set_paragraph_alignment(trans_dst, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
                _mirror_paragraph_indents_from_orig(orig_dst, trans_dst)
                _sync_translation_run_font_from_orig(orig_dst, trans_dst)
            stats["title_layout_preserved"] += 1
            continue

        _set_paragraph_alignment(orig_dst, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        stats["alignment_normalized"] += 1
        first_after_heading = False
        if is_abstract:
            _set_paragraph_indents(
                orig_dst,
                left=int(target["left"]),
                right=int(target["right"]),
                first_line=0,
            )
            _clear_first_line_indent(orig_dst)
        elif _paragraph_has_hanging_indent(src_p):
            _preserve_paragraph_layout(src_p, orig_dst)
            _set_paragraph_alignment(orig_dst, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        else:
            first_after_heading = _is_first_paragraph_after_heading(
                src_paras,
                src_i,
                body_start,
            )
            _apply_body_indent_from_source(
                orig_dst,
                src_p,
                body_profile,
                is_first_after_heading=first_after_heading,
            )
        stats["indents_normalized"] += 1

        if trans_dst is not None:
            if is_abstract:
                _set_paragraph_alignment(trans_dst, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
                _set_paragraph_indents(
                    trans_dst,
                    left=int(target["left"]),
                    right=int(target["right"]),
                    first_line=0,
                )
                _clear_first_line_indent(trans_dst)
            elif _paragraph_has_hanging_indent(src_p):
                _preserve_paragraph_layout(src_p, trans_dst)
                _set_paragraph_alignment(trans_dst, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            else:
                _mirror_paragraph_indents_from_orig(orig_dst, trans_dst)
            _sync_translation_run_font_from_orig(orig_dst, trans_dst)
            stats["alignment_normalized"] += 1
            stats["indents_normalized"] += 1

    return stats


def recover_docx_layout(
    source_docx: str,
    translated_docx: str,
    *,
    pdf_path: Optional[str] = None,
    analysis: Optional[Dict[str, Any]] = None,
    bilingual_mode: Optional[str] = None,
    layout_mode: Optional[str] = None,
) -> Dict[str, int]:
    """Restore paragraph properties from pre-translation DOCX (safe, index-based)."""
    if layout_mode is None:
        layout_mode = resolve_pdf_layout_mode(analysis)
    regional = uses_regional_layout(layout_mode)
    bi_mode = normalize_bilingual_mode(bilingual_mode)

    stats: Dict[str, int] = {
        "changed": 0,
        "paragraphs_synced": 0,
        "table_cells_synced": 0,
        "header_footer_synced": 0,
        "mismatched_paragraphs": 0,
    }

    if not os.path.isfile(source_docx) or not os.path.isfile(translated_docx):
        return stats
    if not _env_bool("PDF_DOCX_LAYOUT_SYNC", True):
        return stats

    src = docx.Document(source_docx)
    dst = docx.Document(translated_docx)

    if bi_mode == "newline":
        stats["mismatched_paragraphs"] = abs(len(src.paragraphs) - len(dst.paragraphs))
        stats["paragraphs_synced"] = _sync_bilingual_newline_layout(src, dst)
        stats["table_cells_synced"] = _sync_table_layout_from_source(src, dst)
        stats["header_footer_synced"] = _sync_header_footer_layout(src, dst)
        table_stats = _recover_tables_and_images(dst)
        stats["table_rows_relaxed"] = int(table_stats.get("table_rows_relaxed", 0))
        stats["images_resized"] = int(table_stats.get("images_resized", 0))
        stats["inline_artifacts_stripped"] = _strip_inline_pdf_artifacts(dst)
        post_stats = _normalize_newline_bilingual_layout_in_doc(
            dst,
            src,
            pdf_path=pdf_path,
        )
        stats["post_alignment_normalized"] = int(post_stats.get("alignment_normalized", 0))
        stats["post_indents_normalized"] = int(post_stats.get("indents_normalized", 0))
        stats["title_layout_preserved"] = int(post_stats.get("title_layout_preserved", 0))
        if any(int(stats.get(k, 0) or 0) for k in (
            "paragraphs_synced", "table_cells_synced", "header_footer_synced",
            "table_rows_relaxed", "images_resized", "inline_artifacts_stripped",
            "post_alignment_normalized", "post_indents_normalized", "title_layout_preserved",
        )):
            stats["changed"] = 1
            dst.save(translated_docx)
        return stats

    pair_count = min(len(src.paragraphs), len(dst.paragraphs))
    stats["mismatched_paragraphs"] = abs(len(src.paragraphs) - len(dst.paragraphs))
    sync_fn = (
        _sync_paragraph_style_only if regional else _sync_paragraph_properties_only
    )
    for i in range(pair_count):
        sync_fn(src.paragraphs[i], dst.paragraphs[i])
        stats["paragraphs_synced"] += 1
        stats["run_styles_synced"] = stats.get("run_styles_synced", 0) + _sync_run_styles_from_source(
            src.paragraphs[i],
            dst.paragraphs[i],
        )
        stats["markers_preserved"] = stats.get("markers_preserved", 0) + _preserve_leading_marker_from_source(
            src.paragraphs[i],
            dst.paragraphs[i],
        )

    stats["lines_merged"] = merge_pdf_continuation_paragraphs(dst)
    stats["symbols_repaired"] = ensure_symbol_font_runs_in_doc(dst)
    stats["latin_font_repaired"] = repair_latin_runs_in_symbol_font(dst)
    for para in dst.paragraphs:
        stats["markers_deduped"] = stats.get("markers_deduped", 0) + _dedupe_leading_markers_in_paragraph(para)
    stats["table_cells_synced"] = _sync_table_layout_from_source(src, dst)
    stats["header_footer_synced"] = _sync_header_footer_layout(src, dst)

    table_stats = _recover_tables_and_images(dst)
    stats["table_rows_relaxed"] = int(table_stats.get("table_rows_relaxed", 0))
    stats["images_resized"] = int(table_stats.get("images_resized", 0))

    if bi_mode == "none":
        if regional and not should_preserve_pdf_lines(layout_mode):
            post_stats = normalize_converted_docx_layout_in_doc(
                dst,
                pdf_path=pdf_path,
                src_doc=src,
                layout_mode=layout_mode,
            )
            stats["post_alignment_normalized"] = int(post_stats.get("alignment_normalized", 0))
            stats["post_indents_normalized"] = int(post_stats.get("indents_normalized", 0))

            if pdf_path and os.path.isfile(pdf_path) and _env_bool("PDF_DOCX_PDF_GEOMETRY_PROFILE", True):
                body_profile = _get_fallback_body_profile(dst)
                pdf_profiles = _derive_layout_profiles_from_pdf(pdf_path, dst)
                if pdf_profiles:
                    body_profile = pdf_profiles[0]
                stats["pdf_indent_fixed"] = _fix_indents_from_pdf_geometry(
                    dst,
                    src,
                    pdf_path,
                    body_profile,
                )
        elif pdf_path and os.path.isfile(pdf_path):
            stats["inline_artifacts_stripped"] = _strip_inline_pdf_artifacts(dst)
            stats["fonts_unified"] = unify_pdf_docx_fonts(dst)
        stats["sizes_normalized"] = _normalize_body_run_sizes(dst)

    stats["inline_artifacts_stripped"] = stats.get("inline_artifacts_stripped", 0) or _strip_inline_pdf_artifacts(dst)
    if _env_bool("PDF_DOCX_UNIFY_FONTS", True) and not stats.get("fonts_unified"):
        stats["fonts_unified"] = unify_pdf_docx_fonts(dst)
        stats["sizes_normalized"] = _normalize_body_run_sizes(dst)

    if (
        stats["paragraphs_synced"]
        or stats["table_cells_synced"]
        or stats["header_footer_synced"]
        or stats.get("table_rows_relaxed")
        or stats.get("images_resized")
        or stats.get("post_alignment_normalized")
        or stats.get("post_indents_normalized")
        or stats.get("pdf_indent_fixed")
        or stats.get("pdf_formats_applied")
        or stats.get("fonts_unified")
        or stats.get("inline_artifacts_stripped")
        or stats.get("pdf_bold_synced")
        or stats.get("run_styles_synced")
        or stats.get("markers_preserved")
        or stats.get("lines_merged")
        or stats.get("symbols_repaired")
    ):
        stats["changed"] = 1
        dst.save(translated_docx)

    return stats
